import streamlit as st
import io
import os
import re
import time
import random
import zipfile
from collections import Counter
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT

from utils import (
    TIMECODE_REGEX, RED_COLOR, DEFAULT_NON_SPEAKER_PHRASES,
    clean_and_normalize_text, calculate_duration_sec, find_all_speaker_tags,
    get_speaker_color_config, apply_speaker_styling_to_run,
    is_valid_speaker_name, round_seconds_to_int_minutes, srt_timecode_to_ass,
    generate_actor_docx, generate_vibrant_rgb_colors_excluding, DEFAULT_FIXED_SPEAKER_COLORS,
    get_paragraph_text_with_html, clean_file_name_for_output
)

def apply_html_and_phonetic_to_paragraph(paragraph, text, enable_phonetic):
    working_text = text
    if enable_phonetic and 'custom_phonetics' in st.session_state:
        for eng_word, pho_word in st.session_state['custom_phonetics'].items():
            pattern = r'\b' + re.escape(eng_word) + r'\b'
            working_text = re.sub(pattern, f"{eng_word} ({pho_word})", working_text, flags=re.IGNORECASE)

    tokens = re.split(r'(<i>.*?</i>|<b>.*?</b>|<u>.*?</u>)', working_text, flags=re.IGNORECASE)
    for token in tokens:
        if not token: continue
        r = paragraph.add_run()
        low = token.lower()
        if low.startswith('<i>') and low.endswith('</i>'):
            r.text = token[3:-4]; r.italic = True
        elif low.startswith('<b>') and low.endswith('</b>'):
            r.text = token[3:-4]; r.bold = True
        elif low.startswith('<u>') and low.endswith('</u>'):
            r.text = token[3:-4]; r.underline = True
        else:
            r.text = token
        r.font.name = 'Times New Roman'

def process_tab1_docx(uploaded_file, file_name_without_ext, enable_colors, enable_phonetic, enable_cast, font_size_pt=12):
    speaker_color_map = {}
    fixed_colors = st.session_state.get('fixed_speaker_colors', DEFAULT_FIXED_SPEAKER_COLORS)
    fixed_rgb_set = {tuple(v.get("text_color")) for v in fixed_colors.values() if isinstance(v, dict) and v.get("text_color")}
    
    available_rgb_tuples = generate_vibrant_rgb_colors_excluding(fixed_rgb_set, 200)
    random.shuffle(available_rgb_tuples)

    stats_counter = Counter()
    seen_speakers_first_time = set()
    actor_dialogue_map = {}
    qc_warnings = []        
    
    custom_speakers = st.session_state.get('custom_speakers', set())
    custom_non_speakers = st.session_state.get('custom_non_speakers', set())

    original_document = Document(io.BytesIO(uploaded_file.getvalue()))
    raw_paragraphs = [p for p in original_document.paragraphs]
    
    processed_strings = []
    for p in raw_paragraphs:
        txt = get_paragraph_text_with_html(p).strip()
        if txt: processed_strings.append(txt)
    
    document = Document()
    title_text = file_name_without_ext.upper()
    title_paragraph = document.add_paragraph(title_text)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.runs[0].font.name = 'Times New Roman'; title_paragraph.runs[0].font.size = Pt(20); title_paragraph.runs[0].bold = True
    
    document.add_paragraph()
    progress_bar = st.progress(0); status_text = st.empty()
    total_paras = len(processed_strings)
    
    for idx, text in enumerate(processed_strings):
        if total_paras > 0 and idx % max(1, total_paras // 10) == 0:
            progress_bar.progress(int((idx / total_paras) * 100))
            status_text.text(f"Đang phân tích kịch bản {idx}/{total_paras}...")

        if TIMECODE_REGEX.match(text):
            p_tc = document.add_paragraph(text)
            p_tc.runs[0].font.bold = True; p_tc.runs[0].font.name = 'Times New Roman'; p_tc.runs[0].font.size = Pt(font_size_pt)
            p_tc.paragraph_format.space_before = Pt(0); p_tc.paragraph_format.space_after = Pt(0)
        else:
            p_line = document.add_paragraph()
            apply_html_and_phonetic_to_paragraph(p_line, text, enable_phonetic)
            p_line.paragraph_format.space_before = Pt(0); p_line.paragraph_format.space_after = Pt(4)

    progress_bar.progress(100); status_text.text("Xử lý hoàn tất!")
    time.sleep(0.3); progress_bar.empty(); status_text.empty()
    
    docx_file = io.BytesIO(); document.save(docx_file); docx_file.seek(0)
    return docx_file

def render_tab1(enable_colors, enable_phonetic, enable_cast):
    st.subheader("🎬 XỬ LÝ KỊCH BẢN GỐC (SUBTITLE & DIALOGUE)")
    st.markdown("Tải file kịch bản Tiếng Anh (.docx) để tự động phân vai, gán màu nhân vật, chèn phiên âm và xuất file Word chuẩn 12pt.")
    
    up_key = f"tab1_file_{st.session_state.get('uploader_key', 0)}"
    uploaded_file = st.file_uploader("Tải file Kịch bản gốc (.docx):", type=['docx'], key=up_key)
    
    if uploaded_file is not None:
        file_name_without_ext = os.path.splitext(uploaded_file.name)[0]
        
        if st.button("🚀 1. BẮT ĐẦU ĐỊNH DẠNG TỰ ĐỘNG", type="primary", use_container_width=True):
            try:
                with st.spinner("Đang định dạng kịch bản..."):
                    out_docx = process_tab1_docx(uploaded_file, file_name_without_ext, enable_colors, enable_phonetic, enable_cast)
                    st.session_state['processed_docx'] = out_docx
                    st.session_state['processed_name'] = file_name_without_ext
                st.success("✅ Đã xử lý định dạng kịch bản thành công!")
            except Exception as e:
                st.error(f"Đã có lỗi xảy ra: {e}")
                
        if 'processed_docx' in st.session_state:
            st.markdown("---")
            st.markdown("### 📥 Tải Xuất Dữ Liệu Sau Khi Xử Lý")
            fn = clean_file_name_for_output(st.session_state.get('processed_name', 'KichBan'), tag="_edit", ext=".docx")
            st.download_button(
                label=f"⬇️ TẢI FILE WORD ĐÃ ĐỊNH DẠNG ({fn})",
                data=st.session_state['processed_docx'],
                file_name=fn,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True
            )
