import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_COLOR_INDEX
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import io
import os
import re
import random
from collections import Counter
import time
import pandas as pd
import json
from gtts import gTTS
import zipfile
from datetime import datetime

# ==========================================
# 1. CẤU HÌNH TRANG CHỦ STREAMLIT
# ==========================================
st.set_page_config(
    page_title="ScriptPro Enterprise - Subtitle & Script Editor",
    page_icon="🎬",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==========================================
# 2. DATABASE MẶC ĐỊNH & HÀM ĐỌC/GHI JSON
# ==========================================
NON_SPEAKER_DB_FILE = "custom_non_speakers.json"
SPEAKER_DB_FILE = "custom_speakers.json"
PHONETIC_DB_FILE = "custom_phonetics.json"
CAST_DB_FILE = "custom_cast_mapping.json"
TRACKER_DB_FILE = "dubbing_tracker.json"
RATES_DB_FILE = "payroll_rates.json"
PRONOUN_REL_DB_FILE = "custom_pronoun_relationships.json"
SPEAKER_COLOR_DB_FILE = "custom_speaker_colors.json"

DEFAULT_CAST_MAPPING = {
    "BRI": "TRÚC", "CHASE": "THIỆN", "PRESTON": "KHÁNH", "SCOTT": "THÔNG",
    "KEELEY": "TÚ", "DAKA": "QUANG", "STEPHEN": "QUANG", "VINCE": "QUANG",
    "JOSH": "THÔNG", "YOMI": "TÚ", "LARRY": "A.TRUNG", "CHLOE": "TÚ",
    "STEVEN": "QUANG (HOẶC THÔNG)", "LOGAN": "THÔNG", "NATHAN": "A.TRUNG",
    "RILEY": "PHỤNG", "ZHONG": "THÔNG", "MICHELLE (YOUTUBER)": "TÚ",
    "BEN AZELART": "QUANG", "ZONG": "THÔNG", "BA BRI": "THÔNG", "TYLER": "PHÁT",
    "COBY": "THIỆN", "CORY": "KHÁNH", "SPARKY": "A.TRUNG", "GARRETT": "C.DŨNG",
    "CODY": "CƯỜNG", "BETHANY": "TRÚC", "KRISTIN": "PHỤNG", "AMY": "TÚ",
    "ALLISON": "TÚ", "AUBREY": "PHỤNG", "TY JACKSON": "THÔNG", "HLV": "CƯỜNG",
    "JACKSON OLSON": "THÔNG", "DANNY": "QUANG", "KYLE": "QUANG", "BILL": "HITA", "TIM": "HITA"
}

# 🎨 BẢNG MÀU CHỮ & HIGHLIGHT CỐ ĐỊNH CHÍNH XÁC (CÓ HỖ TRỢ CẢ MÀU CHỮ LẪN HIGHLIGHT NỀN)
DEFAULT_FIXED_SPEAKER_COLORS = {
    "ALL": {"text_color": (255, 0, 0), "highlight_color": (255, 255, 0)},
    "BEN AZELART": {"text_color": (21, 96, 130), "highlight_color": None},
    "BETHANY": {"text_color": (255, 192, 0), "highlight_color": None},
    "BRI": {"text_color": (116, 166, 123), "highlight_color": None},
    "CALEB": {"text_color": (116, 56, 25), "highlight_color": None},
    "CHASE": {"text_color": (255, 0, 255), "highlight_color": None},
    "CHUNKZ": {"text_color": (21, 96, 130), "highlight_color": None},
    "COBY": {"text_color": (255, 0, 255), "highlight_color": None},
    "CODY": {"text_color": (233, 113, 50), "highlight_color": None},
    "CORY": {"text_color": (255, 0, 0), "highlight_color": None},
    "DAKA": {"text_color": (110, 196, 229), "highlight_color": None},
    "GARRETT": {"text_color": (243, 243, 243), "highlight_color": None},
    "JOSH": {"text_color": (160, 43, 147), "highlight_color": None},
    "KEELEY": {"text_color": (255, 192, 0), "highlight_color": None},
    "LARRY": {"text_color": (243, 243, 243), "highlight_color": None},
    "LOGAN": {"text_color": (216, 170, 211), "highlight_color": None},
    "NATHAN": {"text_color": (21, 96, 130), "highlight_color": None},
    "NICK": {"text_color": (255, 0, 0), "highlight_color": None},
    "PRESTON": {"text_color": (255, 0, 0), "highlight_color": None},
    "RILEY": {"text_color": (116, 166, 123), "highlight_color": None},
    "SCOTT": {"text_color": (233, 113, 50), "highlight_color": None},
    "SPARKY": {"text_color": (116, 56, 25), "highlight_color": None},
    "STEPHEN": {"text_color": (0, 51, 204), "highlight_color": None},
    "STEVEN": {"text_color": (0, 255, 255), "highlight_color": None},
    "TYLER": {"text_color": (160, 43, 147), "highlight_color": None},
    "YOMI": {"text_color": (0, 153, 153), "highlight_color": None}
}

DEFAULT_SOUTH_VIETNAM_PHONETICS = {
    "MCDONALD": "Mắc-đô-nồ", "MCDONALD'S": "Mắc-đô-nồ", "LONDON": "Luân Đôn", "TNT": "Ti-èn-ti",
    "NOOB": "Núp", "GOLEM": "Gô-lùm", "GOOGLE": "Gú-gồ", "FACEBOOK": "Phây-sbúc",
    "KFC": "Ke-ép-xi", "STARBUCKS": "Xì-ta-bắc", "APPLE": "Ép-pồ", "BURGER": "Bơ-gơ",
    "PARIS": "Ba Lê", "WASHINGTON": "Hoa Thịnh Đốn", "CALIFORNIA": "Cả-li", "PRO": "Prồ",
    "LAG": "Lác", "BUFF": "Bớp", "VIP": "Vi-ai-pi", "FBI": "Ép-bi-ai"
}

VN_SYLLABLES = {
    "a", "ai", "an", "ang", "anh", "ao", "ap", "at", "au", "ay", "ba", "bac", "bai", "ban", "bang", "bao", 
    "bat", "bay", "be", "ben", "beng", "beo", "bi", "bic", "bien", "biet", "binh", "bo", "boc", "boi", "bon", 
    "bong", "bot", "bu", "bua", "bui", "bun", "buoc", "buon", "ca", "cac", "cai", "cam", "can", "cang", "cao", 
    "cap", "cat", "cau", "cay", "cha", "chai", "cham", "chan", "chang", "chao", "chap", "chat", "chau", "chay", 
    "che", "chen", "cheo", "chi", "chia", "chiem", "chieu", "chin", "chinh", "cho", "choc", "choi", "chon", 
    "chong", "chu", "chua", "chuan", "chuc", "chui", "chum", "chun", "chuoc", "chuon", "chura", "chuong", "co", 
    "coc", "coi", "con", "cong", "cot", "cu", "cua", "cuc", "cui", "cum", "cung", "cuoc", "cuoi", "cuon", 
    "cuong", "cuot", "da", "dac", "dai", "dam", "dan", "dang", "dao", "dap", "dat", "dau", "day", "de", "den", 
    "deo", "di", "dia", "diem", "dien", "diet", "dieu", "dinh", "do", "doc", "doi", "don", "dong", "dot", "du", 
    "dua", "duc", "dui", "dum", "dung", "duoc", "duoi", "duon", "duong", "em", "ga", "gac", "gai", "gam", "gan", 
    "gang", "gao", "gap", "gat", "gau", "gay", "ge", "gen", "gi", "gia", "giac", "giai", "giam", "gian", "giang", 
    "giao", "giap", "giat", "giau", "giay", "gio", "gioc", "gioi", "gion", "giong", "giu", "giua", "go", "goc", 
    "goi", "gon", "gong", "gu", "gua", "guc", "gui", "gum", "gung", "ha", "hac", "hai", "ham", "han", "hang", 
    "hao", "hap", "hat", "hau", "hay", "he", "hen", "heo", "hi", "hien", "hiet", "hieu", "hinh", "ho", "hoc", 
    "hoi", "hon", "hong", "hot", "hu", "hua", "huc", "hui", "hum", "hung", "huoc", "huong", "huot", "i", "ic", 
    "it", "khac", "khai", "kham", "khan", "khang", "khao", "khap", "khat", "khau", "khay", "khe", "khen", "kheo", 
    "khi", "khia", "khien", "khieu", "kho", "khoc", "khoi", "khon", "khong", "khot", "khu", "khua", "khuc", 
    "khui", "khum", "khung", "khuon", "khura", "khuong", "la", "lac", "lai", "lam", "lan", "lang", "lao", "lap", 
    "lat", "lau", "lay", "le", "len", "leo", "li", "lia", "liem", "lien", "liet", "lieu", "linh", "lo", "loc", 
    "loi", "lon", "long", "lot", "lu", "lua", "luc", "lui", "lum", "lung", "luoc", "luoi", "luon", "luong", 
    "luot", "ma", "mac", "mai", "mam", "man", "mang", "mao", "map", "mat", "mau", "may", "me", "men", "meo", 
    "mi", "mia", "mien", "mieu", "minh", "mo", "moc", "moi", "mon", "mong", "mot", "mu", "mua", "muc", "mui", 
    "mum", "mung", "muoc", "muoi", "muon", "muong", "na", "nac", "nai", "nam", "nan", "nang", "nao", "nap", 
    "nat", "nau", "nay", "ne", "nen", "neo", "ni", "nia", "niem", "nien", "nieu", "ninh", "no", "noc", 
    "noi", "non", "nong", "not", "nu", "nua", "nuc", "nui", "num", "nung", "nuoc", "nuoi", "nuon", "nuong", 
    "nga", "ngac", "ngai", "ngam", "ngan", "ngang", "ngao", "ngap", "ngat", "ngau", "ngay", "nge", "ngen", 
    "nghe", "nghen", "ngheo", "nghi", "nghia", "nghiem", "nghien", "nghiet", "nghieu", "nghinh", "ngo", "ngoc", 
    "ngoi", "ngon", "ngong", "ngot", "ngu", "ngua", "nguc", "ngui", "ngum", "ngung", "nguoc", "nguoi", "nguon", 
    "nguong", "nha", "nhac", "nhai", "nham", "nhan", "nhang", "nhao", "nhap", "nhat", "nhau", "nhay", "nhe", 
    "nhen", "nheo", "nhi", "nhia", "nhiem", "nhien", "nhiet", "nhieu", "nhinh", "nho", "nhoc", "nhoi", "nhon", 
    "nhong", "nhot", "nhu", "nhua", "nhuc", "nhui", "nhum", "nhung", "nhuoc", "nhuoi", "nhuon", "nhuong", "oa", 
    "oai", "oam", "oan", "oang", "oat", "oay", "oe", "oen", "oi", "om", "on", "ong", "ot", "pa", "pha", "phac", 
    "phai", "pham", "phan", "phang", "phao", "phap", "phat", "phau", "phay", "phe", "phen", "pheo", "phi", 
    "phia", "phiem", "phien", "phiet", "phieu", "phinh", "pho", "phoc", "phoi", "phon", "phong", "phot", "phu", 
    "phua", "phuc", "phui", "phum", "phung", "phuoc", "phuong", "phuot", "qua", "quac", "quai", "quam", "quan", 
    "quang", "quao", "quap", "quat", "quay", "que", "quen", "queo", "qui", "quie", "quien", "quieu", "quinh", 
    "quo", "quoc", "quoi", "quon", "quong", "ra", "rac", "rai", "ram", "ran", "rang", "rao", "rap", "rat", 
    "rau", "ray", "re", "ren", "reo", "ri", "ria", "rim", "rin", "rinh", "ro", "roc", "roi", "ron", "rong", 
    "rot", "ru", "rua", "ruc", "rui", "rum", "rung", "ruoc", "ruoi", "ruon", "ruong", "sa", "sac", "sai", 
    "sam", "san", "sang", "sao", "sap", "sat", "sau", "say", "se", "sen", "seo", "si", "sia", "siem", "sien", 
    "siet", "sieu", "sinh", "so", "soc", "soi", "son", "song", "sot", "su", "sua", "suc", "sui", "sum", 
    "sung", "suoc", "suoi", "suon", "suong", "ta", "tac", "tai", "tam", "tan", "tang", "tao", "tap", "tat", 
    "tau", "tay", "te", "ten", "teo", "tha", "thac", "thai", "tham", "than", "thang", "thao", "thap", "that", 
    "thau", "thay", "the", "then", "theo", "thi", "thia", "thiem", "thien", "thiet", "thieu", "thinh", "tho", 
    "thoc", "thoi", "thon", "thong", "thot", "thu", "thua", "thuc", "thui", "thum", "thung", "thuoc", "thuoi", 
    "thuon", "thuong", "thuot", "to", "toc", "toi", "ton", "tong", "tot", "tra", "trac", "trai", "tram", "tran", 
    "trang", "trao", "trap", "trat", "trau", "tray", "tre", "tren", "treo", "tri", "tria", "triem", "trien", 
    "triet", "trieu", "trinh", "tro", "troc", "troi", "tron", "trong", "trot", "tru", "trua", "truc", "trui", 
    "trum", "trung", "truoc", "truoi", "truon", "truong", "tu", "tua", "tuc", "tui", "tum", "tung", "tuoc", 
    "tuoi", "tuon", "tuong", "tuot", "va", "vac", "vai", "vam", "van", "vang", "vao", "vap", "vat", "vau", 
    "vay", "ve", "ven", "veo", "vi", "via", "viem", "vien", "viet", "vieu", "vinh", "vo", "voc", "voi", "von", 
    "vong", "vot", "vu", "vua", "vuc", "vui", "vum", "vung", "vuoc", "vuoi", "vuon", "vuong", "xa", "xac", 
    "xai", "xam", "xan", "xang", "xao", "xap", "xat", "xau", "xay", "xe", "xen", "xeo", "xi", "xia", "xiem", 
    "xien", "xiet", "xieu", "xinh", "xo", "xoc", "xoi", "xon", "xong", "xot", "xu", "xua", "xuc", "xui", 
    "xum", "xung", "xuoc", "xuoi", "xuon", "xuong", "y", "eu", "yeu"
}

DEFAULT_NON_SPEAKER_PHRASES = {
    "AND REMEMBER", "OFFICIAL DISTANCE", "GOOD NEWS FOR THEIR TEAMMATES", 
    "LL BE HONEST", "FIRST AND FOREMOST", "I SAID", "THE ONLY THING LEFT TO SETTLE", 
    "QUESTION IS", "FINALISTS", "WHISPERS", "SRT CONVERSION", 
    "WILL RED THRIVE OR WILL RED BE DEAD", "BUT REMEMBER", "THE RESULTS ARE IN", 
    "WE CHALLENGED", "I THINK", "IN THEIR DEFENSE", "THE PEAK OF HIS LIFE WAS DOING THE SPACETHING",
    "THE ROCKETS ARE BIGGER", "THE DISTANCE SHOULD BE FURTHER", "GET CRAFTY", "THAT WAS SO SICK",
    "OUT OF 100 CONTESTANTS", "THE FIRST ROUND IS BRUTAL", "YOU KNOW WHICH END GOES",
    "THE GAME IS ON", "THAT'S A GOOD THROW", "HE'S GOING FOR IT", "WE GOT THIS",
    "LAUNCH", "OH NO", "OH", "AH", "YEP", "WAIT", "YEAH", "WOO", "OKAY", "YES", "I ANH", "O BRI", "NG", "THE ONLY PROBLEM", "NOTE", "WARNING", "THINGS", "AND ON THE WAY WE CAME ACROSS THIS", 
    "THIS IS THE HIGHEST SWING IN EUROPE", "AND I SWEAR", "WHICH MEANT", "THE ONLY THING IS", 
    "HERE WE GO", "NEXT UP", "STEP 1", "STEP 2", "STEP 3", "AND STEP 3", "FIRST UP", 
    "SO THE QUESTION IS", "I WAS GROWING UP", "YOU MIGHT BE WONDERING", "UPDATE", 
    "NASHVILLE TO MIAMI", "ALL I KNOW IS", "UNLIKE JUDY", "THE GOOD NEWS IS", 
    "AER LINGUS SEAT", "THE TRUE TEST IS", "JUST AS I SUSPECTED", "LIKE I SAID", 
    "STAR REVIEW AND SAID", "I TOLD THEM ALL", "AND BEST OF ALL", "THE POINT IS", 
    "AMERICANS", "I WAS THINKING", "AND THEY GO", "FIRST OF ALL", "SECOND", 
    "ARE YOU LIKE", "AS A REMINDER", "ROUND 2", "ROUND 1", "ROUND 3", "ROUND 4", 
    "ROUND 5", "WELCOME TO ROUND 3", "THE QUESTION IS", "QUICK REMINDER", 
    "IN 2ND PLACE", "COMING UP", "FIRST STOP", "NEXT STEP", "AND THAT MEANS", 
    "HASHTAG", "SO TO BE CLEAR", "YOUR SECOND WORD", "WELCOME TO ROUND 6", 
    "BATTLE FINALE TIME", "NUMBER 1", "NUMBER 2", "BUT THE TRUTH IS", 
    "SCORE TO BEAT", "AND YOUR WINNER", "\"CRAFTY\" AND \"BETCHA\". COMING UP", 
    "NEXT ONE", "KEEP IN MIND", "AND IT SAYS", "YOU COULD SAY", "WELCOME TO ROUND 2", 
    "AND THE BEST PART", "ONTO ROUND 2", "THE RIDE WE CHOSE", "GOOD NEWS IS", 
    "BAD NEWS", "GOOD NEWS", "HE THOUGHT", "3 TEAMS REMAIN", "QUICK UPDATE", "DISTORTED", "MY FIRST QUESTION IS", "AND THE BEST PART IS", "BUT AS GORDON RAMSAY MIGHT SAY", "BUT THE GOOD NEWS IS", "LET ME JUST SAY", "BUT THE BEST PART", "I WILL SAY THOUGH", "LL SAY IS, UPDATE", "LL SAY IS", "UPDATE", "TO ALL OF YOU WATCHING WITH ME RIGHT NOW", "THIS IS THE LIFE", "I HAVE A QUESTION", "I WILL BE HONEST", "SO I CAN UPDATE MY INSTAGRAM BIO TO SAY"
}

VN_SELF_PRONOUNS = ["tui", "tôi", "mình", "tao", "ta", "em", "anh", "chị", "cháu", "con", "tại hạ", "bản thân"]
VN_TARGET_PRONOUNS = ["ông", "bạn", "mày", "anh", "chị", "chú", "bác", "cậu", "bà", "cưng", "em", "ní", "mấy ní", "sư huynh", "huynh", "đệ"]

def hex_to_rgb(hex_str):
    if not hex_str or not isinstance(hex_str, str) or not hex_str.startswith('#'):
        return None
    h = hex_str.lstrip('#')
    if len(h) != 6: return None
    try:
        return (int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    except ValueError:
        return None

def normalize_speaker_color_db(db_data):
    normalized = {}
    if not isinstance(db_data, dict): return normalized
    for spk, val in db_data.items():
        spk_clean = str(spk).upper().strip()
        if isinstance(val, (list, tuple)) and len(val) == 3:
            normalized[spk_clean] = {"text_color": tuple(val), "highlight_color": None}
        elif isinstance(val, dict):
            tc = tuple(val.get("text_color")) if val.get("text_color") else None
            hc = tuple(val.get("highlight_color")) if val.get("highlight_color") else None
            normalized[spk_clean] = {"text_color": tc, "highlight_color": hc}
    return normalized

def load_json_db(filepath, default_data=None):
    if os.path.exists(filepath):
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
                if filepath == SPEAKER_COLOR_DB_FILE:
                    return normalize_speaker_color_db(data)
                elif isinstance(data, dict) and isinstance(default_data, dict):
                    res = {}
                    for k, v in data.items():
                        if isinstance(v, (list, tuple)) and len(v) == 3: res[k] = tuple(v)
                        else: res[k] = v
                    return res
                return set(data) if isinstance(data, list) and isinstance(default_data, set) else data
        except Exception: pass
    return default_data if default_data is not None else (set() if not isinstance(default_data, dict) else {})

def save_json_db(filepath, data_container):
    try:
        with open(filepath, "w", encoding="utf-8") as f:
            if isinstance(data_container, set): json.dump(list(data_container), f, ensure_ascii=False, indent=2)
            elif isinstance(data_container, dict):
                out_dict = {}
                for k, v in data_container.items():
                    if isinstance(v, (tuple, list)): out_dict[k] = list(v)
                    elif isinstance(v, dict):
                        sub_dict = {}
                        for sub_k, sub_v in v.items():
                            if isinstance(sub_v, (tuple, list)): sub_dict[sub_k] = list(sub_v)
                            else: sub_dict[sub_k] = sub_v
                        out_dict[k] = sub_dict
                    else: out_dict[k] = v
                json.dump(out_dict, f, ensure_ascii=False, indent=2)
            else: json.dump(data_container, f, ensure_ascii=False, indent=2)
    except Exception as e: st.error(f"Không thể lưu vào Database: {e}")

# Khởi tạo Session State
if 'uploader_key' not in st.session_state: st.session_state['uploader_key'] = 0
if 'resync_uploader_key' not in st.session_state: st.session_state['resync_uploader_key'] = 0
if 'spk_input_key' not in st.session_state: st.session_state['spk_input_key'] = 0
if 'ns_input_key' not in st.session_state: st.session_state['ns_input_key'] = 0
if 'pho_input_key' not in st.session_state: st.session_state['pho_input_key'] = 0
if 'cast_input_key' not in st.session_state: st.session_state['cast_input_key'] = 0
if 'pronoun_input_key' not in st.session_state: st.session_state['pronoun_input_key'] = 0
if 'color_input_key' not in st.session_state: st.session_state['color_input_key'] = 0
if 'textarea_clean_output' not in st.session_state: st.session_state['textarea_clean_output'] = ""

if 'custom_non_speakers' not in st.session_state: st.session_state['custom_non_speakers'] = load_json_db(NON_SPEAKER_DB_FILE, set())
if 'custom_speakers' not in st.session_state: st.session_state['custom_speakers'] = load_json_db(SPEAKER_DB_FILE, set())

if 'custom_phonetics' not in st.session_state:
    loaded_pho = load_json_db(PHONETIC_DB_FILE, DEFAULT_SOUTH_VIETNAM_PHONETICS)
    merged_pho = {**DEFAULT_SOUTH_VIETNAM_PHONETICS, **loaded_pho}
    st.session_state['custom_phonetics'] = merged_pho

if 'custom_cast_mapping' not in st.session_state:
    loaded_cast = load_json_db(CAST_DB_FILE, DEFAULT_CAST_MAPPING)
    merged_cast = {**DEFAULT_CAST_MAPPING, **loaded_cast}
    st.session_state['custom_cast_mapping'] = merged_cast

if 'fixed_speaker_colors' not in st.session_state:
    loaded_colors = load_json_db(SPEAKER_COLOR_DB_FILE, DEFAULT_FIXED_SPEAKER_COLORS)
    merged_colors = {**normalize_speaker_color_db(DEFAULT_FIXED_SPEAKER_COLORS), **normalize_speaker_color_db(loaded_colors)}
    st.session_state['fixed_speaker_colors'] = merged_colors

if 'custom_pronoun_rel' not in st.session_state:
    default_pronouns = {
        "TYLER|BILL": {"self": "tui", "target": "ông"},
        "CORY|EASTON": {"self": "tui", "target": "ông"},
        "COBY|COACH RAC": {"self": "tui", "target": "ông"}
    }
    st.session_state['custom_pronoun_rel'] = load_json_db(PRONOUN_REL_DB_FILE, default_pronouns)

if 'dubbing_tracker' not in st.session_state: st.session_state['dubbing_tracker'] = load_json_db(TRACKER_DB_FILE, [])

if 'payroll_rates' not in st.session_state:
    default_rates = {"mode": "minute", "unit_rate": 30000}
    st.session_state['payroll_rates'] = load_json_db(RATES_DB_FILE, default_rates)

NON_SPEAKER_PHRASES = DEFAULT_NON_SPEAKER_PHRASES.union(st.session_state['custom_non_speakers'])
TIMECODE_REGEX = re.compile(r"^\d{2}:\d{2}:\d{2},\d{3}\s+-->\s+\d{2}:\d{2}:\d{2},\d{3}$")
ENGLISH_WORD_REGEX = re.compile(r"\b[A-Za-z][A-Za-z0-9'-]*\b")
RED_COLOR = RGBColor(255, 0, 0)

# ==========================================
# 3. UNIFIED SIDEBAR (CONTROL PANEL)
# ==========================================
st.sidebar.markdown("### ⚡ Control Panel")

ui_theme_choice = st.sidebar.radio(
    "Lựa chọn Skin hiển thị:",
    options=["Mai Han Standard (Mặc định)", "Enterprise Pro (Tối ưu tương phản)"],
    index=0,
    help="Chế độ 'Mai Han Standard' giữ nguyên 100% giao diện truyền thống. Chế độ 'Enterprise Pro' mang lại phong cách Studio hiện đại, rõ nét và tương phản cao."
)

if st.sidebar.button("🔄 Reset phiên làm việc", use_container_width=True, type="primary"):
    for key in ['processed_docx', 'processed_ass', 'processed_srt', 'actor_zip', 'r_processed_docx', 'r_processed_ass', 'r_processed_srt', 'r_actor_zip']:
        if key in st.session_state:
            del st.session_state[key]
    st.session_state['uploader_key'] += 1
    st.session_state['resync_uploader_key'] += 1
    st.rerun()

st.sidebar.markdown("---")
st.sidebar.markdown("#### 🎛️ Bật/Tắt Tính năng")
enable_colors = st.sidebar.toggle("🌈 Tô màu nhân vật", value=True)
enable_phonetic = st.sidebar.toggle("🗣️ Phiên âm giọng Nam", value=True, help="Tự động chèn phiên âm giọng Nam trước từ Tiếng Anh (ngoặc đơn + tô màu vàng)")
enable_cast = st.sidebar.toggle("🎭 Phân vai lồng tiếng", value=True, help="Hiển thị thông tin diễn viên lồng tiếng ở đầu trang và lần xuất hiện đầu tiên của nhân vật")

st.sidebar.markdown("---")
st.sidebar.markdown("#### 💾 Database Quản Lý Cụm Từ")

with st.sidebar.expander("🎭 Database Người nói (Whitelist)", expanded=False):
    manual_spk_input = st.text_area("Nhập thủ công:", height=80, key=f"spk_manual_{st.session_state['spk_input_key']}")
    upload_spk_file = st.file_uploader("Tải file (.txt, .docx, .xlsx)", type=['txt', 'docx', 'xlsx'], key=f"spk_uploader_{st.session_state['spk_input_key']}")
    
    if st.button("Lưu Người Nói", use_container_width=True):
        new_spks = set()
        if manual_spk_input:
            parts = re.split(r'[,\n]', manual_spk_input)
            new_spks.update([p.strip() for p in parts if p.strip()])
        if upload_spk_file:
            new_spks.update(extract_phrases_from_file(upload_spk_file, upload_spk_file.name))
            
        if new_spks:
            st.session_state['custom_speakers'].update(new_spks)
            save_json_db(SPEAKER_DB_FILE, st.session_state['custom_speakers'])
            st.session_state['spk_input_key'] += 1
            st.success(f"✅ Đã lưu {len(new_spks)} người nói!"); time.sleep(1); st.rerun()

with st.sidebar.expander("🚫 Database Từ nhiễu (Non-speaker)", expanded=False):
    manual_input = st.text_area("Nhập thủ công:", height=80, key=f"ns_manual_{st.session_state['ns_input_key']}")
    upload_non_speaker = st.file_uploader("Tải file (.txt, .docx, .xlsx)", type=['txt', 'docx', 'xlsx'], key=f"ns_uploader_{st.session_state['ns_input_key']}")
    
    if st.button("Lưu Từ Nhiễu", use_container_width=True):
        new_phrases = set()
        if manual_input:
            parts = re.split(r'[,\n]', manual_input)
            new_phrases.update([p.strip().upper() for p in parts if p.strip()])
        if upload_non_speaker:
            new_phrases.update([p.upper() for p in extract_phrases_from_file(upload_non_speaker, upload_non_speaker.name)])
            
        if new_phrases:
            st.session_state['custom_non_speakers'].update(new_phrases)
            save_json_db(NON_SPEAKER_DB_FILE, st.session_state['custom_non_speakers'])
            st.session_state['ns_input_key'] += 1
            st.success(f"✅ Đã lưu {len(new_phrases)} từ nhiễu!"); time.sleep(1); st.rerun()

# ==========================================
# 4. DYNAMIC CSS INJECTION THEO CHẾ ĐỘ SKIN
# ==========================================
if "Enterprise Pro" in ui_theme_choice:
    st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');
        
        html, body, [class*="css"] {
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
            color: #0F172A;
        }
        
        .hero-container {
            background: linear-gradient(135deg, #0F172A 0%, #1E293B 100%);
            padding: 2.2rem 2rem;
            border-radius: 14px;
            color: #FFFFFF;
            margin-bottom: 1.8rem;
            border-left: 6px solid #38BDF8;
            box-shadow: 0 10px 25px -5px rgba(15, 23, 42, 0.2);
        }
        .hero-title { font-size: 2.3rem; font-weight: 800; margin: 0; letter-spacing: -0.02em; color: #FFFFFF; }
        .hero-subtitle { font-size: 1.05rem; color: #94A3B8; margin-top: 0.4rem; font-weight: 400; }
        .badge-pro {
            background-color: #0284C7; color: #FFFFFF; padding: 4px 12px;
            border-radius: 6px; font-size: 0.75rem; font-weight: 700;
            text-transform: uppercase; letter-spacing: 0.05em; display: inline-block; margin-bottom: 0.6rem;
        }
        
        [data-testid="stVerticalBlockBorderWrapper"] {
            background-color: #FFFFFF !important;
            border: 1px solid #CBD5E1 !important;
            border-radius: 12px !important;
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.05) !important;
        }
        
        .metric-card {
            background: #F8FAFC; border: 1px solid #E2E8F0; border-radius: 10px;
            padding: 1.25rem; box-shadow: 0 1px 3px rgba(0,0,0,0.04);
            transition: all 0.2s ease-in-out;
        }
        .metric-card:hover { border-color: #0284C7; transform: translateY(-2px); box-shadow: 0 8px 15px rgba(0,0,0,0.06); }
        .metric-label { font-size: 0.8rem; color: #475569; font-weight: 700; text-transform: uppercase; letter-spacing: 0.04em; }
        .metric-value { font-size: 1.8rem; font-weight: 800; color: #0F172A; margin-top: 0.2rem; }
        
        .stTabs [data-baseweb="tab-list"] {
            gap: 8px; border-bottom: 2px solid #E2E8F0; padding-bottom: 2px;
        }
        .stTabs [data-baseweb="tab"] {
            height: 48px; border-radius: 8px 8px 0 0; font-weight: 700;
            padding: 0 20px; color: #475569; font-size: 0.95rem;
        }
        .stTabs [aria-selected="true"] {
            background-color: #F1F5F9 !important; color: #0284C7 !important;
            border-bottom: 3px solid #0284C7 !important;
        }
        
        .qc-card-warning {
            background-color: #FEF2F2; border-left: 5px solid #DC2626;
            color: #991B1B; padding: 12px 16px; border-radius: 8px;
            margin-bottom: 10px; font-size: 0.92rem; font-weight: 500;
        }
        .saas-footer {
            text-align: center; padding: 2rem 0; color: #64748B;
            font-size: 0.85rem; border-top: 1px solid #E2E8F0; margin-top: 3rem; font-weight: 500;
        }
    </style>
    """, unsafe_allow_html=True)
else:
    st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
        html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
        .hero-container {
            background: linear-gradient(135deg, #4F46E5 0%, #7C3AED 100%);
            padding: 2.5rem 2rem; border-radius: 16px; color: white; margin-bottom: 2rem;
            box-shadow: 0 10px 25px -5px rgba(79, 70, 229, 0.3);
        }
        .hero-title { font-size: 2.4rem; font-weight: 800; margin: 0; letter-spacing: -0.02em; }
        .hero-subtitle { font-size: 1.05rem; opacity: 0.9; margin-top: 0.5rem; }
        .badge-pro {
            background-color: rgba(255, 255, 255, 0.2); backdrop-filter: blur(8px);
            padding: 4px 12px; border-radius: 9999px; font-size: 0.8rem; font-weight: 600;
            text-transform: uppercase; letter-spacing: 0.05em; display: inline-block; margin-bottom: 0.8rem;
        }
        .metric-card {
            background: white; border: 1px solid #E2E8F0; border-radius: 12px; padding: 1.25rem;
            box-shadow: 0 1px 3px 0 rgba(0, 0, 0, 0.05); transition: transform 0.2s, box-shadow 0.2s;
        }
        .metric-card:hover { transform: translateY(-2px); box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.05); }
        .metric-label { font-size: 0.85rem; color: #64748B; font-weight: 500; text-transform: uppercase; }
        .metric-value { font-size: 1.8rem; font-weight: 700; color: #0F172A; margin-top: 0.25rem; }
        .stTabs [data-baseweb="tab-list"] { gap: 8px; border-bottom: 2px solid #E2E8F0; }
        .stTabs [data-baseweb="tab"] { height: 50px; border-radius: 8px 8px 0 0; font-weight: 600; padding: 0 20px; }
        .qc-card-warning { background-color: #FEF2F2; border-left: 4px solid #EF4444; padding: 12px 16px; border-radius: 8px; margin-bottom: 8px; font-size: 0.9rem; }
        .saas-footer { text-align: center; padding: 2rem 0 1rem 0; color: #94A3B8; font-size: 0.85rem; border-top: 1px solid #E2E8F0; margin-top: 3rem; }
    </style>
    """, unsafe_allow_html=True)

# ==========================================
# 5. HÀM BÓC TÁCH BẰNG REGEX & THUẬT TOÁN
# ==========================================
def generate_vibrant_rgb_colors_excluding(excluded_rgbs, count=200):
    colors = []
    used_set = set(tuple(x) for x in excluded_rgbs)
    attempts = 0
    while len(colors) < count and attempts < 10000:
        attempts += 1
        h = random.random(); s = 0.9; v = 0.8
        i = int(h * 6.0); f = h * 6.0 - i; p = v * (1.0 - s); q = v * (1.0 - s * f); t = v * (1.0 - s * (1.0 - f))
        if i % 6 == 0: r, g, b = v, t, p
        elif i % 6 == 1: r, g, b = q, v, p
        elif i % 6 == 2: r, g, b = p, v, t
        elif i % 6 == 3: r, g, b = p, q, v
        elif i % 6 == 4: r, g, b = t, p, v
        else: r, g, b = v, p, q
        r_int, g_int, b_int = int(r * 255), int(g * 255), int(b * 255)
        rgb_tuple = (r_int, g_int, b_int)
        if rgb_tuple not in used_set:
            used_set.add(rgb_tuple)
            colors.append(rgb_tuple)
    return colors

def get_speaker_color_config(speaker_name, speaker_color_map, available_rgb_tuples):
    spk_upper = speaker_name.strip().upper()
    fixed_colors = st.session_state.get('fixed_speaker_colors', DEFAULT_FIXED_SPEAKER_COLORS)
    
    # 1. Nếu là nhân vật cố định trong Database
    if spk_upper in fixed_colors:
        cfg = fixed_colors[spk_upper]
        if isinstance(cfg, dict):
            return cfg
        elif isinstance(cfg, (tuple, list)):
            return {"text_color": tuple(cfg), "highlight_color": None}
            
    # 2. Nếu là nhân vật phụ / vãng lai ➔ Sinh màu chữ ngẫu nhiên KHÔNG TRÙNG LẶP
    if spk_upper not in speaker_color_map:
        if available_rgb_tuples:
            r, g, b = available_rgb_tuples.pop()
        else:
            r, g, b = (random.randint(50, 220), random.randint(50, 220), random.randint(50, 220))
        speaker_color_map[spk_upper] = {"text_color": (r, g, b), "highlight_color": None}
        
    return speaker_color_map[spk_upper]

def get_speaker_color(speaker_name, speaker_color_map, available_rgb_tuples):
    cfg = get_speaker_color_config(speaker_name, speaker_color_map, available_rgb_tuples)
    tc = cfg.get("text_color")
    if tc: return RGBColor(tc[0], tc[1], tc[2])
    return RGBColor(79, 70, 229)

def apply_speaker_styling_to_run(run, text_color_tuple, highlight_color_tuple):
    if text_color_tuple:
        r, g, b = text_color_tuple
        run.font.color.rgb = RGBColor(r, g, b)
    if highlight_color_tuple:
        hr, hg, hb = highlight_color_tuple
        hex_fill = f"{hr:02X}{hg:02X}{hb:02X}"
        shd_xml = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_fill}"/>')
        run._r.get_or_add_rPr().append(shd_xml)

def clean_and_normalize_text(text, strip_all_tags=False, fix_punctuation=True, normalize_spaces=True, capitalize_first=True, remove_leading_dash=True):
    if not text or not isinstance(text, str): return ""
    res = text
    
    # 1. Bóc tách thẻ ASS formatting
    res = re.sub(r'\{\\[^}]*\}', '', res)
    res = re.sub(r'\\N', '\n', res, flags=re.IGNORECASE)
    
    # 2. Bóc tách thẻ HTML rác
    if strip_all_tags:
        res = re.sub(r'<[^>]*>', '', res)
    else:
        res = re.sub(r'<(?!/?(i|b|u)\b)[^>]*>', '', res, flags=re.IGNORECASE)
        
    # 3. Xóa dấu gạch đầu dòng thừa
    if remove_leading_dash:
        res = re.sub(r'^\s*[-–—]\s*', '', res)
        res = re.sub(r'(\n)\s*[-–—]\s*', r'\1', res)
        
    # 4. Sửa lỗi dấu câu Tiếng Việt
    if fix_punctuation:
        res = re.sub(r'\s+([,!?:;\.\)])', r'\1', res)
        res = re.sub(r'([,!?:;])([A-Za-zÀ-ỹ0-9])', r'\1 \2', res)
        res = re.sub(r'(\.\.\.)([A-Za-zÀ-ỹ0-9])', r'\1 \2', res)
        res = re.sub(r'"\s*(.*?)\s*"', r'"\1"', res)
        
    # 5. Thu gọn khoảng trắng thừa
    if normalize_spaces:
        res = re.sub(r'[ \t]+', ' ', res)
        res = re.sub(r'\s*\n\s*', '\n', res)
        res = res.strip()
        
    # 6. Viết hoa đầu câu & Sau tên người nói
    if capitalize_first and res:
        lines = res.split('\n')
        cap_lines = []
        for line in lines:
            m = re.match(r'^([^:]+:\s*)([a-zà-ỹ])(.*)$', line)
            if m: line = m.group(1) + m.group(2).upper() + m.group(3)
            elif line and line[0].islower(): line = line[0].upper() + line[1:]
            cap_lines.append(line)
        res = '\n'.join(cap_lines)
        
    return res

def calculate_duration_sec(timecode_line):
    m = re.match(r"(\d{2}):(\d{2}):(\d{2})[,.](\d{3})\s+-->\s+(\d{2}):(\d{2}):(\d{2})[,.](\d{3})", timecode_line.strip())
    if not m: return 1.0, 0.0, 0.0
    h1, m1, s1, ms1, h2, m2, s2, ms2 = map(int, m.groups())
    t1 = h1 * 3600 + m1 * 60 + s1 + ms1 / 1000.0
    t2 = h2 * 3600 + m2 * 60 + s2 + ms2 / 1000.0
    dur = t2 - t1
    return dur if dur > 0 else 0.5, t1, t2

def timecode_to_sec(tc):
    m = re.match(r"(\d{2}):(\d{2}):(\d{2})[,.](\d{3})", str(tc).strip())
    if not m: return 0.0
    h, mins, s, ms = map(int, m.groups())
    return h * 3600 + mins * 60 + s + ms / 1000.0

def calculate_time_overlap(s1, e1, s2, e2):
    latest_start = max(s1, s2)
    earliest_end = min(e1, e2)
    return max(0.0, earliest_end - latest_start)

def is_valid_speaker_boundary(prefix, start_idx):
    if start_idx == 0: return True
    prev_char = prefix[start_idx - 1]
    curr_char = prefix[start_idx]
    if not prev_char.isalnum(): return True
    if prev_char.islower() and curr_char.isupper(): return True
    return False

def find_all_speaker_tags(text, custom_speakers=None, non_speakers=None):
    if custom_speakers is None: custom_speakers = st.session_state.get('custom_speakers', set())
    if non_speakers is None: non_speakers = st.session_state.get('custom_non_speakers', set())
    non_speakers_upper = {s.upper() for s in non_speakers}.union({s.upper() for s in DEFAULT_NON_SPEAKER_PHRASES})
    
    custom_names_upper = {s.upper(): s for s in custom_speakers if s.strip()}
    pattern = r"([A-Za-z0-9À-ỹ \t&\-\(\)\.]{1,35}):\s*"
    
    matches = []
    for m in re.finditer(pattern, text):
        raw_prefix = m.group(1)
        extracted_spk = None
        
        # A. Ưu tiên Whitelist Người nói
        for spk_upper, orig_spk in sorted(custom_names_upper.items(), key=lambda x: len(x[0]), reverse=True):
            if raw_prefix.upper().endswith(spk_upper):
                start_idx = len(raw_prefix) - len(spk_upper)
                if is_valid_speaker_boundary(raw_prefix, start_idx):
                    extracted_spk = orig_spk
                    spk_start = m.start(1) + start_idx
                    spk_end = m.end()
                    matches.append((spk_start, spk_end, extracted_spk, m.group(0)))
                    break
                    
        if extracted_spk: continue
            
        # B. Nhận diện tổng quát đuôi chuỗi
        match_tail = re.search(r"(?:^|[^\w]|(?<=[a-zà-ỹ]))([A-ZÀ-Ỹ0-9][A-Za-z0-9À-ỹ \t&\-\(\)\.]{0,24})$", raw_prefix)
        if match_tail:
            cand = match_tail.group(1).strip(".,!?:;- ")
            if cand and len(cand) <= 25 and not cand.isdigit():
                if not (cand.startswith('(') or cand.endswith(')')):
                    if cand.upper() not in non_speakers_upper:
                        if len(cand.split()) <= 4:
                            start_idx = len(raw_prefix) - len(match_tail.group(1))
                            if is_valid_speaker_boundary(raw_prefix, start_idx):
                                spk_start = m.start(1) + start_idx
                                spk_end = m.end()
                                matches.append((spk_start, spk_end, cand, m.group(0)))

    matches.sort(key=lambda x: x[0])
    filtered_matches = []
    last_end = -1
    for start, end, spk, raw_m in matches:
        if start >= last_end:
            filtered_matches.append((start, end, spk, raw_m))
            last_end = end
            
    return filtered_matches

def is_valid_speaker_name(name):
    clean = name.strip()
    if not clean or len(clean) > 25 or clean.isdigit() or re.match(r'^\d+[\d\s:]*$', clean): return False
    if is_stage_direction(clean) or clean.upper() in NON_SPEAKER_PHRASES: return False
    if any(char in clean for char in ['/', '?', '!', ',', '.', '-->', '(', ')']): return False
    if len(clean.split()) > 4: return False
    return True

def process_srt_to_docx(uploaded_file, file_name_without_ext):
    srt_content = uploaded_file.getvalue().decode('utf-8', errors='ignore')
    blocks = re.split(r'\n\s*\n', srt_content.strip())
    document = Document()

    for block in blocks:
        lines = [l.strip() for l in block.strip().split('\n') if l.strip()]
        if not lines: continue
        
        idx_str = ""; tc_str = ""; text_lines = []
        if lines[0].isdigit():
            idx_str = lines[0]
            if len(lines) > 1 and "-->" in lines[1]:
                tc_str = lines[1]; text_lines = lines[2:]
            else: text_lines = lines[1:]
        elif "-->" in lines[0]:
            tc_str = lines[0]; text_lines = lines[1:]
        else: text_lines = lines

        if idx_str:
            p_index = document.add_paragraph(idx_str)
            p_index.runs[0].font.name = 'Times New Roman'; p_index.runs[0].font.size = Pt(12)
            p_index.paragraph_format.space_after = Pt(0)

        if tc_str:
            p_timecode = document.add_paragraph(tc_str)
            p_timecode.runs[0].font.name = 'Times New Roman'; p_timecode.runs[0].font.size = Pt(12)
            p_timecode.paragraph_format.space_after = Pt(0)
            
        if text_lines:
            raw_text = "\n".join(text_lines)
            clean_text = clean_and_normalize_text(raw_text, strip_all_tags=True)
            p_content = document.add_paragraph(clean_text)
            p_content.runs[0].font.name = 'Times New Roman'; p_content.runs[0].font.size = Pt(12)
            p_content.paragraph_format.space_after = Pt(12)

    modified_file = io.BytesIO()
    document.save(modified_file)
    modified_file.seek(0)
    return modified_file

def process_docx_to_srt(uploaded_file):
    document = Document(uploaded_file)
    lines = [p.text.strip() for p in document.paragraphs if p.text.strip() != ""]
    srt_content = ""
    timecode_pattern = re.compile(r'\d{2}:\d{2}:\d{2}[,.]\d{3}\s*-->\s*\d{2}:\d{2}:\d{2}[,.]\d{3}')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.isdigit() and i + 1 < len(lines) and timecode_pattern.search(lines[i+1]):
            if srt_content != "": srt_content += "\n"
            srt_content += line + "\n" + lines[i+1] + "\n"
            i += 2 
            while i < len(lines):
                if lines[i].isdigit() and i + 1 < len(lines) and timecode_pattern.search(lines[i+1]): break 
                else: 
                    clean_l = clean_and_normalize_text(lines[i], strip_all_tags=True)
                    srt_content += clean_l + "\n"; i += 1
        else: i += 1
            
    return srt_content.strip().encode('utf-8-sig')

# --- SRT TO EXCEL CONVERTER ---
EXCEL_COLOR_PALETTE = [
    'background-color: #ADD8E6; color: #000000',
    'background-color: #90EE90; color: #000000',
    'background-color: #FFB6C1; color: #000000',
    'background-color: #FFFFE0; color: #000000',
    'background-color: #DDA0DD; color: #000000',
    'background-color: #AFEEEE; color: #000000',
    'background-color: #F0E68C; color: #000000',
    'background-color: #FFA07A; color: #000000',
    'background-color: #E0FFFF; color: #000000',
    'background-color: #F5F5DC; color: #000000',
    'background-color: #2F4F4F; color: #FFFFFF',
    'background-color: #191970; color: #FFFFFF',
    'background-color: #006400; color: #FFFFFF',
    'background-color: #800000; color: #FFFFFF',
    'background-color: #4B0082; color: #FFFFFF',
    'background-color: #556B2F; color: #FFFFFF',
    'background-color: #8B4513; color: #FFFFFF',
    'background-color: #36454F; color: #FFFFFF',
]

def clean_dialogue_text_for_excel(text):
    text = clean_and_normalize_text(text, strip_all_tags=True)
    return text

def parse_srt_to_dataframe(srt_content, custom_speakers=None, non_speakers=None, default_speaker="Unknown"):
    if custom_speakers is None: custom_speakers = st.session_state.get('custom_speakers', set())
    if non_speakers is None: non_speakers = st.session_state.get('custom_non_speakers', set())

    fallback_spk = default_speaker.strip() if default_speaker and default_speaker.strip() else "Unknown"

    data = []
    blocks = re.split(r'\n\s*\n', srt_content.strip())
    last_known_speaker = fallback_spk
    last_is_explicit = False

    for block in blocks:
        lines = [l.strip() for l in block.strip().split('\n') if l.strip()]
        if len(lines) < 2: continue

        time_line = ""; time_idx = -1
        for idx, line in enumerate(lines[:2]):
            if "-->" in line: time_line = line; time_idx = idx; break
        if not time_line: continue

        time_match = re.match(r'(\d{2}:\d{2}:\d{2}[,.]\d{3})\s*-->\s*(\d{2}:\d{2}:\d{2}[,.]\d{3})', time_line)
        if not time_match: continue

        time_start = time_match.group(1).replace('.', ',')
        time_end = time_match.group(2).replace('.', ',')

        dialogue_lines = lines[time_idx + 1:]
        if not dialogue_lines: continue

        block_text = "\n".join(dialogue_lines)
        speaker_tags = find_all_speaker_tags(block_text, custom_speakers, non_speakers)

        if not speaker_tags:
            clean_text = clean_dialogue_text_for_excel(block_text)
            if clean_text: data.append([time_start, time_end, last_known_speaker, clean_text, last_is_explicit])
        else:
            last_idx = 0
            for i, (start_pos, end_pos, spk_name, raw_m) in enumerate(speaker_tags):
                leading_text = block_text[last_idx:start_pos].strip()
                clean_leading = clean_dialogue_text_for_excel(leading_text)
                if clean_leading: data.append([time_start, time_end, last_known_speaker, clean_leading, last_is_explicit])

                next_start = speaker_tags[i+1][0] if i + 1 < len(speaker_tags) else len(block_text)
                segment_text = block_text[end_pos:next_start]
                clean_seg = clean_dialogue_text_for_excel(segment_text)
                if clean_seg:
                    data.append([time_start, time_end, spk_name, clean_seg, True])
                    last_known_speaker = spk_name
                    last_is_explicit = True
                else:
                    last_known_speaker = spk_name
                    last_is_explicit = True
                last_idx = next_start

    return pd.DataFrame(data, columns=['Start', 'End', 'Speaker', 'Dialogue', 'Is_Explicit'])

def apply_excel_styles(df):
    unique_speakers = df['Speaker'].unique()
    color_map = {speaker: EXCEL_COLOR_PALETTE[i % len(EXCEL_COLOR_PALETTE)] for i, speaker in enumerate(unique_speakers)}
    def highlight_speaker(row):
        color_style = color_map.get(row['Speaker'], 'background-color: #FFFFFF; color: #000000')
        return [color_style] * len(row)
    try: return df.style.apply(highlight_speaker, axis=1)
    except Exception: return df

# --- DUAL ENGLISH COMPARISON & VIETNAMESE QC HELPERS ---
def parse_any_script_file_to_df(file_bytes, filename, custom_speakers=None, non_speakers=None, default_speaker="Unknown"):
    if filename.lower().endswith('.srt'):
        try: content_str = file_bytes.decode('utf-8')
        except UnicodeDecodeError: content_str = file_bytes.decode('latin-1')
        return parse_srt_to_dataframe(content_str, custom_speakers, non_speakers, default_speaker)
    elif filename.lower().endswith('.docx'):
        doc = Document(io.BytesIO(file_bytes))
        paragraphs_text = [p.text.strip() for p in doc.paragraphs if p.text.strip() != ""]
        timecode_pattern = re.compile(r'\d{2}:\d{2}:\d{2}[,.]\d{3}\s*-->\s*\d{2}:\d{2}:\d{2}[,.]\d{3}')
        
        srt_lines = []
        i = 0
        while i < len(paragraphs_text):
            line = paragraphs_text[i]
            if line.isdigit() and i + 1 < len(paragraphs_text) and timecode_pattern.search(paragraphs_text[i+1]):
                srt_lines.append(line)
                srt_lines.append(paragraphs_text[i+1])
                i += 2
                while i < len(paragraphs_text):
                    if paragraphs_text[i].isdigit() and i + 1 < len(paragraphs_text) and timecode_pattern.search(paragraphs_text[i+1]): break
                    else: srt_lines.append(paragraphs_text[i]); i += 1
            else: i += 1
        content_str = "\n".join(srt_lines)
        return parse_srt_to_dataframe(content_str, custom_speakers, non_speakers, default_speaker)
    return pd.DataFrame(columns=['Start', 'End', 'Speaker', 'Dialogue', 'Is_Explicit'])

def align_and_compare_english_scripts(df_mh_eng, df_off_eng, df_vn=None, default_speaker="Unknown"):
    aligned_rows = []
    fallback_spk = default_speaker.strip() if default_speaker and default_speaker.strip() else "Unknown"
    
    for idx_mh, row_mh in df_mh_eng.iterrows():
        s_mh = timecode_to_sec(row_mh['Start'])
        e_mh = timecode_to_sec(row_mh['End'])
        dur_mh = max(0.5, e_mh - s_mh)
        
        is_explicit_mh = row_mh.get('Is_Explicit', False)
        
        best_off_matches = []
        for idx_off, row_off in df_off_eng.iterrows():
            s_off = timecode_to_sec(row_off['Start'])
            e_off = timecode_to_sec(row_off['End'])
            overlap = calculate_time_overlap(s_mh, e_mh, s_off, e_off)
            if overlap > 0.1:
                best_off_matches.append((idx_off, overlap, row_off))
                
        if best_off_matches:
            off_indices = [m[0] for m in best_off_matches]
            min_off_idx = min(off_indices)
            max_off_idx = max(off_indices)
            
            prev_off_text = str(df_off_eng.iloc[min_off_idx-1]['Dialogue']) if min_off_idx > 0 and pd.notna(df_off_eng.iloc[min_off_idx-1]['Dialogue']) else ""
            next_off_text = str(df_off_eng.iloc[max_off_idx+1]['Dialogue']) if max_off_idx < len(df_off_eng)-1 and pd.notna(df_off_eng.iloc[max_off_idx+1]['Dialogue']) else ""
            
            off_dialogues = [str(m[2]['Dialogue']) for m in best_off_matches if pd.notna(m[2]['Dialogue'])]
            off_text_combined = " ".join(off_dialogues)
            off_spk = best_off_matches[0][2]['Speaker']
            
            off_window_text = f"{prev_off_text} {off_text_combined} {next_off_text}".strip()
        else:
            off_text_combined = ""
            off_window_text = ""
            off_spk = ""

        vn_text_combined = ""
        vn_spk = ""
        if df_vn is not None and not df_vn.empty:
            best_vn_matches = []
            for idx_vn, row_vn in df_vn.iterrows():
                s_vn = timecode_to_sec(row_vn['Start'])
                e_vn = timecode_to_sec(row_vn['End'])
                overlap_vn = calculate_time_overlap(s_mh, e_mh, s_vn, e_vn)
                if overlap_vn > 0.1:
                    best_vn_matches.append((overlap_vn, row_vn))
            if best_vn_matches:
                vn_text_combined = " ".join([str(m[1]['Dialogue']) for m in best_vn_matches if pd.notna(m[1]['Dialogue'])])
                vn_spk = best_vn_matches[0][1]['Speaker']
        else:
            vn_spk = str(row_mh['Speaker']) if pd.notna(row_mh['Speaker']) else ""

        qc_status = "🟢 Khớp chuẩn"
        qc_details = "Nội dung Tiếng Anh khớp chuẩn nghĩa"
        
        mh_diag_clean = str(row_mh['Dialogue']).strip() if pd.notna(row_mh['Dialogue']) else ""
        off_diag_clean = off_text_combined.strip()
        
        if not off_diag_clean:
            qc_status = "🔴 Thiếu câu gốc Khách"
            qc_details = "File Mai Han có câu này nhưng file Khách không thấy có"
        elif not mh_diag_clean:
            qc_status = "🔴 Thiếu câu Mai Han"
            qc_details = "File Khách có câu này nhưng Mai Han nghe bị bỏ sót"
        else:
            mh_words = set(re.findall(r'\b\w+\b', mh_diag_clean.lower()))
            off_window_words = set(re.findall(r'\b\w+\b', off_window_text.lower()))
            
            if mh_words:
                missing_in_window = mh_words - off_window_words
                found_ratio = (len(mh_words) - len(missing_in_window)) / len(mh_words)
                
                if found_ratio < 0.75:
                    qc_status = "🟡 Khác từ vựng Tiếng Anh"
                    qc_details = f"Thiếu các từ gốc: {', '.join(list(missing_in_window)[:5])} -> Kiểm tra lại câu Tiếng Việt!"

            spk_mh_clean = str(row_mh['Speaker']).strip().upper() if pd.notna(row_mh['Speaker']) else ""
            spk_off_clean = str(off_spk).strip().upper() if off_spk else ""
            
            if spk_mh_clean and spk_off_clean and spk_mh_clean != spk_off_clean:
                if spk_mh_clean not in ["UNKNOWN", fallback_spk.upper()] and spk_off_clean not in ["UNKNOWN", fallback_spk.upper()]:
                    qc_status = "🔵 Lệch người nói"
                    qc_details = f"Mai Han: {row_mh['Speaker']} vs Khách: {off_spk}"

        spk_display_mh = str(row_mh['Speaker']) if pd.notna(row_mh['Speaker']) and str(row_mh['Speaker']) != "Unknown" else fallback_spk
        spk_display_off = off_spk if off_spk and off_spk != "Unknown" else fallback_spk
        spk_display_vn = vn_spk if vn_spk and vn_spk != "Unknown" else fallback_spk

        aligned_rows.append({
            "Stt": idx_mh + 1,
            "Timecode": f"{row_mh['Start']} --> {row_mh['End']}",
            "Start": row_mh['Start'],
            "End": row_mh['End'],
            "Tiếng Anh Mai Han (AI/Heard)": f"{spk_display_mh}: {row_mh['Dialogue']}",
            "Tiếng Anh Khách (Official)": f"{spk_display_off}: {off_text_combined}" if spk_display_off else off_text_combined,
            "Dịch Tiếng Việt (Cần Sửa)": f"{spk_display_vn}: {vn_text_combined}" if spk_display_vn else vn_text_combined,
            "Speaker_MH": spk_display_mh,
            "Is_Explicit_MH": is_explicit_mh,
            "Speaker_Off": spk_display_off,
            "Speaker_VN": spk_display_vn,
            "Dialogue_MH": row_mh['Dialogue'],
            "Dialogue_Off": off_text_combined,
            "Dialogue_VN": vn_text_combined,
            "Trạng thái QC": qc_status,
            "Ghi chú QC": qc_details
        })

    return pd.DataFrame(aligned_rows)

def generate_qc_dual_excel(df_aligned):
    out_buf = io.BytesIO()
    def highlight_qc_rows(row):
        st_val = str(row['Trạng thái QC'])
        if '🔴' in st_val: return ['background-color: #FEE2E2; color: #991B1B'] * len(row)
        elif '🟡' in st_val: return ['background-color: #FEF3C7; color: #92400E'] * len(row)
        elif '🔵' in st_val: return ['background-color: #E0F2FE; color: #075985'] * len(row)
        return ['background-color: #FFFFFF; color: #000000'] * len(row)

    styled_df = df_aligned.style.apply(highlight_qc_rows, axis=1)
    with pd.ExcelWriter(out_buf, engine='openpyxl') as writer:
        styled_df.to_excel(writer, index=False, sheet_name="Doi_Chieu_English_QC")
    out_buf.seek(0)
    return out_buf

def generate_aligned_docx_file(df_aligned, title_text, enable_colors=True, enable_phonetic=True, enable_cast=True, hide_default_spk=True, fallback_spk_name="Unknown", font_size_pt=12):
    document = Document()
    p_title = document.add_paragraph(title_text.upper())
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.runs[0].font.name = 'Times New Roman'; p_title.runs[0].font.size = Pt(20); p_title.runs[0].bold = True
    
    unique_speakers = []
    for spk in df_aligned['Speaker_MH']:
        if spk and spk not in unique_speakers and spk != "Unknown": unique_speakers.append(spk)
            
    if unique_speakers and not (hide_default_spk and len(unique_speakers) == 1 and unique_speakers[0].upper() == fallback_spk_name.upper()):
        speaker_list_text = "VAI: " + ", ".join(unique_speakers)
        p = document.add_paragraph(speaker_list_text)
        p.runs[0].font.name = 'Times New Roman'; p.runs[0].font.size = Pt(font_size_pt); p.runs[0].bold = True
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(6)

    document.add_paragraph()
    TAB_STOP = Inches(1.0)
    
    speaker_color_map = {}
    fixed_colors = st.session_state.get('fixed_speaker_colors', DEFAULT_FIXED_SPEAKER_COLORS)
    fixed_rgb_set = {tuple(v.get("text_color")) for v in fixed_colors.values() if isinstance(v, dict) and v.get("text_color")}
    available_rgb_tuples = generate_vibrant_rgb_colors_excluding(fixed_rgb_set, 200)
    random.shuffle(available_rgb_tuples)

    for idx, row in df_aligned.iterrows():
        tc_line = row['Timecode']
        spk = row['Speaker_MH'] if row['Speaker_MH'] else "Unknown"
        diag = row['Dialogue_VN'] if row['Dialogue_VN'] else ""
        is_explicit = row.get('Is_Explicit_MH', True)
        
        p_tc = document.add_paragraph(tc_line)
        p_tc.runs[0].font.name = 'Times New Roman'; p_tc.runs[0].font.size = Pt(font_size_pt); p_tc.runs[0].bold = True
        p_tc.paragraph_format.space_before = Pt(0); p_tc.paragraph_format.space_after = Pt(0)
        
        p_line = document.add_paragraph()
        p_line.paragraph_format.left_indent = TAB_STOP
        p_line.paragraph_format.first_line_indent = Inches(-1.0)
        p_line.paragraph_format.tab_stops.add_tab_stop(TAB_STOP, WD_TAB_ALIGNMENT.LEFT)
        
        should_show_spk = True
        if hide_default_spk and (not is_explicit or spk.upper() == fallback_spk_name.upper() or spk.upper() == "UNKNOWN"):
            should_show_spk = False

        if should_show_spk:
            r_spk = p_line.add_run(f"{spk}:")
            r_spk.font.name = 'Times New Roman'; r_spk.font.size = Pt(font_size_pt); r_spk.font.bold = True
            if enable_colors:
                spk_cfg = get_speaker_color_config(spk, speaker_color_map, available_rgb_tuples)
                apply_speaker_styling_to_run(r_spk, spk_cfg.get("text_color"), spk_cfg.get("highlight_color"))
            p_line.add_run("\t")
        else:
            p_line.add_run("\t")

        if diag: apply_html_and_phonetic_to_paragraph(p_line, diag, enable_phonetic)
            
        p_line.paragraph_format.space_before = Pt(0); p_line.paragraph_format.space_after = Pt(4)
        
    for p in document.paragraphs:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for r in p.runs:
            r.font.name = 'Times New Roman'
            if r.font.size is None: r.font.size = Pt(font_size_pt)
        
    buf = io.BytesIO(); document.save(buf); buf.seek(0)
    return buf

# --- HÀM TẠO MARKER TIMELINE CHO DAW ---
def sec_to_reaper_tc(sec):
    h = int(sec // 3600); mins = int((sec % 3600) // 60); s = int(sec % 60)
    ms = int(round((sec - int(sec)) * 1000))
    if ms >= 1000: s += 1; ms -= 1000
    return f"{h:02d}:{mins:02d}:{s:02d}.{ms:03d}"

def sec_to_fps_tc(sec, fps=25):
    h = int(sec // 3600); mins = int((sec % 3600) // 60); s = int(sec % 60)
    frames = int(round((sec - int(sec)) * fps))
    if frames >= fps: s += 1; frames -= fps
    return f"{h:02d}:{mins:02d}:{s:02d}:{frames:02d}"

def generate_reaper_region_csv(df):
    rows = ["#,Name,Start,End,Length,Color"]
    for idx, r in df.iterrows():
        s_sec = timecode_to_sec(r['Start']); e_sec = timecode_to_sec(r['End'])
        dur_sec = max(0.1, e_sec - s_sec)
        s_tc = sec_to_reaper_tc(s_sec); e_tc = sec_to_reaper_tc(e_sec); l_tc = sec_to_reaper_tc(dur_sec)
        name_clean = f"{r['Speaker']}: {r['Dialogue']}".replace('"', '""')
        rows.append(f'R{idx+1},"{name_clean}",{s_tc},{e_tc},{l_tc},')
    return "\n".join(rows)

def generate_pro_tools_csv(df):
    rows = ["Marker Name,Timecode In,Timecode Out,Comment"]
    for idx, r in df.iterrows():
        s_sec = timecode_to_sec(r['Start']); e_sec = timecode_to_sec(r['End'])
        s_tc = sec_to_fps_tc(s_sec, 25); e_tc = sec_to_fps_tc(e_sec, 25)
        spk = r['Speaker'].replace('"', '""'); diag = r['Dialogue'].replace('"', '""')
        rows.append(f'"{spk}",{s_tc},{e_tc},"{diag}"')
    return "\n".join(rows)

def generate_cmx3600_edl(df):
    lines = ["TITLE: SCRIPTPRO_DAW_MARKERS", "FCM: NON-DROP FRAME", ""]
    for idx, r in df.iterrows():
        s_sec = timecode_to_sec(r['Start']); e_sec = timecode_to_sec(r['End'])
        s_tc = sec_to_fps_tc(s_sec, 25); e_tc = sec_to_fps_tc(e_sec, 25)
        lines.append(f"{idx+1:03d}  AX       V     C        {s_tc} {e_tc} {s_tc} {e_tc}")
        lines.append(f"* FROM CLIP: {r['Speaker']}")
        lines.append(f"* COMMENT: {r['Dialogue']}")
        lines.append("")
    return "\n".join(lines)

def generate_english_audio(text_to_speak, accent='com'):
    try:
        tts = gTTS(text=text_to_speak, lang='en', tld=accent)
        fp = io.BytesIO(); tts.write_to_fp(fp); fp.seek(0)
        return fp
    except Exception as e:
        st.error(f"Không thể tải âm thanh: {e}")
        return None

def is_candidate_english_word(word):
    clean_w = word.strip(".,!?:;\"'()[]{}")
    if not clean_w or len(clean_w) <= 1 or clean_w.isdigit(): return False
    lower_w = clean_w.lower()
    if clean_w.upper() in st.session_state['custom_phonetics']: return True
    if lower_w in VN_SYLLABLES: return False
    if any(char in lower_w for char in ['f', 'j', 'w', 'z']): return True
        
    eng_patterns = [
        r"(bb|cc|dd|ff|gg|ll|mm|nn|pp|rr|ss|tt|zz)",
        r"(sh|ck|th|wh|ph|gh)",
        r"(tion|ment|ness|less|ing|ed|able|ible|ally|ce|ge|ck)$",
        r"^([A-Z]{2,})$"
    ]
    for pattern in eng_patterns:
        if re.search(pattern, clean_w): return True
    if clean_w[0].isupper() and lower_w not in VN_SYLLABLES: return True
    return False

def generate_actor_salary_slip_docx(actor_name, week_name, video_rows, total_pay, current_mode):
    doc = Document()
    p_title = doc.add_paragraph("MAI HAN STUDIO - PHIẾU BÁO CÁO THÙ LAO LỒNG TIẾNG")
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.runs[0].font.name = 'Times New Roman'; p_title.runs[0].font.size = Pt(16); p_title.runs[0].bold = True
    p_sub = doc.add_paragraph(f"Diễn viên Lồng tiếng: {actor_name} | {week_name}")
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.runs[0].font.name = 'Times New Roman'; p_sub.runs[0].font.size = Pt(12); p_sub.runs[0].font.italic = True
    
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=5); table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells; hdr_names = ["Stt", "Tiêu đề video", "Khối lượng", "Đơn giá áp dụng", "Thành tiền"]
    for i, name in enumerate(hdr_names):
        hdr_cells[i].text = name; hdr_cells[i].paragraphs[0].runs[0].font.bold = True; hdr_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        
    for r in video_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(r["Stt"]); row_cells[1].text = str(r["Tiêu đề video"]); row_cells[2].text = str(r["Thời lượng"])
        row_cells[3].text = str(r["Đơn giá"]); row_cells[4].text = str(r["Thành tiền"])
        for c in row_cells:
            for p in c.paragraphs:
                for run in p.runs: run.font.name = 'Times New Roman'; run.font.size = Pt(11)

    doc.add_paragraph()
    p_total = doc.add_paragraph(f"👉 TỔNG THÙ LAO THANH TOÁN: {total_pay:,.0f} VNĐ")
    p_total.runs[0].font.name = 'Times New Roman'; p_total.runs[0].font.size = Pt(13); p_total.runs[0].bold = True
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

def round_seconds_to_int_minutes(total_sec):
    if total_sec <= 0: return 1
    mins = int(total_sec // 60); secs = int(round(total_sec % 60))
    if secs >= 30: mins += 1
    return max(1, mins)

def srt_timecode_to_ass(timecode_line):
    m = re.match(r"(\d{2}):(\d{2}):(\d{2}),(\d{3})\s+-->\s+(\d{2}):(\d{2}):(\d{2}),(\d{3})", timecode_line.strip())
    if not m: return None, None
    h1, m1, s1, ms1, h2, m2, s2, ms2 = m.groups()
    ass_start = f"{int(h1)}:{m1}:{s1}.{int(ms1)//10:02d}"
    ass_end = f"{int(h2)}:{m2}:{s2}.{int(ms2)//10:02d}"
    return ass_start, ass_end

def rgb_to_ass_hex(rgb_obj):
    if not rgb_obj: return "&H00FFFFFF&"
    try:
        r, g, b = rgb_obj[0], rgb_obj[1], rgb_obj[2]
        return f"&H00{b:02X}{g:02X}{r:02X}&"
    except Exception: return "&H00FFFFFF&"

def is_stage_direction(name):
    clean = name.strip()
    return clean.startswith('(') or clean.endswith(')')

def preprocess_raw_paragraphs(raw_paragraphs, custom_speakers, non_speakers):
    cleaned_paras = []; i = 0; total = len(raw_paragraphs)
    while i < total:
        raw_text = get_paragraph_text_with_html(raw_paragraphs[i])
        text = re.sub(r'\t+', ' ', raw_text).strip()
        if not text:
            i += 1; continue
        spk_tags = find_all_speaker_tags(text, custom_speakers, non_speakers)
        if spk_tags:
            last_match_end = spk_tags[-1][1]
            content_after = text[last_match_end:].strip()
            real_content = re.sub(r'</?[ibuIBU]>', '', content_after).strip()
            if not real_content:
                next_i = i + 1
                while next_i < total:
                    next_raw_text = get_paragraph_text_with_html(raw_paragraphs[next_i])
                    next_text = re.sub(r'\t+', ' ', next_raw_text).strip()
                    if next_text: break
                    next_i += 1
                if next_i < total:
                    next_raw_text = get_paragraph_text_with_html(raw_paragraphs[next_i])
                    next_text = re.sub(r'\t+', ' ', next_raw_text).strip()
                    is_timecode = TIMECODE_REGEX.match(next_text)
                    is_number = re.fullmatch(r"^\s*\d+\s*$", next_text)
                    is_srt = next_text.lower().startswith("srt conversion") or next_text.lower().startswith("vai:")
                    next_spk_tags = find_all_speaker_tags(next_text, custom_speakers, non_speakers)
                    if not (is_timecode or is_number or is_srt or next_spk_tags):
                        text = f"{text}{next_text}" if text.endswith('>') else f"{text} {next_text}"
                        i = next_i
        cleaned_paras.append(text); i += 1
    return cleaned_paras

def scan_candidate_speakers(uploaded_file, custom_speakers, non_speakers):
    doc = Document(io.BytesIO(uploaded_file.getvalue()))
    raw_paragraphs = [p for p in doc.paragraphs]
    processed_strings = preprocess_raw_paragraphs(raw_paragraphs, custom_speakers, non_speakers)
    candidates = Counter()
    for text in processed_strings:
        if not text or text.lower().startswith("srt conversion"): continue
        tags = find_all_speaker_tags(text, custom_speakers, non_speakers)
        for _, _, spk_name, _ in tags:
            candidates[spk_name] += 1
    return candidates

def scan_english_words_in_dialogue(uploaded_file, custom_speakers, non_speakers):
    doc = Document(io.BytesIO(uploaded_file.getvalue()))
    raw_paragraphs = [p for p in doc.paragraphs]
    processed_strings = preprocess_raw_paragraphs(raw_paragraphs, custom_speakers, non_speakers)
    eng_found = set()
    for text in processed_strings:
        if not text or text.lower().startswith("srt conversion") or TIMECODE_REGEX.match(text): continue
        tags = find_all_speaker_tags(text, custom_speakers, non_speakers)
        dialogue_content = ""
        if not tags: dialogue_content = text
        else:
            last_idx = 0
            for start_pos, end_pos, _, _ in tags:
                dialogue_content += " " + text[last_idx:start_pos]
                last_idx = end_pos
            dialogue_content += " " + text[last_idx:]

        for match in ENGLISH_WORD_REGEX.finditer(dialogue_content):
            word = match.group(0).strip()
            if is_candidate_english_word(word): eng_found.add(word)
    return sorted(list(eng_found), key=lambda x: x.upper())

def extract_phrases_from_file(file_io, file_name):
    phrases = set()
    try:
        if file_name.endswith('.txt'):
            content = file_io.getvalue().decode("utf-8")
            phrases.update([line.strip() for line in content.split('\n') if line.strip()])
        elif file_name.endswith('.docx'):
            doc = Document(io.BytesIO(file_io.getvalue()))
            for p in doc.paragraphs:
                p_text = re.sub(r'\t+', ' ', p.text).strip()
                if p_text:
                    parts = re.split(r'[,\n]', p_text)
                    phrases.update([part.strip() for part in parts if part.strip()])
        elif file_name.endswith('.xlsx'):
            df = pd.read_excel(file_io, header=None)
            for col in df.columns:
                for item in df[col].dropna():
                    parts = re.split(r'[,\n]', str(item))
                    phrases.update([part.strip() for part in parts if part.strip()])
    except Exception as e: st.error(f"Lỗi đọc file: {e}")
    return phrases

# --- HERO BANNER HEADER ---
st.markdown(f"""
<div class="hero-container">
    <div class="badge-pro">{ui_theme_choice}</div>
    <div class="hero-title">🎬 ScriptPro Enterprise Studio</div>
    <div class="hero-subtitle">Hệ thống xử lý kịch bản lồng tiếng, chuẩn hóa định dạng Word, phân vai & báo cáo thù lao cá nhân thông minh.</div>
</div>
""", unsafe_allow_html=True)

# --- MÀN HÌNH CHÍNH TÁCH 9 TABS ---
tab_script, tab_resync, tab_dub_tracker, tab_cast_db, tab_phonetic_db, tab_dual_align, tab_consistency, tab_cleaner, tab_tools = st.tabs([
    "🎬 Xử lý Kịch bản Gốc", 
    "🔄 Re-Sync Kịch Bản Biên Tập",
    "📋 Theo dõi & Báo cáo Lương",
    "🎭 Bảng Phân Vai & Mã Màu", 
    "📚 Kho Database Phiên Âm Giọng Nam",
    "🔀 Đối Chiếu 2 File Tiếng Anh & QC Dịch",
    "🔎 Soát Bất Nhất Thuật Ngữ & Xưng Hô",
    "🧹 Dọn Dẹp & Chuẩn Hóa Phụ Đề",
    "🧰 Bộ Công Cụ Chuyển Đổi"
])

# ==========================================
# TAB 1: XỬ LÝ KỊCH BẢN GỐC
# ==========================================
with tab_script:
    col1, col2 = st.columns([1.6, 1])

    with col1:
        with st.container(border=True):
            st.markdown("### 📁 Tải lên file Kịch bản Word gốc (.docx)")
            uploaded_file = st.file_uploader(
                "Kéo thả file .docx gốc của bạn vào đây", 
                type=['docx'], 
                key=f"main_uploader_{st.session_state['uploader_key']}"
            )

        if uploaded_file is None:
            st.info("📌 **Vui lòng tải file kịch bản (.docx) ở trên để hiển thị công cụ biên tập.**")
        else:
            original_filename = uploaded_file.name
            file_name_without_ext = os.path.splitext(original_filename)[0] 
            st.success(f"📄 Đã nhận file thành công: **{original_filename}**")

            custom_spks = st.session_state.get('custom_speakers', set())
            custom_non_spks = st.session_state.get('custom_non_speakers', set())

            candidates = scan_candidate_speakers(uploaded_file, custom_spks, custom_non_spks)

            detected_speakers_names = [name for name in candidates.keys() if name.upper() not in NON_SPEAKER_PHRASES]
            detected_non_speakers_names = [name for name in candidates.keys() if name.upper() in NON_SPEAKER_PHRASES]

            detected_speakers = [f"{name} ({candidates[name]} lần)" for name in detected_speakers_names]
            detected_non_speakers = [f"{name} ({candidates[name]} lần)" for name in detected_non_speakers_names]

            with st.container(border=True):
                st.markdown("### 🎭 Phân Vai Lồng Tiếng Cho Kịch Bản Hiện Tại")
                st.caption("Xem và gán người lồng tiếng Việt cho từng nhân vật trong file kịch bản này:")

                if detected_speakers_names:
                    cast_table_data = []
                    for spk_name in detected_speakers_names:
                        current_actor = st.session_state['custom_cast_mapping'].get(spk_name.upper(), "")
                        cast_table_data.append({
                            "Nhân vật (Tiếng Anh)": spk_name,
                            "Diễn viên Lồng tiếng (Tiếng Việt)": current_actor,
                            "Nạp vào Database": True
                        })

                    df_cast = pd.DataFrame(cast_table_data)

                    edited_cast_df = st.data_editor(
                        df_cast,
                        column_config={
                            "Nhân vật (Tiếng Anh)": st.column_config.TextColumn("Nhân vật (Kịch bản gốc)", disabled=True),
                            "Diễn viên Lồng tiếng (Tiếng Việt)": st.column_config.TextColumn("Diễn viên lồng tiếng (Sửa trực tiếp)"),
                            "Nạp vào Database": st.column_config.CheckboxColumn("Lưu Database?", default=True)
                        },
                        disabled=["Nhân vật (Tiếng Anh)"],
                        hide_index=True,
                        use_container_width=True,
                        key="script_cast_editor_table"
                    )

                    if st.button("💾 Lưu Bảng Phân Vai Kịch Bản Này Vào Database", type="secondary", use_container_width=True):
                        updated_cast_count = 0
                        for _, row in edited_cast_df.iterrows():
                            if row["Nạp vào Database"]:
                                spk_k = str(row["Nhân vật (Tiếng Anh)"]).upper().strip()
                                act_v = str(row["Diễn viên Lồng tiếng (Tiếng Việt)"]).strip().upper()
                                if act_v:
                                    st.session_state['custom_cast_mapping'][spk_k] = act_v
                                    updated_cast_count += 1
                        save_json_db(CAST_DB_FILE, st.session_state['custom_cast_mapping'])
                        st.success(f"✅ Đã lưu phân vai cho {updated_cast_count} nhân vật vào Database!")
                        time.sleep(1); st.rerun()

            with st.container(border=True):
                st.markdown("### 🔍 Soát Lỗi Nhận Diện Tên Người Nói")
                tab_spk, tab_non_spk = st.tabs(["🎭 Nhận diện là NGƯỜI NÓI", "🚫 Đang bị xem là TỪ NHIỄU"])
                
                with tab_spk:
                    if detected_speakers:
                        st.write(", ".join([f"`{s}`" for s in detected_speakers]))
                        to_move_to_ns = st.multiselect(
                            "Phát hiện từ nào bị nhận diện sai? Chọn để LƯU VÀO DATABASE TỪ NHIỄU:",
                            options=[name for name in candidates.keys() if name.upper() not in NON_SPEAKER_PHRASES],
                            key="select_to_ns"
                        )
                        if st.button("➡️ Đưa vào Database TỪ NHIỄU", type="secondary"):
                            if to_move_to_ns:
                                new_items = [item.upper() for item in to_move_to_ns]
                                st.session_state['custom_non_speakers'].update(new_items)
                                save_json_db(NON_SPEAKER_DB_FILE, st.session_state['custom_non_speakers'])
                                st.success(f"✅ Đã lưu {len(new_items)} từ vào Database Từ Nhiễu!")
                                time.sleep(1); st.rerun()
                    else: st.info("Chưa tìm thấy cụm từ người nói nào.")

                with tab_non_spk:
                    if detected_non_speakers:
                        st.write(", ".join([f"`{s}`" for s in detected_non_speakers]))
                        to_move_to_spk = st.multiselect(
                            "Từ nào thực ra là NGƯỜI NÓI? Chọn để LƯU VÀO DATABASE NGƯỜI NÓI:",
                            options=[name for name in candidates.keys() if name.upper() in NON_SPEAKER_PHRASES],
                            key="select_to_spk"
                        )
                        if st.button("➡️ Đưa vào Database NGƯỜI NÓI", type="secondary"):
                            if to_move_to_spk:
                                st.session_state['custom_speakers'].update(to_move_to_spk)
                                save_json_db(SPEAKER_DB_FILE, st.session_state['custom_speakers'])
                                for item in to_move_to_spk: st.session_state['custom_non_speakers'].discard(item.upper())
                                save_json_db(NON_SPEAKER_DB_FILE, st.session_state['custom_non_speakers'])
                                st.success(f"✅ Đã lưu {len(to_move_to_spk)} tên vào Database Người Nói!")
                                time.sleep(1); st.rerun()
                    else: st.info("Không có cụm từ nào bị loại vào danh sách từ nhiễu.")

            with st.container(border=True):
                st.markdown("### 🗣️ Từ Tiếng Anh Xuất Hiện Trong Kịch Bản")
                st.caption("Quét và điều chỉnh phiên âm riêng cho kịch bản này (Đã qua bộ lọc thông minh):")

                detected_eng_words = scan_english_words_in_dialogue(uploaded_file, custom_spks, custom_non_spks)

                if detected_eng_words:
                    st.markdown("#### 🔊 Trình nghe phát âm chuẩn giọng bản xứ (Google US/UK)")
                    col_listen1, col_listen2, col_listen3 = st.columns([2.5, 1.5, 1.5])
                    
                    with col_listen1:
                        word_to_listen = st.selectbox("Chọn từ cần nghe phát âm:", options=detected_eng_words, key="script_listen_select")
                    with col_listen2:
                        accent_choice = st.radio("Giọng phát âm:", options=["Giọng Mỹ (US)", "Giọng Anh (UK)"], horizontal=True, key="script_accent_radio")
                    with col_listen3:
                        st.markdown("<div style='margin-top: 28px;'></div>", unsafe_allow_html=True)
                        listen_btn = st.button("🔊 Nghe phát âm", type="secondary", use_container_width=True)

                    if listen_btn and word_to_listen:
                        tld = 'com' if "Mỹ" in accent_choice else 'co.uk'
                        audio_fp = generate_english_audio(word_to_listen, accent=tld)
                        if audio_fp: st.audio(audio_fp, format="audio/mp3", autoplay=True)

                    st.markdown("---")
                    
                    table_data = []
                    for word in detected_eng_words:
                        current_pho = st.session_state['custom_phonetics'].get(word.upper(), word)
                        table_data.append({
                            "Từ Tiếng Anh": word,
                            "Phiên âm hiện tại": current_pho,
                            "Đề xuất chỉnh sửa của bạn": current_pho,
                            "Nạp vào Database": True
                        })

                    df_eng = pd.DataFrame(table_data)

                    edited_df = st.data_editor(
                        df_eng,
                        column_config={
                            "Từ Tiếng Anh": st.column_config.TextColumn("Từ Tiếng Anh gốc", disabled=True),
                            "Phiên âm hiện tại": st.column_config.TextColumn("Phiên âm gán hiện tại", disabled=True),
                            "Đề xuất chỉnh sửa của bạn": st.column_config.TextColumn("Đề xuất phiên âm mới"),
                            "Nạp vào Database": st.column_config.CheckboxColumn("Lưu Database?", default=True)
                        },
                        disabled=["Từ Tiếng Anh", "Phiên âm hiện tại"],
                        hide_index=True,
                        use_container_width=True,
                        key="phonetic_script_table"
                    )

                    if st.button("💾 Nạp chỉnh sửa kịch bản này vào Database Phiên Âm", type="secondary", use_container_width=True):
                        updated_count = 0
                        for _, row in edited_df.iterrows():
                            if row["Nạp vào Database"]:
                                eng_k = str(row["Từ Tiếng Anh"]).upper().strip()
                                pho_v = str(row["Đề xuất chỉnh sửa của bạn"]).strip()
                                if pho_v:
                                    st.session_state['custom_phonetics'][eng_k] = pho_v
                                    updated_count += 1
                        
                        save_json_db(PHONETIC_DB_FILE, st.session_state['custom_phonetics'])
                        st.success(f"✅ Đã cập nhật {updated_count} từ phiên âm vào Database!")
                        time.sleep(1); st.rerun()
                else: st.info("Không phát hiện từ Tiếng Anh / Tên riêng nước ngoài nào trong phần lời thoại kịch bản này.")

            st.markdown("---")
            if st.button("✨ 2. BẮT ĐẦU ĐỊNH DẠNG TỰ ĐỘNG", use_container_width=True, type="primary"):
                try:
                    modified_docx, ass_f, srt_f, act_zip, stats = process_docx(uploaded_file, file_name_without_ext, enable_colors, enable_phonetic, enable_cast, is_resync=False, font_size_pt=12)
                    
                    st.session_state['processed_docx'] = modified_docx
                    st.session_state['processed_ass'] = ass_f
                    st.session_state['processed_srt'] = srt_f
                    st.session_state['actor_zip'] = act_zip
                    st.session_state['docx_name'] = clean_file_name_for_output(original_filename, tag="_edit", ext=".docx")
                    st.session_state['ass_name'] = clean_file_name_for_output(original_filename, tag="_edit", ext=".ass")
                    st.session_state['srt_name'] = clean_file_name_for_output(original_filename, tag="_edit", ext=".srt")
                    st.session_state['zip_name'] = clean_file_name_for_output(original_filename, tag="_KichBan_TachVai", ext=".zip")
                    st.session_state['stats'] = stats
                    
                except Exception as e: st.error(f"Đã có lỗi xảy ra: {e}")

            if 'processed_docx' in st.session_state:
                st.markdown("---")
                qc_warns = st.session_state['stats'].get("qc_warnings", [])
                if qc_warns:
                    with st.expander("🔍 BÁO CÁO CẢNH BÁO CHẤT LƯỢNG (QC & CPS CHECKER)", expanded=True):
                        st.caption("Danh sách cảnh báo về tốc độ đọc thoại hoặc gán phân vai để BTV rà soát:")
                        for w in qc_warns[:10]: st.markdown(f"<div class='qc-card-warning'>{w}</div>", unsafe_allow_html=True)
                        if len(qc_warns) > 10: st.info(f"...và thêm {len(qc_warns)-10} cảnh báo khác.")
                
                st.markdown("### ⬇️ 3. TẢI VỀ CÁC FILE ĐÃ XỬ LÝ HOÀN HẢO")
                col_dl1, col_dl2, col_dl3 = st.columns(3)
                with col_dl1:
                    st.download_button(
                        label="📄 FILE WORD KỊCH BẢN (.DOCX)",
                        data=st.session_state['processed_docx'],
                        file_name=st.session_state['docx_name'],
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary", use_container_width=True
                    )
                with col_dl2:
                    st.download_button(
                        label="🎬 PHỤ ĐỀ CAO CẤP (.ASS)",
                        data=st.session_state['processed_ass'],
                        file_name=st.session_state['ass_name'],
                        mime="text/plain", use_container_width=True
                    )
                with col_dl3:
                    st.download_button(
                        label="📝 PHỤ ĐỀ CHUẨN (.SRT)",
                        data=st.session_state['processed_srt'],
                        file_name=st.session_state['srt_name'],
                        mime="text/plain", use_container_width=True
                    )
                    
                st.markdown("---")
                st.markdown("#### 🎙️ KỊCH BẢN TÁCH VAI RIÊNG CHO PHÒNG THU LỒNG TIẾNG")
                st.caption("Mỗi diễn viên chỉ nhận đúng câu thoại của mình, giúp thu âm nhanh và không xao nhãng:")
                
                act_map = st.session_state['stats'].get("actor_dialogue_map", {})
                if act_map:
                    col_act1, col_act2 = st.columns([2, 1])
                    with col_act1:
                        selected_actor = st.selectbox("Chọn Diễn viên lồng tiếng để tải file riêng:", options=list(act_map.keys()))
                        if selected_actor:
                            act_buf = generate_actor_docx(st.session_state['stats']['video_title'], selected_actor, act_map[selected_actor], font_size_pt=12)
                            st.download_button(
                                label=f"⬇️ TẢI FILE WORD RIÊNG CHO {selected_actor} (.DOCX)",
                                data=act_buf,
                                file_name=f"KichBan_{selected_actor}_{st.session_state['stats']['video_title']}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                    with col_act2:
                        st.markdown("<div style='margin-top: 28px;'></div>", unsafe_allow_html=True)
                        st.download_button(
                            label="📦 TẢI TRỌN BỘ KỊCH BẢN TÁCH VAI (.ZIP)",
                            data=st.session_state['actor_zip'],
                            file_name=st.session_state['zip_name'],
                            mime="application/zip", type="secondary", use_container_width=True
                        )
                st.balloons()

    with col2:
        st.markdown("### 📊 SaaS Analytics")
        if 'stats' in st.session_state:
            stats = st.session_state['stats']
            st.markdown(f"""
            <div class="metric-card" style="margin-bottom: 12px;">
                <div class="metric-label">🎭 Tổng số Nhân vật</div>
                <div class="metric-value">{stats["total_speakers"]}</div>
            </div>
            <div class="metric-card" style="margin-bottom: 12px;">
                <div class="metric-label">💬 Tổng số Câu thoại</div>
                <div class="metric-value">{stats["total_lines"]}</div>
            </div>
            <div class="metric-card" style="margin-bottom: 12px;">
                <div class="metric-label">⏱️ Độ dài Video</div>
                <div class="metric-value">{stats["video_duration_min"]} phút</div>
            </div>
            """, unsafe_allow_html=True)
            top_name, top_count = stats["top_speaker"]
            st.info(f"👑 **Nhân vật thoại nhiều nhất:** \n\n**{top_name}** với {top_count} câu thoại.")
        else: st.info("Bảng phân tích dữ liệu kịch bản sẽ xuất hiện tại đây sau khi bạn xử lý file.")

# ==========================================
# TAB 2: RE-SYNC KỊCH BẢN ĐÃ BIÊN TẬP (CỠ PHÔNG 14PT)
# ==========================================
with tab_resync:
    col_r1, col_r2 = st.columns([1.6, 1])
    
    with col_r1:
        with st.container(border=True):
            st.markdown("### 🔄 Tải lên file Kịch bản ĐÃ BIÊN TẬP THỦ CÔNG (.docx)")
            st.caption("Dành riêng cho file kịch bản đã được team biên tập chỉnh sửa lời thoại. Tự động phục hồi màu sắc, phân vai và **xuất phông chữ 14pt cực nét** cho phòng thu.")
            
            resync_file = st.file_uploader(
                "Kéo thả file .docx đã biên tập vào đây", 
                type=['docx'], 
                key=f"resync_uploader_{st.session_state['resync_uploader_key']}"
            )
            project_week_input = st.text_input("📌 Gán Tuần Dự Án cho video này:", value="Tuần 1", help="VD: Tuần 1, Tuần 2, Tuần 1 - Đợt Phim A...")
            
        if resync_file is not None:
            r_filename = resync_file.name
            r_name_no_ext = os.path.splitext(r_filename)[0]
            st.success(f"📄 Đã nhận file kịch bản biên tập: **{r_filename}**")
            
            st.markdown("---")
            if st.button("✨ 2. BẮT ĐẦU RE-SYNC & CHUẨN HÓA LẠI ĐỊNH DẠNG (14PT)", use_container_width=True, type="primary", key="btn_resync_start"):
                try:
                    # Chú ý: font_size_pt=14 dành riêng cho Re-Sync theo yêu cầu
                    r_docx, r_ass, r_srt, r_zip, r_stats = process_docx(resync_file, r_name_no_ext, enable_colors, enable_phonetic, enable_cast, is_resync=True, font_size_pt=14)
                    
                    st.session_state['r_processed_docx'] = r_docx
                    st.session_state['r_processed_ass'] = r_ass
                    st.session_state['r_processed_srt'] = r_srt
                    st.session_state['r_actor_zip'] = r_zip
                    st.session_state['r_docx_name'] = clean_file_name_for_output(r_filename, tag="_final", ext=".docx")
                    st.session_state['r_ass_name'] = clean_file_name_for_output(r_filename, tag="_final", ext=".ass")
                    st.session_state['r_srt_name'] = clean_file_name_for_output(r_filename, tag="_final", ext=".srt")
                    st.session_state['r_zip_name'] = clean_file_name_for_output(r_filename, tag="_KichBan_TachVai_Final", ext=".zip")
                    st.session_state['resync_stats'] = r_stats
                    
                    video_title = r_stats.get("video_title", r_name_no_ext)
                    actors_list = r_stats.get("actors_list", [])
                    actors_str = ", ".join(actors_list) if actors_list else "Chưa có thông tin"
                    actor_breakdown = r_stats.get("actor_stats_breakdown", {})
                    total_lines = r_stats.get("total_lines", 0)
                    video_dur_min = r_stats.get("video_duration_min", 1)
                    
                    today_str = time.strftime("%d/%m/%Y")
                    assigned_week = project_week_input.strip() if project_week_input else "Tuần 1"
                    
                    tracker_list = st.session_state['dubbing_tracker']
                    existing_entry = next((item for item in tracker_list if item['video_title'].upper() == video_title.upper()), None)
                    
                    curr_def_rate = st.session_state['payroll_rates'].get("unit_rate", 30000)
                    custom_actor_rates = {a.upper(): curr_def_rate for a in actors_list}
                    
                    entry_data = {
                        "video_title": video_title, "actors": actors_str, "actor_breakdown": actor_breakdown,
                        "total_lines": total_lines, "video_duration_min": video_dur_min, "date": today_str,
                        "project_week": assigned_week, "custom_actor_rates": custom_actor_rates
                    }
                    
                    if existing_entry:
                        if "custom_actor_rates" in existing_entry: entry_data["custom_actor_rates"] = existing_entry["custom_actor_rates"]
                        existing_entry.update(entry_data)
                    else: tracker_list.append(entry_data)
                    
                    st.session_state['dubbing_tracker'] = tracker_list
                    save_json_db(TRACKER_DB_FILE, tracker_list)
                    
                except Exception as e: st.error(f"Lỗi xảy ra khi Re-Sync: {e}")
                    
            if 'r_processed_docx' in st.session_state:
                st.markdown("---")
                r_qc_warns = st.session_state['resync_stats'].get("qc_warnings", [])
                if r_qc_warns:
                    with st.expander("🔍 BÁO CÁO CẢNH BÁO CHẤT LƯỢNG (QC & CPS CHECKER)", expanded=True):
                        st.caption("Danh sách cảnh báo về tốc độ đọc thoại hoặc gán phân vai để BTV rà soát:")
                        for w in r_qc_warns[:10]: st.markdown(f"<div class='qc-card-warning'>{w}</div>", unsafe_allow_html=True)
                        if len(r_qc_warns) > 10: st.info(f"...và thêm {len(r_qc_warns)-10} cảnh báo khác.")
                
                st.markdown("### ⬇️ 3. TẢI VỀ CÁC FILE CHUẨN HOÀN HẢO (PHÔNG 14PT)")
                col_rdl1, col_rdl2, col_rdl3 = st.columns(3)
                with col_rdl1:
                    st.download_button(
                        label="📄 FILE WORD KỊCH BẢN (14PT)", data=st.session_state['r_processed_docx'],
                        file_name=st.session_state['r_docx_name'], mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary", use_container_width=True, key="btn_resync_dl_docx"
                    )
                with col_rdl2:
                    st.download_button(
                        label="🎬 PHỤ ĐỀ CAO CẤP (.ASS)", data=st.session_state['r_processed_ass'],
                        file_name=st.session_state['r_ass_name'], mime="text/plain", use_container_width=True, key="btn_resync_dl_ass"
                    )
                with col_rdl3:
                    st.download_button(
                        label="📝 PHỤ ĐỀ CHUẨN (.SRT)", data=st.session_state['r_processed_srt'],
                        file_name=st.session_state['r_srt_name'], mime="text/plain", use_container_width=True, key="btn_resync_dl_srt"
                    )
                    
                st.markdown("---")
                st.markdown("#### 🎙️ KỊCH BẢN TÁCH VAI RIÊNG CHO PHÒNG THU LỒNG TIẾNG (14PT)")
                st.caption("Mỗi diễn viên chỉ nhận đúng câu thoại của mình, giúp thu âm nhanh và không xao nhãng:")
                
                r_act_map = st.session_state['resync_stats'].get("actor_dialogue_map", {})
                if r_act_map:
                    col_ract1, col_ract2 = st.columns([2, 1])
                    with col_ract1:
                        r_selected_actor = st.selectbox("Chọn Diễn viên lồng tiếng để tải file riêng:", options=list(r_act_map.keys()), key="resync_select_actor")
                        if r_selected_actor:
                            r_act_buf = generate_actor_docx(st.session_state['resync_stats']['video_title'], r_selected_actor, r_act_map[r_selected_actor], font_size_pt=14)
                            st.download_button(
                                label=f"⬇️ TẢI FILE WORD RIÊNG CHO {r_selected_actor} (14PT)", data=r_act_buf,
                                file_name=f"KichBan_{r_selected_actor}_{st.session_state['resync_stats']['video_title']}_Final.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, key="btn_dl_single_actor_resync"
                            )
                    with col_ract2:
                        st.markdown("<div style='margin-top: 28px;'></div>", unsafe_allow_html=True)
                        st.download_button(
                            label="📦 TẢI TRỌN BỘ KỊCH BẢN TÁCH VAI (.ZIP)", data=st.session_state['r_actor_zip'],
                            file_name=st.session_state['r_zip_name'], mime="application/zip", type="secondary", use_container_width=True, key="btn_dl_zip_actor_resync"
                        )
                st.balloons()
        else: st.info("📌 **Hãy tải file kịch bản đã qua chỉnh sửa thủ công để hệ thống phục hồi lại định dạng chuẩn.**")

    with col_r2:
        st.markdown("### 📊 Re-Sync Analytics")
        if 'resync_stats' in st.session_state:
            r_stats = st.session_state['resync_stats']
            st.markdown(f"""
            <div class="metric-card" style="margin-bottom: 12px;">
                <div class="metric-label">🎭 Tổng số Nhân vật</div>
                <div class="metric-value">{r_stats["total_speakers"]}</div>
            </div>
            <div class="metric-card" style="margin-bottom: 12px;">
                <div class="metric-label">💬 Tổng số Câu thoại</div>
                <div class="metric-value">{r_stats["total_lines"]}</div>
            </div>
            <div class="metric-card" style="margin-bottom: 12px;">
                <div class="metric-label">⏱️ Độ dài Video</div>
                <div class="metric-value">{r_stats["video_duration_min"]} phút</div>
            </div>
            """, unsafe_allow_html=True)
            top_name, top_count = r_stats["top_speaker"]
            st.info(f"👑 **Nhân vật thoại nhiều nhất:** \n\n**{top_name}** với {top_count} câu thoại.")
        else: st.info("Thống kê file Re-Sync sẽ xuất hiện tại đây sau khi hoàn tất.")

# ==========================================
# TAB 3: THEO DÕI & BÁO CÁO LƯƠNG
# ==========================================
with tab_dub_tracker:
    st.subheader("📋 BÁO CÁO BẢNG TÍNH LƯƠNG LỒNG TIẾNG THEO TUẦN DỰ ÁN")
    
    with st.expander("⚙️ CẤU HÌNH ĐƠN GIÁ MẶC ĐỊNH SẢN XUẤT", expanded=False):
        c_rate1, c_rate2 = st.columns([2, 3])
        curr_mode = st.session_state['payroll_rates'].get("mode", "minute")
        mode_idx = 0 if curr_mode == "minute" else (1 if curr_mode == "line" else 2)
        
        with c_rate1:
            rate_mode_choice = st.radio(
                "Cách tính thù lao:", options=["Theo Phút video (phút)", "Theo Câu thoại (câu)", "Theo Số từ (từ)"], 
                index=mode_idx, key="radio_payroll_mode"
            )
        with c_rate2:
            default_unit_rate = st.number_input(
                "Đơn giá mặc định khi thêm video mới (VNĐ):", value=int(st.session_state['payroll_rates'].get("unit_rate", 30000)), 
                step=5000, key="num_payroll_unit_rate"
            )
            unit_label = "phút" if "Phút" in rate_mode_choice else ("câu" if "Câu" in rate_mode_choice else "từ")
            st.caption(f"👉 **Đơn giá mặc định áp dụng:** `{default_unit_rate:,.0f}` VNĐ / {unit_label}")
            
        if "Phút" in rate_mode_choice: new_mode = "minute"
        elif "Câu" in rate_mode_choice: new_mode = "line"
        else: new_mode = "word"
            
        if new_mode != st.session_state['payroll_rates'].get("mode") or default_unit_rate != st.session_state['payroll_rates'].get("unit_rate"):
            st.session_state['payroll_rates'] = {"mode": new_mode, "unit_rate": default_unit_rate}
            save_json_db(RATES_DB_FILE, st.session_state['payroll_rates'])

    current_mode = st.session_state['payroll_rates'].get("mode", "minute")
    current_rate = st.session_state['payroll_rates'].get("unit_rate", 30000)

    subtab_video, subtab_custom_rates, subtab_payroll = st.tabs([
        "📹 Bảng Theo Dõi Video Theo Tuần Dự Án", "💵 Chỉnh Đơn Giá Cá Nhân (Video x Diễn viên)", "💰 Báo Cáo Lương Chi Tiết Từng Diễn Viên"
    ])

    tracker_data = st.session_state['dubbing_tracker']

    with subtab_video:
        st.markdown("#### Bảng quản lý danh sách Video và Thù lao (Định dạng chuẩn 6 cột)")
        if tracker_data:
            formatted_editor_list = []
            for idx, item in enumerate(tracker_data, 1):
                p_week = item.get("project_week", "Tuần 1")
                v_dur_min = int(item.get("video_duration_min", 1))
                v_lines = item.get("total_lines", 0)
                v_acts = item.get("actors", "")
                bd = item.get("actor_breakdown", {})
                custom_rates = item.get("custom_actor_rates", {})
                
                raw_acts = [a.strip().upper() for a in v_acts.split(',') if a.strip() and a.strip() != "CHƯA CÓ THÔNG TIN"]
                
                if current_mode == "minute":
                    v_pay = sum(v_dur_min * custom_rates.get(act, current_rate) for act in raw_acts) if raw_acts else v_dur_min * current_rate
                    dur_str = f"{v_dur_min} phút"
                elif current_mode == "line":
                    v_pay = sum((bd.get(act, {}).get("lines", v_lines)) * custom_rates.get(act, current_rate) for act in raw_acts) if raw_acts else v_lines * current_rate
                    dur_str = f"{v_lines} câu"
                else:
                    tot_words = sum(a["words"] for a in bd.values()) if bd else 0
                    v_pay = sum((bd.get(act, {}).get("words", 0)) * custom_rates.get(act, current_rate) for act in raw_acts) if raw_acts else tot_words * current_rate
                    dur_str = f"{tot_words} từ"
                    
                formatted_editor_list.append({
                    "Stt": idx, "Tuần dự án": p_week, "Tiêu đề video": item['video_title'],
                    "Thời lượng": dur_str, "Thành tiền Video": f"{v_pay:,.0f} VNĐ", "Diễn viên": v_acts, "Xóa": False
                })

            df_editor_raw = pd.DataFrame(formatted_editor_list)

            edited_tracker_df = st.data_editor(
                df_editor_raw,
                column_config={
                    "Stt": st.column_config.NumberColumn("Stt", disabled=True, width="small"),
                    "Tuần dự án": st.column_config.TextColumn("Tuần dự án (Sửa trực tiếp)"),
                    "Tiêu đề video": st.column_config.TextColumn("Tiêu đề video (Sửa trực tiếp)"),
                    "Thời lượng": st.column_config.TextColumn("Thời lượng", disabled=True),
                    "Thành tiền Video": st.column_config.TextColumn("Thành tiền (Tự động tính theo giá riêng)", disabled=True),
                    "Diễn viên": st.column_config.TextColumn("Diễn viên (Sửa trực tiếp)"),
                    "Xóa": st.column_config.CheckboxColumn("Xóa?")
                },
                hide_index=True, use_container_width=True, key="dubbing_tracker_editor_main"
            )

            col_tr1, col_tr2 = st.columns([1, 1])
            with col_tr1:
                if st.button("💾 LƯU THAY ĐỔI TRÊN BẢNG VIDEO", type="primary", use_container_width=True):
                    new_tracker = []; deleted_cnt = 0
                    for idx_r, row in edited_tracker_df.iterrows():
                        if row["Xóa"]: deleted_cnt += 1
                        else:
                            orig_item = tracker_data[idx_r]
                            orig_item['project_week'] = str(row["Tuần dự án"]).strip()
                            orig_item['video_title'] = str(row["Tiêu đề video"]).strip()
                            
                            new_acts_raw = [a.strip().upper() for a in str(row["Diễn viên"]).split(',') if a.strip() and a.strip() != "CHƯA CÓ THÔNG TIN"]
                            orig_item['actors'] = ", ".join(new_acts_raw) if new_acts_raw else "CHƯA CÓ THÔNG TIN"
                            
                            old_bd = orig_item.get("actor_breakdown", {}); new_bd = {}
                            c_rates = orig_item.get("custom_actor_rates", {})
                            for act in new_acts_raw:
                                if act in old_bd: new_bd[act] = old_bd[act]
                                else: new_bd[act] = {"lines": 0, "words": 0}
                                if act not in c_rates: c_rates[act] = current_rate
                            
                            orig_item['actor_breakdown'] = new_bd
                            orig_item['custom_actor_rates'] = c_rates
                            new_tracker.append(orig_item)
                            
                    st.session_state['dubbing_tracker'] = new_tracker
                    save_json_db(TRACKER_DB_FILE, new_tracker)
                    st.success(f"✅ Đã lưu thay đổi & đồng bộ dữ liệu!"); time.sleep(1); st.rerun()

            st.markdown("---")
            st.markdown("#### 📑 Xem theo Bảng chia Tuần Dự Án (Có Hàng Tổng)")

            project_weeks_map = {}
            for item in tracker_data:
                pw = item.get("project_week", "Tuần 1")
                if pw not in project_weeks_map: project_weeks_map[pw] = []
                project_weeks_map[pw].append(item)

            total_studio_money = 0; excel_sheets_data = {}

            for pw_name, items in sorted(project_weeks_map.items()):
                st.markdown(f"##### 📅 {pw_name.upper()}")
                week_rows = []; week_tot_dur = 0; week_tot_money = 0
                
                for idx, item in enumerate(items, 1):
                    v_title = item['video_title']; v_dur_min = int(item.get("video_duration_min", 1))
                    v_lines = item.get("total_lines", 0); v_actors = item.get("actors", "")
                    bd = item.get("actor_breakdown", {}); custom_rates = item.get("custom_actor_rates", {})
                    
                    raw_acts = [a.strip().upper() for a in v_actors.split(',') if a.strip() and a.strip() != "CHƯA CÓ THÔNG TIN"]
                    
                    if current_mode == "minute":
                        v_pay = sum(v_dur_min * custom_rates.get(act, current_rate) for act in raw_acts) if raw_acts else v_dur_min * current_rate
                        dur_str = f"{v_dur_min} phút"
                    elif current_mode == "line":
                        v_pay = sum((bd.get(act, {}).get("lines", v_lines)) * custom_rates.get(act, current_rate) for act in raw_acts) if raw_acts else v_lines * current_rate
                        dur_str = f"{v_lines} câu"
                    else:
                        tot_words = sum(a["words"] for a in bd.values()) if bd else 0
                        v_pay = sum((bd.get(act, {}).get("words", 0)) * custom_rates.get(act, current_rate) for act in raw_acts) if raw_acts else tot_words * current_rate
                        dur_str = f"{tot_words} từ"
                        
                    week_tot_dur += v_dur_min; week_tot_money += v_pay
                    
                    week_rows.append({
                        "Stt": idx, "Tiêu đề video": v_title, "Thời lượng": dur_str,
                        "Thành tiền": f"{v_pay:,.0f} VNĐ", "Diễn viên": v_actors
                    })

                total_studio_money += week_tot_money
                week_rows.append({
                    "Stt": "TỔNG", "Tiêu đề video": f"TỔNG CỘNG ({len(items)} video)",
                    "Thời lượng": f"{week_tot_dur} phút" if current_mode == "minute" else "-",
                    "Thành tiền": f"{week_tot_money:,.0f} VNĐ", "Diễn viên": "-"
                })

                df_week_display = pd.DataFrame(week_rows)
                st.dataframe(df_week_display, hide_index=True, use_container_width=True)
                excel_sheets_data[pw_name[:30]] = df_week_display

            with col_tr2:
                excel_payroll_buffer = io.BytesIO()
                with pd.ExcelWriter(excel_payroll_buffer, engine='openpyxl') as writer:
                    for s_name, s_df in excel_sheets_data.items():
                        s_df.to_excel(writer, index=False, sheet_name=s_name.replace(":", "_").replace("/", "_"))
                excel_payroll_buffer.seek(0)

                st.download_button(
                    label="📊 XUẤT TẤT CẢ TUẦN RA EXCEL (.XLSX)", data=excel_payroll_buffer,
                    file_name="Theo_Doi_Video_Long_Tieng_MaiHan.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True
                )
        else: st.info("Chưa có dữ liệu video nào. Hãy chạy Re-Sync ở Tab 2 để tự động ghi nhận video mới!")

    with subtab_custom_rates:
        st.markdown("#### 💵 Bảng Điều Chỉnh Đơn Giá Cá Nhân (Theo Video & Diễn viên)")
        st.caption("Gõ đơn giá riêng cho từng người trong từng video (Ví dụ: Khánh video A 30k, video B 20k; Trúc video A 15k...):")
        
        if tracker_data:
            rate_editor_rows = []
            for v_idx, item in enumerate(tracker_data):
                pw = item.get("project_week", "Tuần 1"); v_title = item['video_title']
                v_dur = int(item.get("video_duration_min", 1)); v_acts = item.get("actors", "")
                custom_rates = item.get("custom_actor_rates", {})
                acting_actors = [a.strip().upper() for a in v_acts.split(',') if a.strip() and a.strip() != "CHƯA CÓ THÔNG TIN"]
                
                for act in acting_actors:
                    act_rate = custom_rates.get(act, current_rate)
                    rate_editor_rows.append({
                        "v_idx": v_idx, "Tuần dự án": pw, "Tiêu đề video": v_title,
                        "Thời lượng (Phút)": v_dur, "Diễn viên": act, "Đơn giá cá nhân (VNĐ)": int(act_rate)
                    })

            if rate_editor_rows:
                df_custom_rates = pd.DataFrame(rate_editor_rows)
                edited_custom_rates_df = st.data_editor(
                    df_custom_rates[["Tuần dự án", "Tiêu đề video", "Thời lượng (Phút)", "Diễn viên", "Đơn giá cá nhân (VNĐ)"]],
                    column_config={
                        "Tuần dự án": st.column_config.TextColumn("Tuần", disabled=True),
                        "Tiêu đề video": st.column_config.TextColumn("Tiêu đề video", disabled=True),
                        "Thời lượng (Phút)": st.column_config.NumberColumn("Độ dài (Phút)", disabled=True),
                        "Diễn viên": st.column_config.TextColumn("Diễn viên", disabled=True),
                        "Đơn giá cá nhân (VNĐ)": st.column_config.NumberColumn("Đơn giá riêng (Gõ để sửa)", format="%d", step=1000)
                    },
                    hide_index=True, use_container_width=True, key="custom_rates_editor_table_main"
                )

                if st.button("💾 LƯU ĐƠN GIÁ CÁ NHÂN", type="primary", use_container_width=True):
                    for idx_r, row in edited_custom_rates_df.iterrows():
                        v_i = df_custom_rates.iloc[idx_r]["v_idx"]; act_n = df_custom_rates.iloc[idx_r]["Diễn viên"]
                        new_r = float(row["Đơn giá cá nhân (VNĐ)"])
                        if "custom_actor_rates" not in tracker_data[v_i]: tracker_data[v_i]["custom_actor_rates"] = {}
                        tracker_data[v_i]["custom_actor_rates"][act_n] = new_r
                        
                    st.session_state['dubbing_tracker'] = tracker_data
                    save_json_db(TRACKER_DB_FILE, tracker_data)
                    st.success("✅ Đã lưu đơn giá riêng cho từng diễn viên thành công!"); time.sleep(1); st.rerun()
            else: st.info("Chưa có thông tin diễn viên lồng tiếng trong danh sách video.")
        else: st.info("Chưa có danh sách video để điều chỉnh đơn giá.")

    with subtab_payroll:
        st.markdown("#### Bảng Tổng kết & Báo cáo Lương Chi Tiết Từng Diễn Viên")
        if tracker_data:
            col_filter1, col_filter2 = st.columns([1, 1])
            all_project_weeks = sorted(list(set(item.get("project_week", "Tuần 1") for item in tracker_data)))
            with col_filter1: selected_week_filter = st.selectbox("🎯 Chọn Tuần Dự Án:", options=["TẤT CẢ CÁC TUẦN"] + all_project_weeks)

            if selected_week_filter == "TẤT CẢ CÁC TUẦN": active_tracker_data = tracker_data
            else: active_tracker_data = [item for item in tracker_data if item.get("project_week", "Tuần 1") == selected_week_filter]

            actor_weekly_map = {}
            for item in active_tracker_data:
                v_title = item['video_title']; v_dur = int(item.get("video_duration_min", 1))
                bd = item.get("actor_breakdown", {}); v_lines = item.get("total_lines", 0)
                custom_rates = item.get("custom_actor_rates", {})
                acting_actors = [a.strip().upper() for a in item.get("actors", "").split(",") if a.strip() and a.strip() != "CHƯA CÓ THÔNG TIN"]

                for act_name_clean in acting_actors:
                    if act_name_clean not in actor_weekly_map:
                        actor_weekly_map[act_name_clean] = {
                            "videos_count": 0, "total_video_mins": 0, "total_lines": 0, "total_words": 0, "video_rows": []
                        }
                    
                    act_rate = custom_rates.get(act_name_clean, current_rate)
                    act_lines = bd[act_name_clean]["lines"] if (bd and act_name_clean in bd) else v_lines
                    act_words = bd[act_name_clean]["words"] if (bd and act_name_clean in bd) else 0
                    
                    if current_mode == "minute": act_pay = v_dur * act_rate; dur_str = f"{v_dur} phút"
                    elif current_mode == "line": act_pay = act_lines * act_rate; dur_str = f"{act_lines} câu"
                    else: act_pay = act_words * act_rate; dur_str = f"{act_words} từ"

                    actor_weekly_map[act_name_clean]["videos_count"] += 1
                    actor_weekly_map[act_name_clean]["total_video_mins"] += v_dur
                    actor_weekly_map[act_name_clean]["total_lines"] += act_lines
                    actor_weekly_map[act_name_clean]["total_words"] += act_words
                    actor_weekly_map[act_name_clean]["video_rows"].append({
                        "Stt": len(actor_weekly_map[act_name_clean]["video_rows"]) + 1,
                        "Tiêu đề video": v_title, "Thời lượng": dur_str,
                        "Đơn giá": f"{act_rate:,.0f} VNĐ", "Thành tiền": f"{act_pay:,.0f} VNĐ", "Pay_Num": act_pay
                    })

            all_actor_names = sorted(list(actor_weekly_map.keys()))
            with col_filter2:
                selected_actor_view = st.selectbox("👤 Chọn Diễn viên để xem & xuất Báo cáo Lương cá nhân:", options=["TẤT CẢ DIỄN VIÊN"] + all_actor_names)

            st.markdown("---")

            if selected_actor_view != "TẤT CẢ DIỄN VIÊN":
                a_info = actor_weekly_map[selected_actor_view]; a_rows = a_info["video_rows"]
                a_tot_pay = sum(r["Pay_Num"] for r in a_rows)
                st.subheader(f"👤 PHIẾU BÁO CÁO THÙ LAO: {selected_actor_view}")
                st.caption(f"Dữ liệu thù lao lồng tiếng cho {selected_actor_view} ({selected_week_filter})")
                
                df_single_actor = pd.DataFrame(a_rows)[["Stt", "Tiêu đề video", "Thời lượng", "Đơn giá", "Thành tiền"]]
                st.dataframe(df_single_actor, hide_index=True, use_container_width=True)
                st.metric(f"💰 TỔNG THÙ LAO DỰ KIẾN TRẢ CHO {selected_actor_view}:", f"{a_tot_pay:,.0f} VNĐ")
                
                col_p_btn1, col_p_btn2 = st.columns(2)
                with col_p_btn1:
                    actor_docx_buf = generate_actor_salary_slip_docx(selected_actor_view, selected_week_filter, a_rows, a_tot_pay, current_mode)
                    st.download_button(
                        label=f"🖨️ IN / TẢI PHIẾU LƯƠNG WORD CỦA {selected_actor_view} (.DOCX)",
                        data=actor_docx_buf, file_name=f"PhieuLuong_{selected_actor_view}_{selected_week_filter}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary", use_container_width=True
                    )
                with col_p_btn2:
                    excel_single_buf = io.BytesIO()
                    with pd.ExcelWriter(excel_single_buf, engine='openpyxl') as writer: df_single_actor.to_excel(writer, index=False, sheet_name="Phieu Luong")
                    excel_single_buf.seek(0)
                    st.download_button(
                        label=f"📊 TẢI PHIẾU LƯƠNG CÁ NHÂN EXCEL (.XLSX)",
                        data=excel_single_buf, file_name=f"PhieuLuong_{selected_actor_view}_{selected_week_filter}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True
                    )
            else:
                actor_payroll_rows = []; grand_actor_pay = 0; unique_video_titles = set(); grand_unique_mins = 0
                for item in active_tracker_data:
                    if item['video_title'] not in unique_video_titles:
                        unique_video_titles.add(item['video_title'])
                        grand_unique_mins += int(item.get("video_duration_min", 1))

                for idx, (act_name, info) in enumerate(sorted(actor_weekly_map.items()), 1):
                    if current_mode == "minute": unit_cnt = info["total_video_mins"]; dur_disp = f"{unit_cnt} phút"
                    elif current_mode == "line": unit_cnt = info["total_lines"]; dur_disp = f"{unit_cnt} câu"
                    else: unit_cnt = info["total_words"]; dur_disp = f"{unit_cnt} từ"

                    tot_p = sum(r["Pay_Num"] for r in info["video_rows"])
                    grand_actor_pay += tot_p

                    actor_payroll_rows.append({
                        "Stt": idx, "Diễn viên Lồng tiếng": act_name, "Số Video đã lồng": info["videos_count"],
                        "Tổng phút video": dur_disp, "Thành tiền Lương": f"{tot_p:,.0f} VNĐ",
                        "Danh sách Video tham gia": ", ".join([r["Tiêu đề video"] for r in info["video_rows"]])
                    })

                actor_payroll_rows.append({
                    "Stt": "TỔNG", "Diễn viên Lồng tiếng": f"TỔNG CỘNG ({len(actor_weekly_map)} Diễn viên)",
                    "Số Video đã lồng": "-", "Tổng phút video": f"{grand_unique_mins} phút" if current_mode == "minute" else "-",
                    "Thành tiền Lương": f"{grand_actor_pay:,.0f} VNĐ", "Danh sách Video tham gia": "-"
                })

                df_act_payroll = pd.DataFrame(actor_payroll_rows)
                st.metric(f"💰 TỔNG LƯƠNG CẦN CHI CHO DIỄN VIÊN ({selected_week_filter}):", f"{grand_actor_pay:,.0f} VNĐ")
                st.dataframe(
                    df_act_payroll[["Stt", "Diễn viên Lồng tiếng", "Số Video đã lồng", "Tổng phút video", "Thành tiền Lương", "Danh sách Video tham gia"]],
                    hide_index=True, use_container_width=True
                )

                excel_actor_buf = io.BytesIO()
                with pd.ExcelWriter(excel_actor_buf, engine='openpyxl') as writer: df_act_payroll.to_excel(writer, index=False, sheet_name="Luong Dien Vien")
                excel_actor_buf.seek(0)
                st.download_button(
                    label=f"📊 TẢI BÁO CÁO LƯƠNG TOÀN BỘ DIỄN VIÊN ({selected_week_filter}) EXCEL (.XLSX)",
                    data=excel_actor_buf, file_name=f"Bao_Cao_Luong_DienVien_{selected_week_filter}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    type="primary", use_container_width=True
                )
        else: st.info("Chưa có dữ liệu lồng tiếng để lập báo cáo lương.")

# ==========================================
# TAB 4: QUẢN LÝ DATABASE PHÂN VAI & MÀU SẮC LỒNG TIẾNG
# ==========================================
with tab_cast_db:
    subtab_cast_map, subtab_color_map = st.tabs([
        "🎭 Bảng Phân Vai Diễn Viên (Global Database)",
        "🎨 Bảng Màu & Highlight Nhân Vật Cố Định (VAI - MAU.xlsx Database)"
    ])

    # SUBTAB 1: PHÂN VAI DIỄN VIÊN
    with subtab_cast_map:
        with st.container(border=True):
            st.subheader("🎭 BẢNG PHÂN VAI LỒNG TIẾNG")
            st.markdown("Nơi thiết lập mặc định nhân vật Tiếng Anh nào sẽ do diễn viên lồng tiếng Việt nào đảm nhận cho Mai Han Team.")

            st.markdown("#### ➕ Thêm / Cập nhật Phân vai mới")
            c_c1, c_c2, c_c3 = st.columns([2, 2, 1.2])
            with c_c1: add_role_eng = st.text_input("Tên Nhân vật (Tiếng Anh):", placeholder="VD: Bri, Chase...", key=f"add_role_eng_{st.session_state['cast_input_key']}")
            with c_c2: add_actor_vn = st.text_input("Diễn viên Lồng tiếng (Tiếng Việt):", placeholder="VD: TRÚC, THIỆN...", key=f"add_actor_vn_{st.session_state['cast_input_key']}")
            with c_c3:
                st.markdown("<div style='margin-top: 28px;'></div>", unsafe_allow_html=True)
                if st.button("➕ Thêm Phân Vai", use_container_width=True, type="primary", key="btn_add_cast"):
                    if add_role_eng and add_actor_vn:
                        k = add_role_eng.upper().strip(); v = add_actor_vn.strip().upper()
                        st.session_state['custom_cast_mapping'][k] = v
                        save_json_db(CAST_DB_FILE, st.session_state['custom_cast_mapping'])
                        st.session_state['cast_input_key'] += 1
                        st.success(f"✅ Đã gán thành công: `{add_role_eng}` ➔ `{add_actor_vn}`"); time.sleep(1); st.rerun()
                    else: st.warning("Vui lòng nhập đầy đủ tên nhân vật và diễn viên!")

        st.markdown("---")
        with st.container(border=True):
            st.markdown("#### 📑 Danh sách Toàn bộ Bảng Phân Vai Đã Lưu")
            search_cast_query = st.text_input("🔍 Tìm kiếm Nhân vật hoặc Diễn viên lồng tiếng:", placeholder="Gõ tên nhân vật hoặc diễn viên...").strip().upper()
            all_cast_dict = st.session_state['custom_cast_mapping']

            if search_cast_query: filtered_cast = {k: v for k, v in all_cast_dict.items() if search_cast_query in k or search_cast_query in v.upper()}
            else: filtered_cast = all_cast_dict

            if filtered_cast:
                cast_db_table = []
                for eng_role, vn_actor in sorted(filtered_cast.items()):
                    cast_db_table.append({"Nhân vật (Tiếng Anh)": eng_role, "Diễn viên Lồng tiếng (Tiếng Việt)": vn_actor, "Xóa khỏi Database": False})

                df_cast_db = pd.DataFrame(cast_db_table)
                st.caption(f"Đang hiển thị **{len(df_cast_db)}** vai lồng tiếng trong kho:")

                edited_cast_db_df = st.data_editor(
                    df_cast_db,
                    column_config={
                        "Nhân vật (Tiếng Anh)": st.column_config.TextColumn("Tên Nhân vật gốc (In hoa)", disabled=True),
                        "Diễn viên Lồng tiếng (Tiếng Việt)": st.column_config.TextColumn("Diễn viên lồng tiếng (Sửa trực tiếp)"),
                        "Xóa khỏi Database": st.column_config.CheckboxColumn("Xóa?")
                    },
                    disabled=["Nhân vật (Tiếng Anh)"], hide_index=True, use_container_width=True, key="global_cast_db_editor"
                )

                if st.button("💾 LƯU TOÀN BỘ CẬP NHẬT PHÂN VAI", type="primary", use_container_width=True, key="btn_save_global_cast"):
                    new_cast_db = {}; deleted_cast_count = 0
                    if search_cast_query:
                        for k, v in all_cast_dict.items():
                            if k not in filtered_cast: new_cast_db[k] = v

                    for _, row in edited_cast_db_df.iterrows():
                        eng_k = str(row["Nhân vật (Tiếng Anh)"]).upper().strip()
                        act_v = str(row["Diễn viên Lồng tiếng (Tiếng Việt)"]).strip().upper()
                        is_del = row["Xóa khỏi Database"]
                        if is_del: deleted_cast_count += 1
                        else:
                            if act_v: new_cast_db[eng_k] = act_v

                    st.session_state['custom_cast_mapping'] = new_cast_db
                    save_json_db(CAST_DB_FILE, new_cast_db)
                    st.success(f"✅ Đã lưu cập nhật thành công! (Đã xóa {deleted_cast_count} vai)"); time.sleep(1); st.rerun()
            else: st.info("Không tìm thấy vai lồng tiếng nào khớp với từ khóa tìm kiếm.")

    # SUBTAB 2: BẢNG MÀU CHỮ LẪN HIGHLIGHT CỐ ĐỊNH (SỬA LỖI STREAMLIT COLORCOLUMN BẰNG BẢNG THẺ MÀU TRỰC QUAN AN TOÀN)
    with subtab_color_map:
        fixed_color_dict = st.session_state['fixed_speaker_colors']
        
        with st.container(border=True):
            st.subheader("🎨 BẢNG MÀU CHỮ & HIGHLIGHT CỐ ĐỊNH (FIXED COLOR & HIGHLIGHT)")
            st.markdown("Quản lý danh sách nhân vật cốt lõi có MÀU SẮC CHỮ và MÀU HIGHLIGHT NỀN cố định cho kịch bản Tab 1 & Tab 2.")

            st.markdown("#### ➕ Bổ sung / Sửa Màu & Highlight Cố định")
            col_col1, col_col2, col_col3, col_col4 = st.columns([1.8, 1.5, 1.5, 1.2])
            with col_col1:
                new_color_spk = st.text_input("Tên Nhân vật Cố định:", placeholder="VD: ALL, CORY, NICK...", key=f"add_color_spk_{st.session_state['color_input_key']}").strip().upper()
            with col_col2:
                enable_tc = st.checkbox("Tô Màu Chữ", value=True, key=f"chk_tc_{st.session_state['color_input_key']}")
                new_text_hex = st.color_picker("Chọn Màu Chữ:", "#FF0000", key=f"add_tc_picker_{st.session_state['color_input_key']}") if enable_tc else None
            with col_col3:
                enable_hc = st.checkbox("Tô Highlight Nền", value=False, key=f"chk_hc_{st.session_state['color_input_key']}")
                new_hl_hex = st.color_picker("Chọn Màu Highlight:", "#FFFF00", key=f"add_hc_picker_{st.session_state['color_input_key']}") if enable_hc else None
            with col_col4:
                st.markdown("<div style='margin-top: 28px;'></div>", unsafe_allow_html=True)
                if st.button("➕ Thêm Quy Tắc Màu", use_container_width=True, type="primary", key="btn_add_fixed_color"):
                    if new_color_spk:
                        tc_tuple = hex_to_rgb(new_text_hex) if enable_tc else None
                        hc_tuple = hex_to_rgb(new_hl_hex) if enable_hc else None
                            
                        st.session_state['fixed_speaker_colors'][new_color_spk] = {
                            "text_color": tc_tuple,
                            "highlight_color": hc_tuple
                        }
                        save_json_db(SPEAKER_COLOR_DB_FILE, st.session_state['fixed_speaker_colors'])
                        st.session_state['color_input_key'] += 1
                        st.success(f"✅ Đã lưu cấu hình Màu Chữ & Highlight cho `{new_color_spk}`!"); time.sleep(1); st.rerun()

            # 👁️ TRÌNH XEM TRƯỚC MÀU THỰC TẾ TRỰC QUAN (LIVE PREVIEW CARD)
            st.markdown("##### 👁️ Xem trước màu hiển thị thực tế trên kịch bản lồng tiếng:")
            p_tc = new_text_hex if (enable_tc and new_text_hex) else "#0F172A"
            p_bg = new_hl_hex if (enable_hc and new_hl_hex) else "transparent"
            spk_disp_name = new_color_spk if new_color_spk else "TÊN_NHÂN_VẬT"

            st.markdown(f"""
            <div style="padding: 12px 18px; border: 1px dashed #0284C7; border-radius: 8px; background-color: #F8FAFC; margin-bottom: 15px;">
                <span style="color: {p_tc}; background-color: {p_bg}; font-family: 'Times New Roman', serif; font-size: 16px; font-weight: bold; padding: 2px 8px; border-radius: 4px;">
                    {spk_disp_name}:
                </span>
                <span style="font-family: 'Times New Roman', serif; font-size: 16px; color: #0F172A;">
                    Mẫu câu thoại hiển thị thực tế khi xuất ra kịch bản Word lồng tiếng.
                </span>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("---")
        with st.container(border=True):
            st.markdown("#### 🎨 BẢNG THẺ MÀU TRỰC QUAN TOÀN BỘ NHÂN VẬT CỐ ĐỊNH")
            st.caption("Danh sách 25+ nhân vật đã lưu kho được hiển thị trực tiếp bằng màu chữ & màu highlight nền thật:")
            
            # GIẢI PHÁP HIỂN THỊ THẺ MÀU TRỰC QUAN VƯỢT TRỘI KHÔNG LỖI
            badge_html_items = []
            for spk_k, cfg in sorted(fixed_color_dict.items()):
                if isinstance(cfg, dict):
                    tc = cfg.get("text_color")
                    hc = cfg.get("highlight_color")
                else:
                    tc = tuple(cfg) if isinstance(cfg, (list, tuple)) else (255, 0, 0)
                    hc = None
                
                tc_css = f"rgb({tc[0]},{tc[1]},{tc[2]})" if tc else "#0F172A"
                hc_css = f"rgb({hc[0]},{hc[1]},{hc[2]})" if hc else "#FFFFFF"
                
                badge_html_items.append(
                    f'<div style="display: inline-block; margin: 5px; padding: 6px 14px; border-radius: 6px; background-color: {hc_css}; border: 1px solid #CBD5E1; box-shadow: 0 1px 2px rgba(0,0,0,0.05);">'
                    f'<span style="color: {tc_css}; font-family: \'Times New Roman\', serif; font-weight: bold; font-size: 15px;">{spk_k}:</span> '
                    f'<span style="font-size: 12px; color: #64748B;">({tc_css})</span>'
                    f'</div>'
                )

            st.markdown(f'<div style="background: #F8FAFC; padding: 15px; border-radius: 10px; border: 1px solid #E2E8F0; margin-bottom: 20px;">{"".join(badge_html_items)}</div>', unsafe_allow_html=True)

            st.markdown("##### 📑 Bảng Chi Tiết Mã Màu & Chỉnh Sửa / Xóa:")
            color_db_table = []
            for spk_k, cfg in sorted(fixed_color_dict.items()):
                if isinstance(cfg, dict):
                    tc = cfg.get("text_color")
                    hc = cfg.get("highlight_color")
                else:
                    tc = tuple(cfg) if isinstance(cfg, (list, tuple)) else (255, 0, 0)
                    hc = None
                    
                tc_hex = f"#{tc[0]:02X}{tc[1]:02X}{tc[2]:02X}" if tc else ""
                hc_hex = f"#{hc[0]:02X}{hc[1]:02X}{hc[2]:02X}" if hc else ""
                
                color_db_table.append({
                    "Nhân vật Cố định": spk_k,
                    "Màu Chữ (Mã Hex)": tc_hex,
                    "Highlight Nền (Mã Hex)": hc_hex,
                    "Xóa": False
                })

            df_color_db = pd.DataFrame(color_db_table)
            edited_color_db_df = st.data_editor(
                df_color_db,
                column_config={
                    "Nhân vật Cố định": st.column_config.TextColumn("Tên Nhân vật", disabled=True),
                    "Màu Chữ (Mã Hex)": st.column_config.TextColumn("Màu chữ (Gõ mã Hex)"),
                    "Highlight Nền (Mã Hex)": st.column_config.TextColumn("Highlight Nền (Gõ mã Hex)"),
                    "Xóa": st.column_config.CheckboxColumn("Xóa?")
                },
                hide_index=True, use_container_width=True, key="fixed_color_db_editor_table"
            )

            if st.button("💾 LƯU BẢNG MÀU VÀO DATABASE", type="secondary", use_container_width=True, key="btn_save_fixed_color_db"):
                new_fixed_map = {}
                for _, row in edited_color_db_df.iterrows():
                    if not row["Xóa"]:
                        spk_k = str(row["Nhân vật Cố định"]).upper().strip()
                        tc_hex = row["Màu Chữ (Mã Hex)"]
                        hc_hex = row["Highlight Nền (Mã Hex)"]
                        new_fixed_map[spk_k] = {
                            "text_color": hex_to_rgb(tc_hex),
                            "highlight_color": hex_to_rgb(hc_hex)
                        }
                st.session_state['fixed_speaker_colors'] = new_fixed_map
                save_json_db(SPEAKER_COLOR_DB_FILE, new_fixed_map)
                st.success("✅ Đã lưu cập nhật Bảng Màu Cố Định vào Database!"); time.sleep(1); st.rerun()

# ==========================================
# TAB 5: KHO DATABASE PHIÊN ÂM GIỌNG NAM (TỔNG HỢP)
# ==========================================
with tab_phonetic_db:
    with st.container(border=True):
        st.subheader("📚 Từ Điển Phiên Âm Giọng Nam (Global Database)")
        st.markdown("Nơi quản lý toàn bộ kho từ vựng Tiếng Anh và các bản phiên âm giọng Nam được lưu trữ lâu dài trên hệ thống.")
        
        st.markdown("#### 🔊 Nghe phát âm thử bất kỳ cụm từ/từ Tiếng Anh nào")
        col_test1, col_test2, col_test3 = st.columns([2.5, 1.5, 1.5])
        with col_test1: free_test_word = st.text_input("Nhập từ/cụm Tiếng Anh cần nghe thử:", placeholder="VD: Starbucks, Hamburger, McDonald's...", key="free_audio_text")
        with col_test2: free_accent = st.radio("Giọng phát âm:", options=["Giọng Mỹ (US)", "Giọng Anh (UK)"], horizontal=True, key="free_audio_accent")
        with col_test3:
            st.markdown("<div style='margin-top: 28px;'></div>", unsafe_allow_html=True)
            free_listen_btn = st.button("🔊 Phát âm thanh", type="secondary", use_container_width=True, key="btn_free_listen")

        if free_listen_btn and free_test_word:
            tld = 'com' if "Mỹ" in free_accent else 'co.uk'
            test_fp = generate_english_audio(free_test_word, accent=tld)
            if test_fp: st.audio(test_fp, format="audio/mp3", autoplay=True)

        st.markdown("---")
        st.markdown("#### ➕ Bổ sung từ phiên âm mới vào Kho")
        c1, c2, c3 = st.columns([2, 2, 1.2])
        with c1: tab_add_eng = st.text_input("Từ Tiếng Anh gốc:", placeholder="VD: Burger", key=f"tab_add_eng_{st.session_state['pho_input_key']}")
        with c2: tab_add_pho = st.text_input("Phiên âm giọng Nam:", placeholder="VD: Bơ-gơ", key=f"tab_add_pho_{st.session_state['pho_input_key']}")
        with c3:
            st.markdown("<div style='margin-top: 28px;'></div>", unsafe_allow_html=True)
            if st.button("➕ Thêm vào Database", use_container_width=True, type="primary"):
                if tab_add_eng and tab_add_pho:
                    k = tab_add_eng.upper().strip(); v = tab_add_pho.strip()
                    st.session_state['custom_phonetics'][k] = v
                    save_json_db(PHONETIC_DB_FILE, st.session_state['custom_phonetics'])
                    st.session_state['pho_input_key'] += 1
                    st.success(f"✅ Đã thêm thành công: `{tab_add_eng}` ➔ `{tab_add_pho}`"); time.sleep(1); st.rerun()
                else: st.warning("Vui lòng điền đủ 2 ô!")

    st.markdown("---")
    with st.container(border=True):
        st.markdown("#### 📑 Danh sách toàn bộ Từ phiên âm đã lưu")
        search_query = st.text_input("🔍 Tìm kiếm từ Tiếng Anh hoặc Từ phiên âm:", placeholder="Gõ từ cần tìm ở đây...").strip().upper()
        all_phonetics_dict = st.session_state['custom_phonetics']
        
        if search_query: filtered_dict = {k: v for k, v in all_phonetics_dict.items() if search_query in k or search_query in v.upper()}
        else: filtered_dict = all_phonetics_dict

        if filtered_dict:
            db_table_data = []
            for eng_key, pho_val in sorted(filtered_dict.items()):
                db_table_data.append({"Từ Tiếng Anh": eng_key, "Phiên âm giọng Nam": pho_val, "Xóa khỏi Database": False})

            df_db = pd.DataFrame(db_table_data)
            st.caption(f"Đang hiển thị **{len(df_db)}** từ phiên âm trong hệ thống:")

            edited_db_df = st.data_editor(
                df_db,
                column_config={
                    "Từ Tiếng Anh": st.column_config.TextColumn("Từ Tiếng Anh gốc (In hoa)", disabled=True),
                    "Phiên âm giọng Nam": st.column_config.TextColumn("Phiên âm giọng Nam (Sửa trực tiếp tại đây)"),
                    "Xóa khỏi Database": st.column_config.CheckboxColumn("Xóa?")
                },
                disabled=["Từ Tiếng Anh"], hide_index=True, use_container_width=True, key="global_phonetic_db_editor"
            )

            if st.button("💾 LƯU TOÀN BỘ CẬP NHẬT TRONG BẢNG", type="primary", use_container_width=True):
                new_db = {}; deleted_count = 0
                if search_query:
                    for k, v in all_phonetics_dict.items():
                        if k not in filtered_dict: new_db[k] = v

                for _, row in edited_db_df.iterrows():
                    eng_k = str(row["Từ Tiếng Anh"]).upper().strip(); pho_v = str(row["Phiên âm giọng Nam"]).strip()
                    is_delete = row["Xóa khỏi Database"]
                    if is_delete: deleted_count += 1
                    else:
                        if pho_v: new_db[eng_k] = pho_v

                st.session_state['custom_phonetics'] = new_db
                save_json_db(PHONETIC_DB_FILE, new_db)
                st.success(f"✅ Đã lưu cập nhật thành công! (Đã xóa {deleted_count} từ)"); time.sleep(1); st.rerun()
        else: st.info("Không tìm thấy từ phiên âm nào khớp với từ khóa tìm kiếm.")

# ==========================================
# TAB 6: ĐỐI CHIẾU 2 FILE TIẾNG ANH & QC DỊCH
# ==========================================
with tab_dual_align:
    st.subheader("🔀 ĐỐI CHIẾU 2 FILE TIẾNG ANH & SOÁT SỬA BẢN DỊCH VIỆT")
    st.markdown("So sánh file Tiếng Anh do Mai Han Team nghe (giữ 100% Timecode chuẩn Subtitle Edit) với file Tiếng Anh Gốc do Khách gửi trễ. Tự động quét vùng mở rộng 3 câu để không bị báo nhầm khi chỉnh timecode và hỗ trợ gán Tên người nói mặc định!")

    col_spk_fb1, col_spk_fb2 = st.columns([1.8, 1.2])
    with col_spk_fb1:
        default_spk_input = st.text_input(
            "🎭 Tên người nói mặc định (dùng cho Báo cáo QC):", 
            placeholder="VD: Nick, Narrator, MC...", 
            value="", 
            help="Điền tên ở đây để Báo cáo QC Excel hiển thị 'Nick' rõ ràng thay vì 'Unknown:'. Tên này sẽ KHÔNG in vào file kịch bản xuất ra (Word/SRT) nếu bật tùy chọn bên phải."
        )

    with col_spk_fb2:
        st.markdown("<div style='margin-top: 28px;'></div>", unsafe_allow_html=True)
        hide_default_spk_export = st.checkbox(
            "🚫 Không in Tên mặc định vào Kịch bản xuất ra (Word/SRT)", 
            value=True,
            help="Tên người nói mặc định (VD: Nick) chỉ dùng để hiện trên Báo cáo QC Excel. File Word và SRT xuất ra sẽ giữ nguyên văn bản gọn gàng không có chữ 'Nick:' ở từng câu."
        )

    col_dual1, col_dual2, col_dual3 = st.columns(3)
    with col_dual1:
        with st.container(border=True):
            st.markdown("##### 📄 1. File Tiếng Anh - Mai Han Team (.srt/docx)")
            st.caption("File Tiếng Anh ngắt câu & kéo mốc thời gian chuẩn do Mai Han Team thực hiện:")
            uploaded_mh_eng = st.file_uploader("Tải file tiếng Anh - Mai Han Team (.srt/docx):", type=['srt', 'docx'], key="uploader_dual_mh_eng")

    with col_dual2:
        with st.container(border=True):
            st.markdown("##### 📄 2. File Tiếng Anh của Khách (.srt/docx)")
            st.caption("File Tiếng Anh gốc do khách hàng gửi sang sau đó:")
            uploaded_off_eng = st.file_uploader("Tải file tiếng Anh của Khách (.srt/docx):", type=['srt', 'docx'], key="uploader_dual_off_eng")

    with col_dual3:
        with st.container(border=True):
            st.markdown("##### 📄 3. File Tiếng Việt Hiện Tại (Vietnamese Script)")
            st.caption("File Dịch Tiếng Việt Mai Han đã dịch từ bản AI nghe trước đây:")
            uploaded_vn_script = st.file_uploader("Tải file Tiếng Việt (.srt/docx):", type=['srt', 'docx'], key="uploader_dual_vn_script")

    if uploaded_mh_eng is not None and uploaded_off_eng is not None:
        c_spks = st.session_state.get('custom_speakers', set())
        c_non_spks = st.session_state.get('custom_non_speakers', set())
        fallback_spk_name = default_spk_input.strip() if default_spk_input.strip() else "Unknown"

        df_mh_eng = parse_any_script_file_to_df(uploaded_mh_eng.getvalue(), uploaded_mh_eng.name, c_spks, c_non_spks, fallback_spk_name)
        df_off_eng = parse_any_script_file_to_df(uploaded_off_eng.getvalue(), uploaded_off_eng.name, c_spks, c_non_spks, fallback_spk_name)
        
        df_vn_script = None
        if uploaded_vn_script is not None:
            df_vn_script = parse_any_script_file_to_df(uploaded_vn_script.getvalue(), uploaded_vn_script.name, c_spks, c_non_spks, fallback_spk_name)

        if not df_mh_eng.empty and not df_off_eng.empty:
            df_aligned = align_and_compare_english_scripts(df_mh_eng, df_off_eng, df_vn_script, fallback_spk_name)

            # Analytics Summary
            tot_rows = len(df_aligned)
            missing_cnt = sum(1 for st_v in df_aligned['Trạng thái QC'] if '🔴' in str(st_v))
            word_diff_cnt = sum(1 for st_v in df_aligned['Trạng thái QC'] if '🟡' in str(st_v))
            mismatch_cnt = sum(1 for st_v in df_aligned['Trạng thái QC'] if '🔵' in str(st_v))

            st.markdown("---")
            st.markdown("#### 📊 Báo Cáo QC So Sánh Nội Dung Tiếng Anh (English Content Audit)")

            col_m1, col_m2, col_m3, col_m4 = st.columns(4)
            with col_m1:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-label">💬 Tổng số câu thoại</div>
                    <div class="metric-value">{tot_rows}</div>
                </div>
                """, unsafe_allow_html=True)
            with col_m2:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-label">🔴 Bỏ sót câu thoại</div>
                    <div class="metric-value" style="color:#DC2626;">{missing_cnt}</div>
                </div>
                """, unsafe_allow_html=True)
            with col_m3:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-label">🟡 Khác từ vựng (AI nghe nhầm)</div>
                    <div class="metric-value" style="color:#D97706;">{word_diff_cnt}</div>
                </div>
                """, unsafe_allow_html=True)
            with col_m4:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-label">🔵 Lệch người nói</div>
                    <div class="metric-value" style="color:#2563EB;">{mismatch_cnt}</div>
                </div>
                """, unsafe_allow_html=True)

            st.markdown("---")
            st.markdown("#### 👁️ Workspace Bảng Đối Chiếu Tiếng Anh & Chỉnh Sửa Dịch Tiếng Việt")
            st.caption("Nhìn ô màu VÀNG/ĐỎ để biết AI nghe sai từ nào, sau đó gõ sửa trực tiếp câu Tiếng Việt ở ô kế bên:")

            qc_filter = st.selectbox("🔍 Lọc danh sách câu theo Trạng thái QC:", options=["TẤT CẢ CÁC CÂU", "🟡 Chỉ xem câu KHÁC TỪ VỰNG (Cần sửa dịch)", "🔴 Chỉ xem câu BỎ SÓT THOẠI", "🔵 Chỉ xem câu LỆCH VAI"])

            if "🔴" in qc_filter: df_display = df_aligned[df_aligned['Trạng thái QC'].str.contains('🔴', na=False)]
            elif "🟡" in qc_filter: df_display = df_aligned[df_aligned['Trạng thái QC'].str.contains('🟡', na=False)]
            elif "🔵" in qc_filter: df_display = df_aligned[df_aligned['Trạng thái QC'].str.contains('🔵', na=False)]
            else: df_display = df_aligned

            edited_aligned_df = st.data_editor(
                df_display[['Stt', 'Timecode', 'Tiếng Anh Mai Han (AI/Heard)', 'Tiếng Anh Khách (Official)', 'Dịch Tiếng Việt (Cần Sửa)', 'Trạng thái QC', 'Ghi chú QC']],
                column_config={
                    "Stt": st.column_config.NumberColumn("Stt", disabled=True, width="small"),
                    "Timecode": st.column_config.TextColumn("Mốc Timecode Mai Han", disabled=True),
                    "Tiếng Anh Mai Han (AI/Heard)": st.column_config.TextColumn("Bản Anh Mai Han nghe", disabled=True),
                    "Tiếng Anh Khách (Official)": st.column_config.TextColumn("Bản Anh Khách gửi chuẩn", disabled=True),
                    "Dịch Tiếng Việt (Cần Sửa)": st.column_config.TextColumn("Kịch bản Tiếng Việt (Sửa trực tiếp tại đây)"),
                    "Trạng thái QC": st.column_config.TextColumn("Trạng thái", disabled=True),
                    "Ghi chú QC": st.column_config.TextColumn("Ghi chú từ vựng khác biệt", disabled=True)
                },
                hide_index=True,
                use_container_width=True,
                key="dual_english_editor_table"
            )

            st.markdown("---")
            st.markdown("#### ⬇️ Xuất File Kịch Bản Tiếng Việt Hoàn Chỉnh (Giữ Timecode Mai Han)")

            col_ex1, col_ex2, col_ex3 = st.columns(3)
            base_out_name = os.path.splitext(uploaded_mh_eng.name)[0]

            with col_ex1:
                qc_excel_buf = generate_qc_dual_excel(df_aligned)
                st.download_button(
                    label="📊 TẢI BÁO CÁO ĐỐI CHIẾU EXCEL (.XLSX)",
                    data=qc_excel_buf,
                    file_name=f"{base_out_name}_DoiChieu_English_QC.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    type="primary", use_container_width=True
                )

            with col_ex2:
                aligned_docx_buf = generate_aligned_docx_file(
                    df_aligned, base_out_name, enable_colors, enable_phonetic, enable_cast,
                    hide_default_spk=hide_default_spk_export, fallback_spk_name=fallback_spk_name,
                    font_size_pt=12
                )
                st.download_button(
                    label="📄 TẢI WORD VIỆT HOÀN CHỈNH (.DOCX)",
                    data=aligned_docx_buf,
                    file_name=f"{base_out_name}_VI_Final.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary", use_container_width=True
                )

            with col_ex3:
                # Generate SRT (Ẩn tên mặc định nếu bật tùy chọn hide_default_spk_export)
                srt_out_lines = []
                for idx_s, r_s in df_aligned.iterrows():
                    spk_val = r_s['Speaker_MH']
                    is_explicit = r_s.get('Is_Explicit_MH', True)
                    
                    should_show_spk = True
                    if hide_default_spk_export and (not is_explicit or spk_val.upper() == fallback_spk_name.upper() or spk_val.upper() == "UNKNOWN"):
                        should_show_spk = False

                    if should_show_spk and spk_val:
                        spk_text = f"{spk_val}: {r_s['Dialogue_VN']}"
                    else:
                        spk_text = r_s['Dialogue_VN']

                    srt_out_lines.append(f"{idx_s+1}\n{r_s['Timecode']}\n{spk_text}\n")
                srt_out_bytes = "\n".join(srt_out_lines).encode('utf-8-sig')

                st.download_button(
                    label="📝 TẢI SUBTITLE SRT VIỆT (.SRT)",
                    data=srt_out_bytes,
                    file_name=f"{base_out_name}_VI_Final.srt",
                    mime="text/plain",
                    type="primary", use_container_width=True
                )

# ==========================================
# TAB 7: SOÁT BẤT NHẤT THUẬT NGỮ & XƯNG HÔ
# ==========================================
with tab_consistency:
    st.subheader("🔎 SOÁT BẤT NHẤT THUẬT NGỮ & QUAN HỆ XƯNG HÔ NHÂN VẬT")
    st.markdown("Vùng làm việc phát hiện các câu thoại bị sượng xưng hô (lệch đại từ tui/ông, mình/bạn) hoặc bất nhất bản dịch thuật ngữ/tên món ăn giữa các đoạn trong kịch bản.")

    subtab_pronoun, subtab_glossary = st.tabs([
        "👥 Quản Lý & Soát Lỗi Xưng Hô Nhân Vật", 
        "📚 Soát Bất Nhất Thuật Ngữ & Món Ăn/Tên Riêng"
    ])

    # 1. SUBTAB PHẦN XƯNG HÔ NHÂN VẬT
    with subtab_pronoun:
        st.markdown("#### 1. Bảng Thiết Lập Quan Hệ Xưng Hô (Lưu Database)")
        st.caption("Thiết lập quy tắc xưng hô mặc định giữa các cặp nhân vật để hệ thống tự động kiểm tra kịch bản:")

        col_p1, col_p2, col_p3, col_p4, col_p5 = st.columns([2, 2, 1.5, 1.5, 1.2])
        with col_p1:
            rel_spk_a = st.text_input("Người Nói (Speaker A):", placeholder="VD: TYLER", key=f"p_spk_a_{st.session_state['pronoun_input_key']}").strip().upper()
        with col_p2:
            rel_spk_b = st.text_input("Người Nghe (Speaker B):", placeholder="VD: BILL", key=f"p_spk_b_{st.session_state['pronoun_input_key']}").strip().upper()
        with col_p3:
            rel_self = st.text_input("Xưng (Self):", placeholder="VD: tui, tôi, mình...", key=f"p_self_{st.session_state['pronoun_input_key']}").strip().lower()
        with col_p4:
            rel_target = st.text_input("Gọi (Target):", placeholder="VD: ông, bạn, anh...", key=f"p_target_{st.session_state['pronoun_input_key']}").strip().lower()
        with col_p5:
            st.markdown("<div style='margin-top: 28px;'></div>", unsafe_allow_html=True)
            if st.button("➕ Thêm Xưng Hô", type="primary", use_container_width=True, key="btn_add_pronoun"):
                if rel_spk_a and rel_spk_b and rel_self and rel_target:
                    key_pair = f"{rel_spk_a}|{rel_spk_b}"
                    st.session_state['custom_pronoun_rel'][key_pair] = {"self": rel_self, "target": rel_target}
                    save_json_db(PRONOUN_REL_DB_FILE, st.session_state['custom_pronoun_rel'])
                    st.session_state['pronoun_input_key'] += 1
                    st.success(f"✅ Đã lưu: {rel_spk_a} ➔ {rel_spk_b} ({rel_self} - {rel_target})"); time.sleep(1); st.rerun()
                else: st.warning("Vui lòng điền đủ 4 ô!")

        # Bảng Data Editor cho Xưng hô
        rel_db_dict = st.session_state['custom_pronoun_rel']
        if rel_db_dict:
            rel_data_table = []
            for pair_key, val in sorted(rel_db_dict.items()):
                spk_a, spk_b = pair_key.split("|") if "|" in pair_key else (pair_key, "ALL")
                rel_data_table.append({
                    "Người Nói": spk_a, "Người Nghe": spk_b,
                    "Xưng (Self)": val.get("self", "tui"), "Gọi (Target)": val.get("target", "ông"),
                    "Xóa": False
                })

            df_rel = pd.DataFrame(rel_data_table)
            edited_rel_df = st.data_editor(
                df_rel,
                column_config={
                    "Người Nói": st.column_config.TextColumn("Người Nói", disabled=True),
                    "Người Nghe": st.column_config.TextColumn("Người Nghe", disabled=True),
                    "Xưng (Self)": st.column_config.TextColumn("Đại từ Xưng (Self)"),
                    "Gọi (Target)": st.column_config.TextColumn("Đại từ Gọi (Target)"),
                    "Xóa": st.column_config.CheckboxColumn("Xóa?")
                },
                hide_index=True, use_container_width=True, key="pronoun_rel_editor_table"
            )

            if st.button("💾 LƯU BẢNG QUY TẮC XƯNG HÔ VÀO DATABASE", type="secondary", use_container_width=True):
                new_rel_db = {}
                for _, row in edited_rel_df.iterrows():
                    if not row["Xóa"]:
                        pk = f"{str(row['Người Nói']).upper()}|{str(row['Người Nghe']).upper()}"
                        new_rel_db[pk] = {"self": str(row["Xưng (Self)"]).lower(), "target": str(row["Gọi (Target)"]).lower()}
                st.session_state['custom_pronoun_rel'] = new_rel_db
                save_json_db(PRONOUN_REL_DB_FILE, new_rel_db)
                st.success("✅ Đã lưu cập nhật Bảng Xưng Hô vào Database!"); time.sleep(1); st.rerun()

        st.markdown("---")
        st.markdown("#### 2. Công Cụ Soát Lỗi Xưng Hô Tự Động Trong Kịch Bản")
        uploaded_pronoun_script = st.file_uploader("Tải file Kịch bản Tiếng Việt (.srt hoặc .docx) để kiểm tra xưng hô:", type=['srt', 'docx'], key="uploader_pronoun_qc")

        if uploaded_pronoun_script is not None:
            c_spks_p = st.session_state.get('custom_speakers', set())
            c_non_spks_p = st.session_state.get('custom_non_speakers', set())
            df_p_script = parse_any_script_file_to_df(uploaded_pronoun_script.getvalue(), uploaded_pronoun_script.name, c_spks_p, c_non_spks_p)

            if not df_p_script.empty:
                pronoun_audit_rows = []
                speaker_pronoun_stats = {}

                for idx, r in df_p_script.iterrows():
                    spk = str(r['Speaker']).strip()
                    diag = str(r['Dialogue']).strip()
                    words = [w.lower() for w in re.findall(r'\b\w+\b', diag)]

                    found_self = [w for w in words if w in VN_SELF_PRONOUNS]
                    found_target = [w for w in words if w in VN_TARGET_PRONOUNS]

                    if spk not in speaker_pronoun_stats:
                        speaker_pronoun_stats[spk] = {"self": Counter(), "target": Counter()}

                    for s_w in found_self: speaker_pronoun_stats[spk]["self"][s_w] += 1
                    for t_w in found_target: speaker_pronoun_stats[spk]["target"][t_w] += 1

                for idx, r in df_p_script.iterrows():
                    spk = str(r['Speaker']).strip()
                    diag = str(r['Dialogue']).strip()
                    words = [w.lower() for w in re.findall(r'\b\w+\b', diag)]

                    found_self = [w for w in words if w in VN_SELF_PRONOUNS]
                    found_target = [w for w in words if w in VN_TARGET_PRONOUNS]

                    top_self = speaker_pronoun_stats[spk]["self"].most_common(1)[0][0] if speaker_pronoun_stats[spk]["self"] else ""
                    top_target = speaker_pronoun_stats[spk]["target"].most_common(1)[0][0] if speaker_pronoun_stats[spk]["target"] else ""

                    is_unusual = False
                    warn_msg = []

                    if found_self and top_self and any(w != top_self for w in found_self):
                        is_unusual = True
                        warn_msg.append(f"Xưng '{', '.join(found_self)}' (Đa số nhân vật xưng '{top_self}')")

                    if found_target and top_target and any(w != top_target for w in found_target):
                        is_unusual = True
                        warn_msg.append(f"Gọi '{', '.join(found_target)}' (Đa số nhân vật gọi '{top_target}')")

                    status_str = "🟡 Nghi vấn lệch xưng hô" if is_unusual else "🟢 Ok"
                    pronoun_audit_rows.append({
                        "Stt": idx + 1, "Timecode": f"{r['Start']} --> {r['End']}",
                        "Nhân vật": spk, "Câu thoại Tiếng Việt": diag,
                        "Trạng thái": status_str, "Chi tiết QC": "; ".join(warn_msg) if warn_msg else "Xưng hô khớp với tần suất chính"
                    })

                df_p_audit = pd.DataFrame(pronoun_audit_rows)
                unusual_cnt = sum(1 for st_v in df_p_audit['Trạng thái'] if '🟡' in str(st_v))

                col_pm1, col_pm2 = st.columns(2)
                with col_pm1:
                    st.markdown(f"""
                    <div class="metric-card">
                        <div class="metric-label">💬 Tổng số câu thoại đã quét</div>
                        <div class="metric-value">{len(df_p_audit)}</div>
                    </div>
                    """, unsafe_allow_html=True)
                with col_pm2:
                    st.markdown(f"""
                    <div class="metric-card">
                        <div class="metric-label">🟡 Câu nghi vấn lệch xưng hô</div>
                        <div class="metric-value" style="color:#D97706;">{unusual_cnt}</div>
                    </div>
                    """, unsafe_allow_html=True)

                st.markdown("##### 👁️ Bảng Báo Cáo Soát Lỗi Xưng Hô Chi Tiết")
                p_filter = st.checkbox("🟡 Chỉ hiển thị các câu nghi vấn lệch xưng hô", value=True)
                df_p_disp = df_p_audit[df_p_audit['Trạng thái'].str.contains('🟡', na=False)] if p_filter else df_p_audit

                st.dataframe(df_p_disp[['Stt', 'Timecode', 'Nhân vật', 'Câu thoại Tiếng Việt', 'Trạng thái', 'Chi tiết QC']], hide_index=True, use_container_width=True)

    # 2. SUBTAB THUẬT NGỮ & TÊN RIÊNG
    with subtab_glossary:
        st.markdown("#### 📚 Soát Bất Nhất Thuật Ngữ & Bản Dịch Tiếng Việt")
        st.caption("Tải file kịch bản để tự động tìm các từ/cụm từ Tiếng Anh lặp lại và kiểm tra xem có bị dịch thành nhiều nghĩa khác nhau hay không:")

        uploaded_glossary_script = st.file_uploader("Tải file Kịch bản (.srt hoặc .docx) để kiểm tra thuật ngữ:", type=['srt', 'docx'], key="uploader_glossary_qc")

        if uploaded_glossary_script is not None:
            c_spks_g = st.session_state.get('custom_speakers', set())
            c_non_spks_g = st.session_state.get('custom_non_speakers', set())
            df_g_script = parse_any_script_file_to_df(uploaded_glossary_script.getvalue(), uploaded_glossary_script.name, c_spks_g, c_non_spks_g)

            if not df_g_script.empty:
                all_dialogues_str = " ".join(df_g_script['Dialogue'].dropna().tolist())
                eng_matches = ENGLISH_WORD_REGEX.findall(all_dialogues_str)
                eng_counts = Counter([w for w in eng_matches if is_candidate_english_word(w)])

                glossary_audit_rows = []
                for word, cnt in eng_counts.most_common(20):
                    matching_lines = df_g_script[df_g_script['Dialogue'].str.contains(r'\b' + re.escape(word) + r'\b', case=False, na=False)]
                    sample_texts = matching_lines['Dialogue'].head(3).tolist()

                    glossary_audit_rows.append({
                        "Thuật ngữ / Tên riêng": word,
                        "Số lần lặp lại": cnt,
                        "Các câu thoại chứa từ này": " | ".join(sample_texts)
                    })

                if glossary_audit_rows:
                    df_g_audit = pd.DataFrame(glossary_audit_rows)
                    st.markdown("##### 📑 Bảng Thống Kê Thuật Ngữ Lặp Lại Trong Kịch Bản")
                    st.dataframe(df_g_audit, hide_index=True, use_container_width=True)
                else: st.info("Không phát hiện thuật ngữ Tiếng Anh nào lặp lại nhiều lần trong kịch bản này.")

# ==========================================
# TAB 8: DỌN DẸP & CHUẨN HÓA PHỤ ĐỀ
# ==========================================
with tab_cleaner:
    st.subheader("🧹 DỌN DẸP & CHUẨN HÓA PHỤ ĐỀ (TEXT NORMALIZER)")
    st.markdown("Tự động giặt sạch kịch bản rác, bóc tách thẻ HTML/ASS rác, sửa lỗi gõ phím, sửa lỗi dấu câu Tiếng Việt và thu gọn khoảng trắng dư thừa.")

    subtab_paste_clean, subtab_file_clean = st.tabs([
        "✍️ Xử Lý & Dọn Dẹp Văn Bản Trực Tiếp", 
        "📁 Dọn Dẹp & Chuẩn Hóa File Hàng Loạt (.srt / .docx)"
    ])

    # 1. PHÂN KHU 1: XỬ LÝ VĂN BẢN TRỰC TIẾP
    with subtab_paste_clean:
        st.markdown("#### 1. Chọn Quy Tắc Giặt Sạch Văn Bản")
        col_opt1, col_opt2, col_opt3, col_opt4 = st.columns(4)
        with col_opt1:
            opt_strip_html = st.checkbox("🏷️ Xóa sạch thẻ HTML/Color", value=True, help="Loại bỏ toàn bộ các thẻ màu <font color=...>, <span...>")
        with col_opt2:
            opt_fix_punct = st.checkbox("✍️ Sửa lỗi dấu câu Tiếng Việt", value=True, help="Sửa khoảng trắng trước/sau dấu phẩy, chấm, ngoặc kép, ba chấm")
        with col_opt3:
            opt_cap_first = st.checkbox("🔤 Viết hoa đầu câu & Sau tên", value=True, help="Viết hoa ký tự đầu tiên của câu thoại và sau nhãn 'Tyler:'")
        with col_opt4:
            opt_remove_dash = st.checkbox("✂️ Xóa dấu gạch đầu dòng '-'", value=True, help="Xóa dấu - ở đầu câu thoại phụ đề")

        st.markdown("---")
        col_text_in, col_text_out = st.columns(2)

        with col_text_in:
            st.markdown("##### 📥 Văn Bản Rác Đầu Vào (Paste văn bản vào đây):")
            input_raw_text = st.text_area(
                "Nội dung cần làm sạch:",
                value="<font color=\"red\">Cory :</font> Tránh xa  đồng đội tui ra !\n\n- <i>Garrett :</i> \" kịch tính quá \"\n\nTyler :chào bạn ,rất vui được  gặp...bạn",
                height=240,
                key="textarea_raw_clean_input"
            )
            
            btn_do_clean = st.button("🧹 THỰC THI DỌN DẸP VĂN BẢN", type="primary", use_container_width=True, key="btn_run_text_clean_manual")

        if btn_do_clean and input_raw_text:
            cleaned_res = clean_and_normalize_text(
                input_raw_text, 
                strip_all_tags=opt_strip_html, 
                fix_punctuation=opt_fix_punct, 
                normalize_spaces=True, 
                capitalize_first=opt_cap_first, 
                remove_leading_dash=opt_remove_dash
            )
            st.session_state['textarea_clean_output'] = cleaned_res
            st.session_state['manual_cleaned_orig_len'] = len(input_raw_text)
            st.session_state['manual_cleaned_res_len'] = len(cleaned_res)

        with col_text_out:
            st.markdown("##### 📤 Văn Bản Đã Làm Sạch Hoàn Hảo:")
            st.text_area(
                "Kết quả sau khi dọn dẹp:",
                height=240,
                key="textarea_clean_output"
            )

        if 'manual_cleaned_orig_len' in st.session_state:
            orig_c = st.session_state['manual_cleaned_orig_len']
            clean_c = st.session_state['manual_cleaned_res_len']
            diff_c = orig_c - clean_c
            st.success(f"✅ Đã dọn dẹp xong! Giảm **{diff_c}** ký tự rác/khoảng trắng thừa (Từ {orig_c} ➔ {clean_c} ký tự).")

    # 2. PHÂN KHU 2: DỌN DẸP FILE HÀNG LOẠT
    with subtab_file_clean:
        st.markdown("#### 📁 Dọn Dẹp File Phụ Đề & Kịch Bản Hàng Loạt")
        st.caption("Tải lên 1 hoặc nhiều file .srt / .docx bị rác định dạng. Hệ thống sẽ làm sạch văn bản từng câu thoại nhưng giữ nguyên 100% Timecode và Số thứ tự câu:")

        uploaded_clean_files = st.file_uploader(
            "Tải file .srt hoặc .docx cần giặt sạch:", 
            type=['srt', 'docx'], 
            accept_multiple_files=True, 
            key="batch_clean_file_uploader"
        )

        if uploaded_clean_files:
            st.info(f"Đã chọn **{len(uploaded_clean_files)}** file cần dọn dẹp.")
            if st.button("✨ BẮT ĐẦU GIẶT SẠCH CÁC FILE TRÊN", type="primary", use_container_width=True, key="btn_run_batch_clean"):
                try:
                    if len(uploaded_clean_files) == 1:
                        f_item = uploaded_clean_files[0]
                        f_name_no_ext = os.path.splitext(f_item.name)[0]
                        f_ext = os.path.splitext(f_item.name)[1].lower()

                        if f_ext == '.srt':
                            try: raw_str = f_item.getvalue().decode('utf-8')
                            except UnicodeDecodeError: raw_str = f_item.getvalue().decode('latin-1')
                            
                            cleaned_srt_lines = []
                            for block in re.split(r'\n\s*\n', raw_str.strip()):
                                lines_b = [l.strip() for l in block.strip().split('\n') if l.strip()]
                                if len(lines_b) < 2: continue
                                tc_idx = -1
                                for idx_b, l_b in enumerate(lines_b[:2]):
                                    if "-->" in l_b: tc_idx = idx_b; break
                                if tc_idx == -1: continue
                                
                                idx_line = lines_b[0] if tc_idx == 1 else ""
                                tc_line = lines_b[tc_idx]
                                diag_lines = lines_b[tc_idx+1:]
                                
                                clean_diag = clean_and_normalize_text("\n".join(diag_lines), strip_all_tags=True)
                                
                                out_b = ""
                                if idx_line: out_b += idx_line + "\n"
                                out_b += tc_line + "\n" + clean_diag
                                cleaned_srt_lines.append(out_b)
                                
                            out_bytes = "\n\n".join(cleaned_srt_lines).encode('utf-8-sig')
                            st.success("✅ Đã làm sạch file SRT thành công!")
                            st.download_button(
                                label=f"⬇️ TẢI FILE SRT SẠCH ({f_name_no_ext}_Clean.srt)",
                                data=out_bytes,
                                file_name=f"{f_name_no_ext}_Clean.srt",
                                mime="text/plain",
                                type="primary",
                                use_container_width=True
                            )
                        elif f_ext == '.docx':
                            doc = Document(io.BytesIO(f_item.getvalue()))
                            new_doc = Document()
                            timecode_pattern = re.compile(r'^\d{2}:\d{2}:\d{2}[,.]\d{3}\s*-->\s*\d{2}:\d{2}:\d{2}[,.]\d{3}$')

                            for p in doc.paragraphs:
                                txt = p.text.strip()
                                if not txt: continue
                                if timecode_pattern.match(txt) or txt.isdigit() or txt.lower().startswith("srt conversion"):
                                    p_out = new_doc.add_paragraph(txt)
                                    p_out.runs[0].font.name = 'Times New Roman'; p_out.runs[0].font.size = Pt(12)
                                    if timecode_pattern.match(txt) or txt.isdigit(): p_out.runs[0].bold = True
                                    p_out.paragraph_format.space_before = Pt(0); p_out.paragraph_format.space_after = Pt(0)
                                else:
                                    cleaned_p = clean_and_normalize_text(txt, strip_all_tags=True)
                                    p_out = new_doc.add_paragraph(cleaned_p)
                                    p_out.runs[0].font.name = 'Times New Roman'; p_out.runs[0].font.size = Pt(12)
                                    p_out.paragraph_format.space_before = Pt(0); p_out.paragraph_format.space_after = Pt(4)

                            doc_buf = io.BytesIO(); new_doc.save(doc_buf); doc_buf.seek(0)
                            st.success("✅ Đã làm sạch file Word DOCX thành công!")
                            st.download_button(
                                label=f"⬇️ TẢI FILE WORD SẠCH ({f_name_no_ext}_Clean.docx)",
                                data=doc_buf,
                                file_name=f"{f_name_no_ext}_Clean.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                type="primary",
                                use_container_width=True
                            )
                    else:
                        zip_clean_buf = io.BytesIO()
                        with zipfile.ZipFile(zip_clean_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
                            for f_item in uploaded_clean_files:
                                f_name_no_ext = os.path.splitext(f_item.name)[0]
                                f_ext = os.path.splitext(f_item.name)[1].lower()
                                if f_ext == '.srt':
                                    try: raw_str = f_item.getvalue().decode('utf-8')
                                    except UnicodeDecodeError: raw_str = f_item.getvalue().decode('latin-1')
                                    
                                    cleaned_srt_lines = []
                                    for block in re.split(r'\n\s*\n', raw_str.strip()):
                                        lines_b = [l.strip() for l in block.strip().split('\n') if l.strip()]
                                        if len(lines_b) < 2: continue
                                        tc_idx = -1
                                        for idx_b, l_b in enumerate(lines_b[:2]):
                                            if "-->" in l_b: tc_idx = idx_b; break
                                        if tc_idx == -1: continue
                                        
                                        idx_line = lines_b[0] if tc_idx == 1 else ""
                                        tc_line = lines_b[tc_idx]
                                        clean_diag = clean_and_normalize_text("\n".join(lines_b[tc_idx+1:]), strip_all_tags=True)
                                        out_b = f"{idx_line}\n" if idx_line else ""
                                        out_b += tc_line + "\n" + clean_diag
                                        cleaned_srt_lines.append(out_b)
                                    zf.writestr(f"{f_name_no_ext}_Clean.srt", "\n\n".join(cleaned_srt_lines).encode('utf-8-sig'))
                        zip_clean_buf.seek(0)
                        st.success(f"✅ Đã dọn dẹp thành công {len(uploaded_clean_files)} file!")
                        st.download_button(
                            label="📦 TẢI TRỌN BỘ FILE SẠCH (.ZIP)",
                            data=zip_clean_buf,
                            file_name="Cleaned_Files_Pack.zip",
                            mime="application/zip",
                            type="primary",
                            use_container_width=True
                        )
                except Exception as e: st.error(f"Lỗi dọn dẹp file: {e}")

# ==========================================
# TAB 9: BỘ CÔNG CỤ CHUYỂN ĐỔI (CONVERTER SUITE)
# ==========================================
with tab_tools:
    subtab_sub_conv, subtab_srt_excel, subtab_daw_markers, subtab_curr, subtab_dist, subtab_speed, subtab_mass_temp = st.tabs([
        "🎬 Kịch Bản Subtitle (SRT ⇄ DOCX)",
        "📊 SRT ➔ Excel (.xlsx)",
        "🎛️ DAW Marker Timeline",
        "💵 Tiền Tệ (Currency)",
        "📏 Khoảng Cách (Distance)",
        "🚀 Vận Tốc (Speed)",
        "⚖️ Khối Lượng & Nhiệt Độ"
    ])

    # 1. BỘ CHUYỂN ĐỔI SUBTITLE KỊCH BẢN (SRT ⇄ DOCX)
    with subtab_sub_conv:
        st.markdown("#### 🎬 Bộ Công Cụ Chuyển Đổi Subtitle Chuyên Nghiệp")
        col_c1, col_c2 = st.columns(2)
        
        with col_c1:
            with st.container(border=True):
                st.markdown("##### 📄 1. Chuyển SRT ➔ Word (.docx)")
                st.caption("Tải 1 hoặc hàng ngàn file SRT để tự động chuyển sang Word (Times New Roman, 12pt):")
                
                batch_srt_files = st.file_uploader(
                    "Tải 1 hoặc nhiều file .srt:", type=['srt'], accept_multiple_files=True, key="tool_srt_to_docx_batch"
                )
                
                if batch_srt_files:
                    st.info(f"Đã chọn **{len(batch_srt_files)}** file SRT.")
                    if st.button("✨ Chuyển SRT Sang Word", use_container_width=True, type="primary"):
                        try:
                            if len(batch_srt_files) == 1:
                                single_f = batch_srt_files[0]
                                s_name_no_ext = os.path.splitext(single_f.name)[0]
                                docx_buf = process_srt_to_docx(single_f, s_name_no_ext)
                                st.success("✅ Chuyển đổi hoàn tất!")
                                st.download_button(
                                    label=f"⬇️ Tải {s_name_no_ext}.docx", data=docx_buf,
                                    file_name=f"{s_name_no_ext}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
                            else:
                                zip_buf = io.BytesIO()
                                with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                                    for srt_f in batch_srt_files:
                                        s_name_no_ext = os.path.splitext(srt_f.name)[0]
                                        docx_buf = process_srt_to_docx(srt_f, s_name_no_ext)
                                        zf.writestr(f"{s_name_no_ext}.docx", docx_buf.getvalue())
                                zip_buf.seek(0)
                                st.success(f"✅ Đã chuyển đổi thành công {len(batch_srt_files)} file!")
                                st.download_button(
                                    label="📦 Tải Trọn Bộ Word (.ZIP)", data=zip_buf.getvalue(),
                                    file_name="Word_Files.zip", mime="application/zip", use_container_width=True
                                )
                        except Exception as e: st.error(f"Lỗi: {e}")

        with col_c2:
            with st.container(border=True):
                st.markdown("##### 📝 2. Chuyển Word (.docx) ➔ SRT (Batch hàng loạt)")
                st.caption("Tải 1 file hoặc hàng ngàn file Word kịch bản để tự động trích xuất SRT:")
                
                batch_docx_files = st.file_uploader(
                    "Tải 1 hoặc nhiều file .docx:", type=['docx'], accept_multiple_files=True, key="tool_docx_to_srt_batch"
                )
                
                if batch_docx_files:
                    st.info(f"Đã chọn **{len(batch_docx_files)}** file Word.")
                    if st.button("✨ Chuyển Hàng Loạt Sang SRT", use_container_width=True, type="primary"):
                        try:
                            if len(batch_docx_files) == 1:
                                single_f = batch_docx_files[0]
                                s_name_no_ext = os.path.splitext(single_f.name)[0]
                                srt_bytes = process_docx_to_srt(single_f)
                                st.success("✅ Chuyển đổi hoàn tất!")
                                st.download_button(
                                    label=f"⬇️ Tải {s_name_no_ext}.srt", data=srt_bytes,
                                    file_name=f"{s_name_no_ext}.srt", mime="text/plain", use_container_width=True
                                )
                            else:
                                zip_buf = io.BytesIO()
                                with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                                    for doc_f in batch_docx_files:
                                        s_name_no_ext = os.path.splitext(doc_f.name)[0]
                                        srt_bytes = process_docx_to_srt(doc_f)
                                        zf.writestr(f"{s_name_no_ext}.srt", srt_bytes)
                                zip_buf.seek(0)
                                st.success(f"✅ Đã chuyển đổi thành công {len(batch_docx_files)} file!")
                                st.download_button(
                                    label="📦 Tải Trọn Bộ Word (.ZIP)", data=zip_buf.getvalue(),
                                    file_name="Word_Files.zip", mime="application/zip", use_container_width=True
                                )
                        except Exception as e: st.error(f"Lỗi: {e}")

    # 2. BỘ CHUYỂN ĐỔI SRT TO EXCEL WITH SPEAKER STYLING
    with subtab_srt_excel:
        st.markdown("#### 📊 Chuyển Đổi File Subtitle SRT ➔ Bảng Tính Excel (.xlsx)")
        st.caption("Tự động nhận diện nhân vật, tô màu phân biệt người nói và xuất file Excel có cấu trúc:")
        
        uploaded_srt_excel = st.file_uploader("Tải file .srt của bạn vào đây:", type=['srt'], key="tool_srt_to_excel")
        if uploaded_srt_excel is not None:
            try:
                try: srt_content_excel = uploaded_srt_excel.read().decode("utf-8")
                except UnicodeDecodeError: srt_content_excel = uploaded_srt_excel.read().decode("latin-1")
            except Exception:
                st.error("Lỗi mã hóa file. Vui lòng đảm bảo file SRT của bạn ở chuẩn mã hóa UTF-8.")
                srt_content_excel = None

            if srt_content_excel:
                custom_spks_ex = st.session_state.get('custom_speakers', set())
                custom_non_spks_ex = st.session_state.get('custom_non_speakers', set())

                srt_speaker_counts = Counter()
                for line_s in srt_content_excel.split('\n'):
                    line_clean = line_s.strip()
                    if not line_clean or TIMECODE_REGEX.match(line_clean) or line_clean.isdigit(): continue
                    spk_tags_found = find_all_speaker_tags(line_clean, custom_spks_ex, custom_non_spks_ex)
                    for _, _, spk_cand, _ in spk_tags_found:
                        srt_speaker_counts[spk_cand] += 1

                detected_srt_spk_names = [name for name in srt_speaker_counts.keys() if name.upper() not in NON_SPEAKER_PHRASES]
                detected_srt_non_spk_names = [name for name in srt_speaker_counts.keys() if name.upper() in NON_SPEAKER_PHRASES]

                detected_srt_spk_disp = [f"{name} ({srt_speaker_counts[name]} lần)" for name in detected_srt_spk_names]
                detected_srt_non_spk_disp = [f"{name} ({srt_speaker_counts[name]} lần)" for name in detected_srt_non_spk_names]

                with st.container(border=True):
                    st.markdown("### 🔍 Soát Lỗi Nhận Diện Tên Người Nói (SRT)")
                    st.caption("Kiểm tra danh sách tên người nói bóc tách từ file SRT. Chọn từ bị nhận diện sai để nạp trực tiếp vào Database:")
                    tab_srt_spk, tab_srt_non_spk = st.tabs(["🎭 Nhận diện là NGƯỜI NÓI", "🚫 Đang bị xem là TỪ NHIỄU"])

                    with tab_srt_spk:
                        if detected_srt_spk_disp:
                            st.write(", ".join([f"`{s}`" for s in detected_srt_spk_disp]))
                            to_move_ns_srt = st.multiselect("Phát hiện từ nào bị nhận diện sai? Chọn để LƯU VÀO DATABASE TỪ NHIỄU:", options=detected_srt_spk_names, key="select_srt_to_ns")
                            if st.button("➡️ Đưa vào Database TỪ NHIỄU", type="secondary", key="btn_srt_to_ns"):
                                if to_move_ns_srt:
                                    new_items = [item.upper() for item in to_move_ns_srt]
                                    st.session_state['custom_non_speakers'].update(new_items)
                                    save_json_db(NON_SPEAKER_DB_FILE, st.session_state['custom_non_speakers'])
                                    st.success(f"✅ Đã lưu {len(new_items)} từ vào Database Từ Nhiễu!")
                                    time.sleep(1); st.rerun()
                        else: st.info("Chưa tìm thấy cụm từ người nói nào trong file SRT.")

                    with tab_srt_non_spk:
                        if detected_srt_non_spk_disp:
                            st.write(", ".join([f"`{s}`" for s in detected_srt_non_spk_disp]))
                            to_move_spk_srt = st.multiselect("Từ nào thực ra là NGƯỜI NÓI? Chọn để LƯU VÀO DATABASE NGƯỜI NÓI:", options=detected_srt_non_spk_names, key="select_srt_to_spk")
                            if st.button("➡️ Đưa vào Database NGƯỜI NÓI", type="secondary", key="btn_srt_to_spk"):
                                if to_move_spk_srt:
                                    st.session_state['custom_speakers'].update(to_move_spk_srt)
                                    save_json_db(SPEAKER_DB_FILE, st.session_state['custom_speakers'])
                                    for item in to_move_spk_srt: st.session_state['custom_non_speakers'].discard(item.upper())
                                    save_json_db(NON_SPEAKER_DB_FILE, st.session_state['custom_non_speakers'])
                                    st.success(f"✅ Đã lưu {len(to_move_spk_srt)} tên vào Database Người Nói!")
                                    time.sleep(1); st.rerun()
                        else: st.info("Không có cụm từ nào bị loại vào danh sách từ nhiễu.")

                st.markdown("---")

                with st.spinner('Đang phân tích dữ liệu SRT...'):
                    df_converted_excel = parse_srt_to_dataframe(srt_content_excel, custom_spks_ex, custom_non_spks_ex)
                
                if df_converted_excel.empty:
                    st.error("Không thể đọc được phụ đề nào từ file SRT này.")
                else:
                    st.markdown("##### 📊 Thống Kê Nhân Vật Trong File Excel")
                    unique_spks = df_converted_excel['Speaker'].unique()
                    actual_spks = [s for s in unique_spks if s not in ["Unknown", ""]]
                    
                    st.success(f"**Tổng số Người nói được nhận dạng:** {len(actual_spks)} người.")
                    if actual_spks: st.markdown(f"**Danh sách Người nói:** {', '.join(actual_spks)}")
                    else: st.info("Không tìm thấy người nói rõ ràng (ngoại trừ các đoạn hội thoại không gắn tên).")

                    st.markdown("##### 👁️ Xem Trước Bảng Dữ Liệu Chuyển Đổi")
                    styled_excel_df = apply_excel_styles(df_converted_excel)
                    st.dataframe(styled_excel_df, use_container_width=True)

                    output_excel = io.BytesIO()
                    styled_excel_df.to_excel(output_excel, index=False, engine='openpyxl')
                    output_excel.seek(0)

                    orig_base_name = uploaded_srt_excel.name.rsplit('.', 1)[0]
                    excel_out_filename = f"{orig_base_name}.xlsx"
                    
                    st.download_button(
                        label=f"💾 TẢI FILE EXCEL (.XLSX): {excel_out_filename}", data=output_excel.getvalue(),
                        file_name=excel_out_filename, mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        type="primary", use_container_width=True
                    )

    # 3. BỘ XUẤT DAW MARKER TIMELINE (PRO TOOLS / REAPER / CMX 3600 EDL)
    with subtab_daw_markers:
        st.markdown("#### 🎛️ Tự Động Tạo File Marker Timeline Cho Phần Mềm Thu Âm DAW")
        st.caption("Chuyển đổi kịch bản (.srt hoặc .docx) thành các điểm mốc Marker phủ màu sẵn cho KTV thu âm trên Pro Tools, Reaper, Premiere, Resolve:")

        uploaded_marker_file = st.file_uploader("Tải file Kịch bản (.srt hoặc .docx) của bạn vào đây:", type=['srt'], key="tool_marker_uploader")
        
        if uploaded_marker_file is not None:
            m_filename = uploaded_marker_file.name
            m_base_name = os.path.splitext(m_filename)[0]
            custom_spks_m = st.session_state.get('custom_speakers', set())
            custom_non_spks_m = st.session_state.get('custom_non_speakers', set())
            
            if m_filename.endswith('.srt'):
                try: m_srt_text = uploaded_marker_file.read().decode("utf-8")
                except UnicodeDecodeError: m_srt_text = uploaded_marker_file.read().decode("latin-1")
                df_markers = parse_srt_to_dataframe(m_srt_text, custom_spks_m, custom_non_spks_m)
            else:
                s_bytes = process_docx_to_srt(uploaded_marker_file)
                m_srt_text = s_bytes.decode("utf-8", errors="ignore")
                df_markers = parse_srt_to_dataframe(m_srt_text, custom_spks_m, custom_non_spks_m)

            if not df_markers.empty:
                st.success(f"✅ Đã trích xuất **{len(df_markers)}** câu thoại từ **{m_filename}**!")
                
                col_m1, col_m2, col_m3 = st.columns(3)
                
                with col_m1:
                    with st.container(border=True):
                        st.markdown("##### 🎧 1. REAPER (Region CSV)")
                        st.caption("Tải file .csv để import trực tiếp vào REAPER Marker Manager:")
                        reaper_csv_str = generate_reaper_region_csv(df_markers)
                        st.download_button(
                            label=f"⬇️ Tải {m_base_name}_Reaper.csv",
                            data=reaper_csv_str.encode('utf-8-sig'),
                            file_name=f"{m_base_name}_Reaper.csv",
                            mime="text/csv", type="primary", use_container_width=True
                        )

                with col_m2:
                    with st.container(border=True):
                        st.markdown("##### 🎛️ 2. PRO TOOLS (Track Markers)")
                        st.caption("Dùng cho Pro Tools 2023.6+ Import Track Markers:")
                        pt_csv_str = generate_pro_tools_csv(df_markers)
                        st.download_button(
                            label=f"⬇️ Tải {m_base_name}_ProTools.csv",
                            data=pt_csv_str.encode('utf-8-sig'),
                            file_name=f"{m_base_name}_ProTools.csv",
                            mime="text/csv", type="primary", use_container_width=True
                        )

                with col_m3:
                    with st.container(border=True):
                        st.markdown("##### 🎬 3. CMX 3600 EDL (Premiere/Resolve)")
                        st.caption("Chuẩn EDL đa năng cho Premiere Pro, Resolve, Vegas, Nuendo:")
                        edl_str = generate_cmx3600_edl(df_markers)
                        st.download_button(
                            label=f"⬇️ Tải {m_base_name}.edl",
                            data=edl_str.encode('utf-8-sig'),
                            file_name=f"{m_base_name}.edl",
                            mime="text/plain", type="primary", use_container_width=True
                        )

    # 4. BỘ CHUYỂN ĐỔI TIỀN TỆ (CURRENCY)
    with subtab_curr:
        st.markdown("#### 💵 Quy Đổi Tiền Tệ Đa Ngoại Tệ")
        rates = {
            "VND": 1.0, "USD": 25400.0, "EUR": 27500.0, "GBP": 32000.0, "JPY": 165.0,
            "CNY": 3500.0, "KRW": 18.5, "AUD": 16800.0, "CAD": 18200.0, "SGD": 18900.0
        }
        c_col1, c_col2, c_col3 = st.columns([2, 1.5, 1.5])
        with c_col1: curr_amount = st.number_input("Số lượng tiền cần đổi:", value=100.0, min_value=0.0, step=10.0)
        with c_col2: from_curr = st.selectbox("Từ đồng tiền:", options=list(rates.keys()), index=1)
        with c_col3: to_curr = st.selectbox("Sang đồng tiền:", options=list(rates.keys()), index=0)
            
        amount_in_vnd = curr_amount * rates[from_curr]
        result_curr = amount_in_vnd / rates[to_curr]
        
        st.markdown("---")
        st.markdown(f"### 🎯 Kết Quả: **{curr_amount:,.2f} {from_curr}** = **{result_curr:,.2f} {to_curr}**")
        st.caption(f"Tỷ giá tham chiếu: 1 USD = {rates['USD']:,.0f} VND | 1 EUR = {rates['EUR']:,.0f} VND | 1 JPY = {rates['JPY']:,.1f} VND")

    # 5. BỘ CHUYỂN ĐỔI KHOẢNG CÁCH (DISTANCE)
    with subtab_dist:
        st.markdown("#### 📏 Quy Đổi Đơn Vị Khoảng Cách")
        dist_factors = {
            "Millimet (mm)": 0.001, "Centimet (cm)": 0.01, "Mét (m)": 1.0, "Kilômét (km)": 1000.0,
            "Inch (in)": 0.0254, "Foot (ft)": 0.3048, "Yard (yd)": 0.9144, "Dặm (Mile)": 1609.344
        }
        d_col1, d_col2, d_col3 = st.columns([2, 1.5, 1.5])
        with d_col1: dist_val = st.number_input("Giá trị khoảng cách:", value=1.0, min_value=0.0, step=1.0)
        with d_col2: from_dist = st.selectbox("Từ đơn vị:", options=list(dist_factors.keys()), index=3)
        with d_col3: to_dist = st.selectbox("Sang đơn vị:", options=list(dist_factors.keys()), index=2)
            
        meters = dist_val * dist_factors[from_dist]
        res_dist = meters / dist_factors[to_dist]
        
        st.markdown("---")
        st.markdown(f"### 🎯 Kết Quả: **{dist_val:,.4f} {from_dist}** = **{res_dist:,.4f} {to_dist}**")

    # 6. BỘ CHUYỂN ĐỔI VẬN TỐC (SPEED)
    with subtab_speed:
        st.markdown("#### 🚀 Quy Đổi Đơn Vị Vận Tốc")
        speed_factors = {
            "Mét/giây (m/s)": 1.0, "Kilômét/giờ (km/h)": 1 / 3.6,
            "Dặm/giờ (mph)": 0.44704, "Hải lý/giờ (Knot)": 0.514444
        }
        s_col1, s_col2, s_col3 = st.columns([2, 1.5, 1.5])
        with s_col1: speed_val = st.number_input("Giá trị vận tốc:", value=100.0, min_value=0.0, step=5.0)
        with s_col2: from_speed = st.selectbox("Từ đồng tiền:", options=list(speed_factors.keys()), index=1)
        with s_col3: to_speed = st.selectbox("Sang đồng tiền:", options=list(speed_factors.keys()), index=0)
            
        ms_val = speed_val * speed_factors[from_speed]
        res_speed = ms_val / speed_factors[to_speed]
        
        st.markdown("---")
        st.markdown(f"### 🎯 Kết Quả: **{speed_val:,.2f} {from_speed}** = **{res_speed:,.2f} {to_speed}**")

    # 7. KHỐI LƯỢNG & NHIỆT ĐỘ
    with subtab_mass_temp:
        m_col1, m_col2 = st.columns(2)
        with m_col1:
            with st.container(border=True):
                st.markdown("##### ⚖️ Quy Đổi Khối Lượng")
                mass_factors = {
                    "Gram (g)": 0.001, "Kilôgram (kg)": 1.0, "Tấn": 1000.0,
                    "Ounce (oz)": 0.0283495, "Pound (lb)": 0.453592
                }
                m_val = st.number_input("Khối lượng:", value=1.0, min_value=0.0, key="m_val_in")
                m_from = st.selectbox("Từ:", options=list(mass_factors.keys()), index=1, key="m_from_sel")
                m_to = st.selectbox("Sang:", options=list(mass_factors.keys()), index=4, key="m_to_sel")
                
                kg_val = m_val * mass_factors[m_from]
                res_mass = kg_val / mass_factors[m_to]
                st.info(f"👉 **{m_val:,.2f} {m_from}** = **{res_mass:,.2f} {m_to}**")
                
        with m_col2:
            with st.container(border=True):
                st.markdown("##### 🌡️ Quy Đổi Nhiệt Độ")
                temp_val = st.number_input("Nhiệt độ:", value=37.0, key="temp_val_in")
                t_from = st.selectbox("Từ:", options=["Độ C (°C)", "Độ F (°F)", "Kelvin (K)"], index=0, key="t_from_sel")
                t_to = st.selectbox("Sang:", options=["Độ C (°C)", "Độ F (°F)", "Kelvin (K)"], index=1, key="t_to_sel")
                
                if "°C" in t_from: c_temp = temp_val
                elif "°F" in t_from: c_temp = (temp_val - 32) * 5 / 9
                else: c_temp = temp_val - 273.15
                
                if "°C" in t_to: res_temp = c_temp
                elif "°F" in t_to: res_temp = (c_temp * 9 / 5) + 32
                else: res_temp = c_temp + 273.15
                
                st.info(f"👉 **{temp_val:,.1f} {t_from}** = **{res_temp:,.1f} {t_to}**")

# ==========================================
# 6. FOOTER SAAS TÙY CHỈNH
# ==========================================
st.markdown("""
<div class="saas-footer">
    Copyright © Mai Han Team. All Rights Reserved.
</div>
""", unsafe_allow_html=True)
