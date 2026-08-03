import streamlit as st
import io
import os
import re
import zipfile
from docx import Document
from docx.shared import Pt
from utils import clean_and_normalize_text

def render_tab8():
    st.subheader("🧹 DỌN DẸP & CHUẨN HÓA PHỤ ĐỀ (TEXT NORMALIZER)")
    st.markdown("Tự động giặt sạch kịch bản rác, bóc tách thẻ HTML/ASS rác, sửa lỗi gõ phím, sửa lỗi dấu câu Tiếng Việt và thu gọn khoảng trắng dư thừa.")

    subtab_paste_clean, subtab_file_clean = st.tabs([
        "✍️ Xử Lý & Dọn Dẹp Văn Bản Trực Tiếp", 
        "📁 Dọn Dẹp & Chuẩn Hóa File Hàng Loạt (.srt / .docx)"
    ])

    with subtab_paste_clean:
        st.markdown("#### 1. Chọn Quy Tắc Giặt Sạch Văn Bản")
        col_opt1, col_opt2, col_opt3, col_opt4 = st.columns(4)
        with col_opt1: opt_strip_html = st.checkbox("🏷️ Xóa sạch thẻ HTML/Color", value=True)
        with col_opt2: opt_fix_punct = st.checkbox("✍️ Sửa lỗi dấu câu Tiếng Việt", value=True)
        with col_opt3: opt_cap_first = st.checkbox("🔤 Viết hoa đầu câu & Sau tên", value=True)
        with col_opt4: opt_remove_dash = st.checkbox("✂️ Xóa dấu gạch đầu dòng '-'", value=True)

        st.markdown("---")
        col_text_in, col_text_out = st.columns(2)

        with col_text_in:
            st.markdown("##### 📥 Văn Bản Rác Đầu Vào (Paste văn bản vào đây):")
            input_raw_text = st.text_area(
                "Nội dung cần làm sạch:",
                value="<font color=\"red\">Cory :</font> Tránh xa  đồng đội tui ra !\n\n- <i>Garrett :</i> \" kịch tính quá \"\n\nTyler :chào bạn ,rất vui được  gặp...bạn",
                height=240,
                key="textarea_raw_clean_input"
            )
            btn_do_clean = st.button("🧹 THỰC THI DỌN DẸP VĂN BẢN", type="primary", use_container_width=True, key="btn_run_text_clean_manual")

        if btn_do_clean and input_raw_text:
            cleaned_res = clean_and_normalize_text(
                input_raw_text, 
                strip_all_tags=opt_strip_html, 
                fix_punctuation=opt_fix_punct, 
                normalize_spaces=True, 
                capitalize_first=opt_cap_first, 
                remove_leading_dash=opt_remove_dash
            )
            st.session_state['textarea_clean_output'] = cleaned_res
            st.session_state['manual_cleaned_orig_len'] = len(input_raw_text)
            st.session_state['manual_cleaned_res_len'] = len(cleaned_res)

        with col_text_out:
            st.markdown("##### 📤 Văn Bản Đã Làm Sạch Hoàn Hảo:")
            out_val = st.session_state.get('textarea_clean_output', "")
            st.text_area("Kết quả sau khi dọn dẹp:", value=out_val, height=240, key="textarea_clean_output")

        if 'manual_cleaned_orig_len' in st.session_state:
            orig_c = st.session_state['manual_cleaned_orig_len']
            clean_c = st.session_state['manual_cleaned_res_len']
            diff_c = orig_c - clean_c
            st.success(f"✅ Đã dọn dẹp xong! Giảm **{diff_c}** ký tự rác/khoảng trắng thừa (Từ {orig_c} ➔ {clean_c} ký tự).")

    with subtab_file_clean:
        st.markdown("#### 📁 Dọn Dẹp File Phụ Đề & Kịch Bản Hàng Loạt")
        uploaded_clean_files = st.file_uploader("Tải file .srt hoặc .docx cần giặt sạch:", type=['srt', 'docx'], accept_multiple_files=True, key="batch_clean_file_uploader")

        if uploaded_clean_files:
            st.info(f"Đã chọn **{len(uploaded_clean_files)}** file cần dọn dẹp.")
            if st.button("✨ BẮT ĐẦU GIẶT SẠCH CÁC FILE TRÊN", type="primary", use_container_width=True, key="btn_run_batch_clean"):
                try:
                    if len(uploaded_clean_files) == 1:
                        f_item = uploaded_clean_files[0]
                        f_name_no_ext = os.path.splitext(f_item.name)[0]
                        f_ext = os.path.splitext(f_item.name)[1].lower()

                        if f_ext == '.srt':
                            try: raw_str = f_item.getvalue().decode('utf-8')
                            except UnicodeDecodeError: raw_str = f_item.getvalue().decode('latin-1')
                            
                            cleaned_srt_lines = []
                            for block in re.split(r'\n\s*\n', raw_str.strip()):
                                lines_b = [l.strip() for l in block.strip().split('\n') if l.strip()]
                                if len(lines_b) < 2: continue
                                tc_idx = -1
                                for idx_b, l_b in enumerate(lines_b[:2]):
                                    if "-->" in l_b: tc_idx = idx_b; break
                                if tc_idx == -1: continue
                                
                                idx_line = lines_b[0] if tc_idx == 1 else ""
                                tc_line = lines_b[tc_idx]
                                clean_diag = clean_and_normalize_text("\n".join(lines_b[tc_idx+1:]), strip_all_tags=True)
                                out_b = f"{idx_line}\n" if idx_line else ""
                                out_b += tc_line + "\n" + clean_diag
                                cleaned_srt_lines.append(out_b)
                                
                            out_bytes = "\n\n".join(cleaned_srt_lines).encode('utf-8-sig')
                            st.success("✅ Đã làm sạch file SRT thành công!")
                            st.download_button(
                                label=f"⬇️ TẢI FILE SRT SẠCH ({f_name_no_ext}_Clean.srt)", data=out_bytes,
                                file_name=f"{f_name_no_ext}_Clean.srt", mime="text/plain", type="primary", use_container_width=True
                            )
                        elif f_ext == '.docx':
                            doc = Document(io.BytesIO(f_item.getvalue()))
                            new_doc = Document()
                            timecode_pattern = re.compile(r'^\d{2}:\d{2}:\d{2}[,.]\d{3}\s*-->\s*\d{2}:\d{2}:\d{2}[,.]\d{3}$')

                            for p in doc.paragraphs:
                                txt = p.text.strip()
                                if not txt: continue
                                if timecode_pattern.match(txt) or txt.isdigit() or txt.lower().startswith("srt conversion"):
                                    p_out = new_doc.add_paragraph(txt)
                                    p_out.runs[0].font.name = 'Times New Roman'; p_out.runs[0].font.size = Pt(12)
                                    if timecode_pattern.match(txt) or txt.isdigit(): p_out.runs[0].bold = True
                                    p_out.paragraph_format.space_before = Pt(0); p_out.paragraph_format.space_after = Pt(0)
                                else:
                                    cleaned_p = clean_and_normalize_text(txt, strip_all_tags=True)
                                    p_out = new_doc.add_paragraph(cleaned_p)
                                    p_out.runs[0].font.name = 'Times New Roman'; p_out.runs[0].font.size = Pt(12)
                                    p_out.paragraph_format.space_before = Pt(0); p_out.paragraph_format.space_after = Pt(4)

                            doc_buf = io.BytesIO(); new_doc.save(doc_buf); doc_buf.seek(0)
                            st.success("✅ Đã làm sạch file Word DOCX thành công!")
                            st.download_button(
                                label=f"⬇️ TẢI FILE WORD SẠCH ({f_name_no_ext}_Clean.docx)", data=doc_buf,
                                file_name=f"{f_name_no_ext}_Clean.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                type="primary", use_container_width=True
                            )
                    else:
                        zip_clean_buf = io.BytesIO()
                        with zipfile.ZipFile(zip_clean_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
                            for f_item in uploaded_clean_files:
                                f_name_no_ext = os.path.splitext(f_item.name)[0]
                                f_ext = os.path.splitext(f_item.name)[1].lower()
                                if f_ext == '.srt':
                                    try: raw_str = f_item.getvalue().decode('utf-8')
                                    except UnicodeDecodeError: raw_str = f_item.getvalue().decode('latin-1')
                                    
                                    cleaned_srt_lines = []
                                    for block in re.split(r'\n\s*\n', raw_str.strip()):
                                        lines_b = [l.strip() for l in block.strip().split('\n') if l.strip()]
                                        if len(lines_b) < 2: continue
                                        tc_idx = -1
                                        for idx_b, l_b in enumerate(lines_b[:2]):
                                            if "-->" in l_b: tc_idx = idx_b; break
                                        if tc_idx == -1: continue
                                        
                                        idx_line = lines_b[0] if tc_idx == 1 else ""
                                        tc_line = lines_b[tc_idx]
                                        clean_diag = clean_and_normalize_text("\n".join(lines_b[tc_idx+1:]), strip_all_tags=True)
                                        out_b = f"{idx_line}\n" if idx_line else ""
                                        out_b += tc_line + "\n" + clean_diag
                                        cleaned_srt_lines.append(out_b)
                                    zf.writestr(f"{f_name_no_ext}_Clean.srt", "\n\n".join(cleaned_srt_lines).encode('utf-8-sig'))
                        zip_clean_buf.seek(0)
                        st.success(f"✅ Đã dọn dẹp thành công {len(uploaded_clean_files)} file!")
                        st.download_button(
                            label="📦 TẢI TRỌN BỘ FILE SẠCH (.ZIP)", data=zip_clean_buf,
                            file_name="Cleaned_Files_Pack.zip", mime="application/zip", type="primary", use_container_width=True
                        )
                except Exception as e: st.error(f"Lỗi dọn dẹp file: {e}")
