import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_COLOR_INDEX
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import io
import os
import re
import random
from collections import Counter
import time
import pandas as pd
import json
from gtts import gTTS
import zipfile

# ==========================================
# CẤU HÌNH DATABASE & FILE PATH
# ==========================================
NON_SPEAKER_DB_FILE = "custom_non_speakers.json"
SPEAKER_DB_FILE = "custom_speakers.json"
PHONETIC_DB_FILE = "custom_phonetics.json"
CAST_DB_FILE = "custom_cast_mapping.json"
TRACKER_DB_FILE = "dubbing_tracker.json"
RATES_DB_FILE = "payroll_rates.json"
PRONOUN_REL_DB_FILE = "custom_pronoun_relationships.json"
SPEAKER_COLOR_DB_FILE = "custom_speaker_colors.json"

DEFAULT_CAST_MAPPING = {
    "BRI": "TRÚC", "CHASE": "THIỆN", "PRESTON": "KHÁNH", "SCOTT": "THÔNG",
    "KEELEY": "TÚ", "DAKA": "QUANG", "STEPHEN": "QUANG", "VINCE": "QUANG",
    "JOSH": "THÔNG", "YOMI": "TÚ", "LARRY": "A.TRUNG", "CHLOE": "TÚ",
    "STEVEN": "QUANG (HOẶC THÔNG)", "LOGAN": "THÔNG", "NATHAN": "A.TRUNG",
    "RILEY": "PHỤNG", "ZHONG": "THÔNG", "MICHELLE (YOUTUBER)": "TÚ",
    "BEN AZELART": "QUANG", "ZONG": "THÔNG", "BA BRI": "THÔNG", "TYLER": "PHÁT",
    "COBY": "THIỆN", "CORY": "KHÁNH", "SPARKY": "A.TRUNG", "GARRETT": "C.DŨNG",
    "CODY": "CƯỜNG", "BETHANY": "TRÚC", "KRISTIN": "PHỤNG", "AMY": "TÚ",
    "ALLISON": "TÚ", "AUBREY": "PHỤNG", "TY JACKSON": "THÔNG", "HLV": "CƯỜNG",
    "JACKSON OLSON": "THÔNG", "DANNY": "QUANG", "KYLE": "QUANG", "BILL": "HITA", "TIM": "HITA"
}

DEFAULT_FIXED_SPEAKER_COLORS = {
    "ALL": {"text_color": (255, 0, 0), "highlight_color": (255, 255, 0)},
    "BEN AZELART": {"text_color": (21, 96, 130), "highlight_color": None},
    "BETHANY": {"text_color": (255, 192, 0), "highlight_color": None},
    "BRI": {"text_color": (116, 166, 123), "highlight_color": None},
    "CALEB": {"text_color": (116, 56, 25), "highlight_color": None},
    "CHASE": {"text_color": (255, 0, 255), "highlight_color": None},
    "CHUNKZ": {"text_color": (21, 96, 130), "highlight_color": None},
    "COBY": {"text_color": (255, 0, 255), "highlight_color": None},
    "CODY": {"text_color": (233, 113, 50), "highlight_color": None},
    "CORY": {"text_color": (255, 0, 0), "highlight_color": None},
    "DAKA": {"text_color": (110, 196, 229), "highlight_color": None},
    "GARRETT": {"text_color": (243, 243, 243), "highlight_color": None},
    "JOSH": {"text_color": (160, 43, 147), "highlight_color": None},
    "KEELEY": {"text_color": (255, 192, 0), "highlight_color": None},
    "LARRY": {"text_color": (243, 243, 243), "highlight_color": None},
    "LOGAN": {"text_color": (216, 170, 211), "highlight_color": None},
    "NATHAN": {"text_color": (21, 96, 130), "highlight_color": None},
    "NICK": {"text_color": (255, 0, 0), "highlight_color": None},
    "PRESTON": {"text_color": (255, 0, 0), "highlight_color": None},
    "RILEY": {"text_color": (116, 166, 123), "highlight_color": None},
    "SCOTT": {"text_color": (233, 113, 50), "highlight_color": None},
    "SPARKY": {"text_color": (116, 56, 25), "highlight_color": None},
    "STEPHEN": {"text_color": (0, 51, 204), "highlight_color": None},
    "STEVEN": {"text_color": (0, 255, 255), "highlight_color": None},
    "TYLER": {"text_color": (160, 43, 147), "highlight_color": None},
    "YOMI": {"text_color": (0, 153, 153), "highlight_color": None}
}

DEFAULT_SOUTH_VIETNAM_PHONETICS = {
    "MCDONALD": "Mắc-đô-nồ", "MCDONALD'S": "Mắc-đô-nồ", "LONDON": "Luân Đôn", "TNT": "Ti-èn-ti",
    "NOOB": "Núp", "GOLEM": "Gô-lùm", "GOOGLE": "Gú-gồ", "FACEBOOK": "Phây-sbúc",
    "KFC": "Ke-ép-xi", "STARBUCKS": "Xì-ta-bắc", "APPLE": "Ép-pồ", "BURGER": "Bơ-gơ",
    "PARIS": "Ba Lê", "WASHINGTON": "Hoa Thịnh Đốn", "CALIFORNIA": "Cả-li", "PRO": "Prồ",
    "LAG": "Lác", "BUFF": "Bớp", "VIP": "Vi-ai-pi", "FBI": "Ép-bi-ai"
}

VN_SYLLABLES = {
    "a", "ai", "an", "ang", "anh", "ao", "ap", "at", "au", "ay", "ba", "bac", "bai", "ban", "bang", "bao", 
    "bat", "bay", "be", "ben", "beng", "beo", "bi", "bic", "bien", "biet", "binh", "bo", "boc", "boi", "bon", 
    "bong", "bot", "bu", "bua", "bui", "bun", "buoc", "buon", "ca", "cac", "cai", "cam", "can", "cang", "cao", 
    "cap", "cat", "cau", "cay", "cha", "chai", "cham", "chan", "chang", "chao", "chap", "chat", "chau", "chay", 
    "che", "chen", "cheo", "chi", "chia", "chiem", "chieu", "chin", "chinh", "cho", "choc", "choi", "chon", 
    "chong", "chu", "chua", "chuan", "chuc", "chui", "chum", "chun", "chuoc", "chuon", "chura", "chuong", "co", 
    "coc", "coi", "con", "cong", "cot", "cu", "cua", "cuc", "cui", "cum", "cung", "cuoc", "cuoi", "cuon", 
    "cuong", "cuot", "da", "dac", "dai", "dam", "dan", "dang", "dao", "dap", "dat", "dau", "day", "de", "den", 
    "deo", "di", "dia", "diem", "dien", "diet", "dieu", "dinh", "do", "doc", "doi", "don", "dong", "dot", "du", 
    "dua", "duc", "dui", "dum", "dung", "duoc", "duoi", "duon", "duong", "em", "ga", "gac", "gai", "gam", "gan", 
    "gang", "gao", "gap", "gat", "gau", "gay", "ge", "gen", "gi", "gia", "giac", "giai", "giam", "gian", "giang", 
    "giao", "giap", "giat", "giau", "giay", "gio", "gioc", "gioi", "gion", "giong", "giu", "giua", "go", "goc", 
    "goi", "gon", "gong", "gu", "gua", "guc", "gui", "gum", "gung", "ha", "hac", "hai", "ham", "han", "hang", 
    "hao", "hap", "hat", "hau", "hay", "he", "hen", "heo", "hi", "hien", "hiet", "hieu", "hinh", "ho", "hoc", 
    "hoi", "hon", "hong", "hot", "hu", "hua", "huc", "hui", "hum", "hung", "huoc", "huong", "huot", "i", "ic", 
    "it", "khac", "khai", "kham", "khan", "khang", "khao", "khap", "khat", "khau", "khay", "khe", "khen", "kheo", 
    "khi", "khia", "khien", "khieu", "kho", "khoc", "khoi", "khon", "khong", "khot", "khu", "khua", "khuc", 
    "khui", "khum", "khung", "khuon", "khura", "khuong", "la", "lac", "lai", "lam", "lan", "lang", "lao", "lap", 
    "lat", "lau", "lay", "le", "len", "leo", "li", "lia", "liem", "lien", "liet", "lieu", "linh", "lo", "loc", 
    "loi", "lon", "long", "lot", "lu", "lua", "luc", "lui", "lum", "lung", "luoc", "luoi", "luon", "luong", 
    "luot", "ma", "mac", "mai", "mam", "man", "mang", "mao", "map", "mat", "mau", "may", "me", "men", "meo", 
    "mi", "mia", "mien", "mieu", "minh", "mo", "moc", "moi", "mon", "mong", "mot", "mu", "mua", "muc", "mui", 
    "mum", "mung", "muoc", "muoi", "muon", "muong", "na", "nac", "nai", "nam", "nan", "nang", "nao", "nap", 
    "nat", "nau", "nay", "ne", "nen", "neo", "ni", "nia", "niem", "nien", "nieu", "ninh", "no", "noc", 
    "noi", "non", "nong", "not", "nu", "nua", "nuc", "nui", "num", "nung", "nuoc", "nuoi", "nuon", "nuong", 
    "nga", "ngac", "ngai", "ngam", "ngan", "ngang", "ngao", "ngap", "ngat", "ngau", "ngay", "nge", "ngen", 
    "nghe", "nghen", "ngheo", "nghi", "nghia", "nghiem", "nghien", "nghiet", "nghieu", "nghinh", "ngo", "ngoc", 
    "ngoi", "ngon", "ngong", "ngot", "ngu", "ngua", "nguc", "ngui", "ngum", "ngung", "nguoc", "nguoi", "nguon", 
    "nguong", "nha", "nhac", "nhai", "nham", "nhan", "nhang", "nhao", "nhap", "nhat", "nhau", "nhay", "nhe", 
    "nhen", "nheo", "nhi", "nhia", "nhiem", "nhien", "nhiet", "nhieu", "nhinh", "nho", "nhoc", "nhoi", "nhon", 
    "nhong", "nhot", "nhu", "nhua", "nhuc", "nhui", "nhum", "nhung", "nhuoc", "nhuoi", "nhuon", "nhuong", "oa", 
    "oai", "oam", "oan", "oang", "oat", "oay", "oe", "oen", "oi", "om", "on", "ong", "ot", "pa", "pha", "phac", 
    "phai", "pham", "phan", "phang", "phao", "phap", "phat", "phau", "phay", "phe", "phen", "pheo", "phi", 
    "phia", "phiem", "phien", "phiet", "phieu", "phinh", "pho", "phoc", "phoi", "phon", "phong", "phot", "phu", 
    "phua", "phuc", "phui", "phum", "phung", "phuoc", "phuong", "phuot", "qua", "quac", "quai", "quam", "quan", 
    "quang", "quao", "quap", "quat", "quay", "que", "quen", "queo", "qui", "quie", "quien", "quieu", "quinh", 
    "quo", "quoc", "quoi", "quon", "quong", "ra", "rac", "rai", "ram", "ran", "rang", "rao", "rap", "rat", 
    "rau", "ray", "re", "ren", "reo", "ri", "ria", "rim", "rin", "rinh", "ro", "roc", "roi", "ron", "rong", 
    "rot", "ru", "rua", "ruc", "rui", "rum", "rung", "ruoc", "ruoi", "ruon", "ruong", "sa", "sac", "sai", 
    "sam", "san", "sang", "sao", "sap", "sat", "sau", "say", "se", "sen", "seo", "si", "sia", "siem", "sien", 
    "siet", "sieu", "sinh", "so", "soc", "soi", "son", "song", "sot", "su", "sua", "suc", "sui", "sum", 
    "sung", "suoc", "suoi", "suon", "suong", "ta", "tac", "tai", "tam", "tan", "tang", "tao", "tap", "tat", 
    "tau", "tay", "te", "ten", "teo", "tha", "thac", "thai", "tham", "than", "thang", "thao", "thap", "that", 
    "thau", "thay", "the", "then", "theo", "thi", "thia", "thiem", "thien", "thiet", "thieu", "thinh", "tho", 
    "thoc", "thoi", "thon", "thong", "thot", "thu", "thua", "thuc", "thui", "thum", "thung", "thuoc", "thuoi", 
    "thuon", "thuong", "thuot", "to", "toc", "toi", "ton", "tong", "tot", "tra", "trac", "trai", "tram", "tran", 
    "trang", "trao", "trap", "trat", "trau", "tray", "tre", "tren", "treo", "tri", "tria", "triem", "trien", 
    "triet", "trieu", "trinh", "tro", "troc", "troi", "tron", "trong", "trot", "tru", "trua", "truc", "trui", 
    "trum", "trung", "truoc", "truoi", "truon", "truong", "tu", "tua", "tuc", "tui", "tum", "tung", "tuoc", 
    "tuoi", "tuon", "tuong", "tuot", "va", "vac", "vai", "vam", "van", "vang", "vao", "vap", "vat", "vau", 
    "vay", "ve", "ven", "veo", "vi", "via", "viem", "vien", "viet", "vieu", "vinh", "vo", "voc", "voi", "von", 
    "vong", "vot", "vu", "vua", "vuc", "vui", "vum", "vung", "vuoc", "vuoi", "vuon", "vuong", "xa", "xac", 
    "xai", "xam", "xan", "xang", "xao", "xap", "xat", "xau", "xay", "xe", "xen", "xeo", "xi", "xia", "xiem", 
    "xien", "xiet", "xieu", "xinh", "xo", "xoc", "xoi", "xon", "xong", "xot", "xu", "xua", "xuc", "xui", 
    "xum", "xung", "xuoc", "xuoi", "xuon", "xuong", "y", "eu", "yeu"
}

DEFAULT_NON_SPEAKER_PHRASES = {
    "AND REMEMBER", "OFFICIAL DISTANCE", "GOOD NEWS FOR THEIR TEAMMATES", 
    "LL BE HONEST", "FIRST AND FOREMOST", "I SAID", "THE ONLY THING LEFT TO SETTLE", 
    "QUESTION IS", "FINALISTS", "WHISPERS", "SRT CONVERSION", 
    "WILL RED THRIVE OR WILL RED BE DEAD", "BUT REMEMBER", "THE RESULTS ARE IN", 
    "WE CHALLENGED", "I THINK", "IN THEIR DEFENSE", "THE PEAK OF HIS LIFE WAS DOING THE SPACETHING",
    "THE ROCKETS ARE BIGGER", "THE DISTANCE SHOULD BE FURTHER", "GET CRAFTY", "THAT WAS SO SICK",
    "OUT OF 100 CONTESTANTS", "THE FIRST ROUND IS BRUTAL", "YOU KNOW WHICH END GOES",
    "THE GAME IS ON", "THAT'S A GOOD THROW", "HE'S GOING FOR IT", "WE GOT THIS",
    "LAUNCH", "OH NO", "OH", "AH", "YEP", "WAIT", "YEAH", "WOO", "OKAY", "YES", "I ANH", "O BRI", "NG", "THE ONLY PROBLEM", "NOTE", "WARNING", "THINGS", "AND ON THE WAY WE CAME ACROSS THIS", 
    "THIS IS THE HIGHEST SWING IN EUROPE", "AND I SWEAR", "WHICH MEANT", "THE ONLY THING IS", 
    "HERE WE GO", "NEXT UP", "STEP 1", "STEP 2", "STEP 3", "AND STEP 3", "FIRST UP", 
    "SO THE QUESTION IS", "I WAS GROWING UP", "YOU MIGHT BE WONDERING", "UPDATE", 
    "NASHVILLE TO MIAMI", "ALL I KNOW IS", "UNLIKE JUDY", "THE GOOD NEWS IS", 
    "AER LINGUS SEAT", "THE TRUE TEST IS", "JUST AS I SUSPECTED", "LIKE I SAID", 
    "STAR REVIEW AND SAID", "I TOLD THEM ALL", "AND BEST OF ALL", "THE POINT IS", 
    "AMERICANS", "I WAS THINKING", "AND THEY GO", "FIRST OF ALL", "SECOND", 
    "ARE YOU LIKE", "AS A REMINDER", "ROUND 2", "ROUND 1", "ROUND 3", "ROUND 4", 
    "ROUND 5", "WELCOME TO ROUND 3", "THE QUESTION IS", "QUICK REMINDER", 
    "IN 2ND PLACE", "COMING UP", "FIRST STOP", "NEXT STEP", "AND THAT MEANS", 
    "HASHTAG", "SO TO BE CLEAR", "YOUR SECOND WORD", "WELCOME TO ROUND 6", 
    "BATTLE FINALE TIME", "NUMBER 1", "NUMBER 2", "BUT THE TRUTH IS", 
    "SCORE TO BEAT", "AND YOUR WINNER", "\"CRAFTY\" AND \"BETCHA\". COMING UP", 
    "NEXT ONE", "KEEP IN MIND", "AND IT SAYS", "YOU COULD SAY", "WELCOME TO ROUND 2", 
    "AND THE BEST PART", "ONTO ROUND 2", "THE RIDE WE CHOSE", "GOOD NEWS IS", 
    "BAD NEWS", "GOOD NEWS", "HE THOUGHT", "3 TEAMS REMAIN", "QUICK UPDATE", "DISTORTED", "MY FIRST QUESTION IS", "AND THE BEST PART IS", "BUT AS GORDON RAMSAY MIGHT SAY", "BUT THE GOOD NEWS IS", "LET ME JUST SAY", "BUT THE BEST PART", "I WILL SAY THOUGH", "LL SAY IS, UPDATE", "LL SAY IS", "UPDATE", "TO ALL OF YOU WATCHING WITH ME RIGHT NOW", "THIS IS THE LIFE", "I HAVE A QUESTION", "I WILL BE HONEST", "SO I CAN UPDATE MY INSTAGRAM BIO TO SAY"
}

NON_SPEAKER_PHRASES = DEFAULT_NON_SPEAKER_PHRASES

EXCEL_COLOR_PALETTE = [
    'background-color: #ADD8E6; color: #000000',
    'background-color: #90EE90; color: #000000',
    'background-color: #FFB6C1; color: #000000',
    'background-color: #FFFFE0; color: #000000',
    'background-color: #DDA0DD; color: #000000',
    'background-color: #AFEEEE; color: #000000',
    'background-color: #F0E68C; color: #000000',
    'background-color: #FFA07A; color: #000000',
    'background-color: #E0FFFF; color: #000000',
    'background-color: #F5F5DC; color: #000000',
    'background-color: #2F4F4F; color: #FFFFFF',
    'background-color: #191970; color: #FFFFFF',
    'background-color: #006400; color: #FFFFFF',
    'background-color: #800000; color: #FFFFFF',
    'background-color: #4B0082; color: #FFFFFF',
    'background-color: #556B2F; color: #FFFFFF',
    'background-color: #8B4513; color: #FFFFFF',
    'background-color: #36454F; color: #FFFFFF',
]

TIMECODE_REGEX = re.compile(r"^\d{2}:\d{2}:\d{2},\d{3}\s+-->\s+\d{2}:\d{2}:\d{2},\d{3}$")
SHORT_TIMECODE_REGEX = re.compile(r"^(?:\d{2}:)?\d{2}:\d{2}(?:[,.]\d{3})?\s*(?:-->\s*(?:\d{2}:)?\d{2}:\d{2}(?:[,.]\d{3})?)?$")
ENGLISH_WORD_REGEX = re.compile(r"\b[A-Za-z][A-Za-z0-9'-]*\b")
RED_COLOR = RGBColor(255, 0, 0)

# ==========================================
# HÀM ĐỌC/GHI JSON DATABASE
# ==========================================
def load_json_db(filepath, default_data=None):
    if os.path.exists(filepath):
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
                return set(data) if isinstance(data, list) and isinstance(default_data, set) else data
        except Exception: pass
    return default_data if default_data is not None else (set() if not isinstance(default_data, dict) else {})

def save_json_db(filepath, data_container):
    try:
        with open(filepath, "w", encoding="utf-8") as f:
            if isinstance(data_container, set): json.dump(list(data_container), f, ensure_ascii=False, indent=2)
            else: json.dump(data_container, f, ensure_ascii=False, indent=2)
    except Exception as e: st.error(f"Không thể lưu vào Database: {e}")

def hex_to_rgb(hex_str):
    if not hex_str: return None
    hex_str = str(hex_str).strip().lstrip('#')
    if len(hex_str) == 6:
        try: return tuple(int(hex_str[i:i+2], 16) for i in (0, 2, 4))
        except ValueError: return None
    return None

def generate_vibrant_rgb_colors_excluding(excluded_rgbs, count=200):
    colors = []
    used_set = set(tuple(x) for x in excluded_rgbs if x)
    attempts = 0
    while len(colors) < count and attempts < 10000:
        attempts += 1
        h = random.random(); s = 0.9; v = 0.8
        i = int(h * 6.0); f = h * 6.0 - i; p = v * (1.0 - s); q = v * (1.0 - s * f); t = v * (1.0 - s * (1.0 - f))
        if i % 6 == 0: r, g, b = v, t, p
        elif i % 6 == 1: r, g, b = q, v, p
        elif i % 6 == 2: r, g, b = p, v, t
        elif i % 6 == 3: r, g, b = p, q, v
        elif i % 6 == 4: r, g, b = t, p, v
        else: r, g, b = v, p, q
        r_int, g_int, b_int = int(r * 255), int(g * 255), int(b * 255)
        rgb_tuple = (r_int, g_int, b_int)
        if rgb_tuple not in used_set:
            used_set.add(rgb_tuple)
            colors.append(rgb_tuple)
    return colors

def generate_vibrant_rgb_colors(count=200):
    colors = set()
    while len(colors) < count:
        h = random.random(); s = 0.9; v = 0.8
        if s == 0.0: r = g = b = v
        else:
            i = int(h * 6.0); f = h * 6.0 - i; p = v * (1.0 - s); q = v * (1.0 - s * f); t = v * (1.0 - s * (1.0 - f))
            if i % 6 == 0: r, g, b = v, t, p
            elif i % 6 == 1: r, g, b = q, v, p
            elif i % 6 == 2: r, g, b = p, v, t
            elif i % 6 == 3: r, g, b = p, q, v
            elif i % 6 == 4: r, g, b = t, p, v
            else: r, g, b = v, p, q
        r, g, b = int(r * 255), int(g * 255), int(b * 255)
        colors.add((r, g, b))
    return list(colors)

FONT_COLORS_RGB_200 = generate_vibrant_rgb_colors(200)

def get_speaker_color_and_highlight(speaker_name, speaker_color_map, used_colors):
    spk_upper = speaker_name.strip().upper()
    fixed_colors = st.session_state.get('fixed_speaker_colors', DEFAULT_FIXED_SPEAKER_COLORS)
    
    if spk_upper in fixed_colors:
        cfg = fixed_colors[spk_upper]
        tc = cfg.get("text_color") if isinstance(cfg, dict) else cfg
        hc = cfg.get("highlight_color") if isinstance(cfg, dict) else None
        text_rgb = RGBColor(tc[0], tc[1], tc[2]) if tc and len(tc) >= 3 else None
        hl_rgb = RGBColor(hc[0], hc[1], hc[2]) if hc and len(hc) >= 3 else None
        return text_rgb, hl_rgb
    
    if spk_upper in speaker_color_map:
        res = speaker_color_map[spk_upper]
        if isinstance(res, tuple) and len(res) == 2:
            return res[0], res[1]
        elif isinstance(res, RGBColor):
            return res, None

    if used_colors: color_object = used_colors.pop()
    else:
        r, g, b = random.choice(FONT_COLORS_RGB_200)
        color_object = RGBColor(r, g, b)
        
    res = (color_object, None)
    speaker_color_map[spk_upper] = res
    return res[0], res[1]

def get_speaker_color(speaker_name, speaker_color_map, used_colors):
    spk_color, _ = get_speaker_color_and_highlight(speaker_name, speaker_color_map, used_colors)
    return spk_color if spk_color else RGBColor(0, 0, 0)

def apply_speaker_styling_to_run(run, text_color_tuple, highlight_color_tuple):
    if text_color_tuple:
        if isinstance(text_color_tuple, RGBColor):
            run.font.color.rgb = text_color_tuple
        elif isinstance(text_color_tuple, (tuple, list)) and len(text_color_tuple) >= 3:
            run.font.color.rgb = RGBColor(text_color_tuple[0], text_color_tuple[1], text_color_tuple[2])
            
    if highlight_color_tuple:
        if isinstance(highlight_color_tuple, (tuple, list, RGBColor)) and len(highlight_color_tuple) >= 3:
            hr, hg, hb = highlight_color_tuple[0], highlight_color_tuple[1], highlight_color_tuple[2]
            hex_fill = f"{hr:02X}{hg:02X}{hb:02X}"
            shd_xml = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_fill}"/>')
            run._r.get_or_add_rPr().append(shd_xml)

def clean_and_normalize_text(text, strip_all_tags=False, fix_punctuation=True, normalize_spaces=True, capitalize_first=True, remove_leading_dash=True):
    if not text or not isinstance(text, str): return ""
    res = text
    res = re.sub(r'\{\\[^}]*\}', '', res)
    res = re.sub(r'\\N', '\n', res, flags=re.IGNORECASE)
    
    if strip_all_tags:
        res = re.sub(r'<[^>]*>', '', res)
    else:
        res = re.sub(r'<(?!/?(i|b|u)\b)[^>]*>', '', res, flags=re.IGNORECASE)
        
    if remove_leading_dash:
        res = re.sub(r'^\s*[-–—]\s+(?=[A-Za-zÀ-ỹ0-9])', '', res)
        res = re.sub(r'(\n)\s*[-–—]\s+(?=[A-Za-zÀ-ỹ0-9])', r'\1', res)
        
    if fix_punctuation:
        res = re.sub(r'(?<!\()\s+([,!?:;\.\)])', r'\1', res)
        res = re.sub(r'([,!?:;])([A-Za-zÀ-ỹ0-9])', r'\1 \2', res)
        res = re.sub(r'(\.\.\.)([A-Za-zÀ-ỹ0-9])', r'\1 \2', res)
        res = re.sub(r'"\s*(.*?)\s*"', r'"\1"', res)
        
    if normalize_spaces:
        res = re.sub(r'[ \t]+', ' ', res)
        res = re.sub(r'\s*\n\s*', '\n', res)
        res = res.strip()
        
    if capitalize_first and res:
        lines = res.split('\n')
        cap_lines = []
        for line in lines:
            m = re.match(r'^([^:]+:\s*)([a-zà-ỹ])(.*)$', line)
            if m: line = m.group(1) + m.group(2).upper() + m.group(3)
            elif line and line[0].islower(): line = line[0].upper() + line[1:]
            cap_lines.append(line)
        res = '\n'.join(cap_lines)
        
    return res

def calculate_duration_sec(timecode_line):
    m = re.match(r"(\d{2}):(\d{2}):(\d{2})[,.](\d{3})\s+-->\s+(\d{2}):(\d{2}):(\d{2})[,.](\d{3})", timecode_line.strip())
    if not m: return 1.0, 0.0, 0.0
    h1, m1, s1, ms1, h2, m2, s2, ms2 = map(int, m.groups())
    t1 = h1 * 3600 + m1 * 60 + s1 + ms1 / 1000.0
    t2 = h2 * 3600 + m2 * 60 + s2 + ms2 / 1000.0
    dur = t2 - t1
    return dur if dur > 0 else 0.5, t1, t2

def timecode_to_sec(tc):
    m = re.match(r"(\d{2}):(\d{2}):(\d{2})[,.](\d{3})", str(tc).strip())
    if not m: return 0.0
    h, mins, s, ms = map(int, m.groups())
    return h * 3600 + mins * 60 + s + ms / 1000.0

def calculate_time_overlap(s1, e1, s2, e2):
    latest_start = max(s1, s2)
    earliest_end = min(e1, e2)
    return max(0.0, earliest_end - latest_start)

def is_valid_speaker_boundary(prefix, start_idx):
    if start_idx == 0: return True
    prev_char = prefix[start_idx - 1]
    curr_char = prefix[start_idx]
    if not prev_char.isalnum(): return True
    if prev_char.islower() and curr_char.isupper(): return True
    return False

def find_all_speaker_tags(text, custom_speakers=None, non_speakers=None):
    if custom_speakers is None: custom_speakers = st.session_state.get('custom_speakers', set())
    if non_speakers is None: non_speakers = st.session_state.get('custom_non_speakers', set())
    non_speakers_upper = {s.upper() for s in non_speakers}.union({s.upper() for s in DEFAULT_NON_SPEAKER_PHRASES})
    
    custom_names_upper = {s.upper(): s for s in custom_speakers if s.strip()}
    pattern = r"([A-Za-z0-9À-ỹ \t&\-\(\)\.\/]{1,45}):\s*"
    
    matches = []
    for m in re.finditer(pattern, text):
        col_pos = m.end() - 1
        prefix_to_colon = text[:col_pos]
        if (prefix_to_colon.count('(') - prefix_to_colon.count(')')) > 0: continue
        if (prefix_to_colon.count('[') - prefix_to_colon.count(']')) > 0: continue

        raw_prefix = m.group(1)
        extracted_spk = None
        
        for spk_upper, orig_spk in sorted(custom_names_upper.items(), key=lambda x: len(x[0]), reverse=True):
            if raw_prefix.upper().endswith(spk_upper):
                start_idx = len(raw_prefix) - len(spk_upper)
                if is_valid_speaker_boundary(raw_prefix, start_idx):
                    extracted_spk = orig_spk
                    spk_start = m.start(1) + start_idx
                    spk_end = m.end()
                    matches.append((spk_start, spk_end, extracted_spk, m.group(0)))
                    break
                    
        if extracted_spk: continue
            
        match_tail = re.search(r"(?:^|[^\w]|(?<=[a-zà-ỹ]))([A-ZÀ-Ỹ0-9][A-Za-z0-9À-ỹ \t&\-\(\)\.\/]{0,30})$", raw_prefix)
        if match_tail:
            cand = match_tail.group(1).strip(".,!?:;- ")
            if cand and len(cand) <= 35 and not cand.isdigit():
                if not (cand.startswith('(') or cand.endswith(')')):
                    if cand.upper() not in non_speakers_upper:
                        if len(cand.split()) <= 6:
                            start_idx = len(raw_prefix) - len(match_tail.group(1))
                            if is_valid_speaker_boundary(raw_prefix, start_idx):
                                spk_start = m.start(1) + start_idx
                                spk_end = m.end()
                                matches.append((spk_start, spk_end, cand, m.group(0)))

    matches.sort(key=lambda x: x[0])
    filtered_matches = []
    last_end = -1
    for start, end, spk, raw_m in matches:
        if start >= last_end:
            filtered_matches.append((start, end, spk, raw_m))
            last_end = end
            
    return filtered_matches

def is_valid_speaker_name(name):
    clean = name.strip()
    if not clean or len(clean) > 35 or clean.isdigit() or re.match(r'^\d+[\d\s:]*$', clean): return False
    if clean.startswith('(') or clean.endswith(')'): return False
    non_speakers_upper = st.session_state.get('custom_non_speakers', set())
    if clean.upper() in non_speakers_upper or clean.upper() in DEFAULT_NON_SPEAKER_PHRASES: return False
    if any(char in clean for char in ['?', '!', ',', '-->']): return False
    if len(clean.split()) > 6: return False
    return True

def is_valid_actor_name_strict(act_str):
    if not act_str: return False
    clean = act_str.strip().upper()
    if not clean or clean in DEFAULT_NON_SPEAKER_PHRASES: return False
    if any(char in clean for char in ['(', ')', '[', ']', '/', '\\', '-', '.']): return False
    if len(clean) > 25 or clean.isdigit(): return False
    return True

def get_paragraph_text_with_html(paragraph):
    text = ""
    for run in paragraph.runs:
        r_text = run.text
        if not r_text: continue
        if run.italic and not ("<i>" in r_text or "</i>" in r_text): text += f"<i>{r_text}</i>"
        else: text += r_text
    text = text.replace("</i><i>", "")
    return text

def round_seconds_to_int_minutes(total_sec):
    if total_sec <= 0: return 1
    mins = int(total_sec // 60); secs = int(round(total_sec % 60))
    if secs >= 30: mins += 1
    return max(1, mins)

def srt_timecode_to_ass(timecode_line):
    if not timecode_line or not isinstance(timecode_line, str): return None, None
    m = re.match(r"(\d{2}):(\d{2}):(\d{2})[,.](\d{3})\s+-->\s+(\d{2}):(\d{2}):(\d{2})[,.](\d{3})", timecode_line.strip())
    if not m: return None, None
    h1, m1, s1, ms1, h2, m2, s2, ms2 = m.groups()
    ass_start = f"{int(h1)}:{m1}:{s1}.{int(ms1)//10:02d}"
    ass_end = f"{int(h2)}:{m2}:{s2}.{int(ms2)//10:02d}"
    return ass_start, ass_end

def rgb_to_ass_hex(rgb_obj):
    if not rgb_obj: return "&H00FFFFFF&"
    try:
        r, g, b = rgb_obj[0], rgb_obj[1], rgb_obj[2]
        return f"&H00{b:02X}{g:02X}{r:02X}&"
    except Exception: return "&H00FFFFFF&"

def clean_file_name_for_output(original_filename, tag="_edit", ext=".docx"):
    name_without_ext = os.path.splitext(original_filename)[0]
    cleaned = re.sub(r'(CONVERTED_|FORMATTED_|\s*\(.*\)$|_edit$|_resync$|_final$)', '', name_without_ext, flags=re.IGNORECASE).strip()
    return f"{cleaned}{tag}{ext}"

def generate_actor_docx(video_title, actor_name, dialogue_list, font_size_pt=12):
    doc = Document()
    p_title = doc.add_paragraph(f"KỊCH BẢN THU ÂM - DIỄN VIÊN: {actor_name}")
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.runs[0].font.name = 'Times New Roman'; p_title.runs[0].font.size = Pt(16); p_title.runs[0].bold = True
    
    p_sub = doc.add_paragraph(f"Video: {video_title}")
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.runs[0].font.name = 'Times New Roman'; p_sub.runs[0].font.size = Pt(12); p_sub.runs[0].font.italic = True
    
    doc.add_paragraph()
    TAB_STOP = Inches(1.0)
    for item in dialogue_list:
        if item.get("timecode"):
            p_tc = doc.add_paragraph(item["timecode"])
            p_tc.runs[0].font.name = 'Times New Roman'; p_tc.runs[0].font.size = Pt(font_size_pt); p_tc.runs[0].bold = True
            p_tc.paragraph_format.space_before = Pt(0); p_tc.paragraph_format.space_after = Pt(0)
            
        p_line = doc.add_paragraph()
        p_line.paragraph_format.left_indent = TAB_STOP
        p_line.paragraph_format.first_line_indent = Inches(-1.0)
        p_line.paragraph_format.tab_stops.add_tab_stop(TAB_STOP, WD_TAB_ALIGNMENT.LEFT)
        
        r_spk = p_line.add_run(f"{item['speaker']}:")
        r_spk.font.name = 'Times New Roman'; r_spk.font.size = Pt(font_size_pt); r_spk.font.bold = True
        r_spk.font.color.rgb = RGBColor(79, 70, 229)
        p_line.add_run("\t")
        r_text = p_line.add_run(item['text'])
        r_text.font.name = 'Times New Roman'; r_text.font.size = Pt(font_size_pt)
        p_line.paragraph_format.space_before = Pt(0); p_line.paragraph_format.space_after = Pt(4)
        
    for p in doc.paragraphs: p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

def generate_actor_salary_slip_docx(actor_name, week_name, video_rows, total_pay, current_mode):
    doc = Document()
    p_title = doc.add_paragraph("MAI HAN STUDIO - PHIẾU BÁO CÁO THÙ LAO LỒNG TIẾNG")
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.runs[0].font.name = 'Times New Roman'; p_title.runs[0].font.size = Pt(16); p_title.runs[0].bold = True
    p_sub = doc.add_paragraph(f"Diễn viên Lồng tiếng: {actor_name} | {week_name}")
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.runs[0].font.name = 'Times New Roman'; p_sub.runs[0].font.size = Pt(12); p_sub.runs[0].font.italic = True
    
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=5); table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells; hdr_names = ["Stt", "Tiêu đề video", "Khối lượng", "Đơn giá áp dụng", "Thành tiền"]
    for i, name in enumerate(hdr_names):
        hdr_cells[i].text = name; hdr_cells[i].paragraphs[0].runs[0].font.bold = True; hdr_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        
    for r in video_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(r["Stt"]); row_cells[1].text = str(r["Tiêu đề video"]); row_cells[2].text = str(r["Thời lượng"])
        row_cells[3].text = str(r["Đơn giá"]); row_cells[4].text = str(r["Thành tiền"])
        for c in row_cells:
            for p in c.paragraphs:
                for run in p.runs: run.font.name = 'Times New Roman'; run.font.size = Pt(11)

    doc.add_paragraph()
    p_total = doc.add_paragraph(f"👉 TỔNG THÙ LAO THANH TOÁN: {total_pay:,.0f} VNĐ")
    p_total.runs[0].font.name = 'Times New Roman'; p_total.runs[0].font.size = Pt(13); p_total.runs[0].bold = True
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

def clean_dialogue_text_for_excel(text):
    text = clean_and_normalize_text(text, strip_all_tags=True)
    return text

def parse_srt_to_dataframe(srt_content, custom_speakers=None, non_speakers=None, default_speaker="Unknown"):
    if custom_speakers is None: custom_speakers = st.session_state.get('custom_speakers', set())
    if non_speakers is None: non_speakers = st.session_state.get('custom_non_speakers', set())

    fallback_spk = default_speaker.strip() if default_speaker and default_speaker.strip() else "Unknown"

    data = []
    blocks = re.split(r'\n\s*\n', srt_content.strip())
    last_known_speaker = fallback_spk
    last_is_explicit = False

    for block in blocks:
        lines = [l.strip() for l in block.strip().split('\n') if l.strip()]
        if len(lines) < 2: continue

        time_line = ""; time_idx = -1
        for idx, line in enumerate(lines[:2]):
            if "-->" in line: time_line = line; time_idx = idx; break
        if not time_line: continue

        time_match = re.match(r'(\d{2}:\d{2}:\d{2}[,.]\d{3})\s*-->\s*(\d{2}:\d{2}:\d{2}[,.]\d{3})', time_line)
        if not time_match: continue

        time_start = time_match.group(1).replace('.', ',')
        time_end = time_match.group(2).replace('.', ',')

        dialogue_lines = lines[time_idx + 1:]
        if not dialogue_lines: continue

        block_text = "\n".join(dialogue_lines)
        speaker_tags = find_all_speaker_tags(block_text, custom_speakers, non_speakers)

        if not speaker_tags:
            clean_text = clean_dialogue_text_for_excel(block_text)
            if clean_text: data.append([time_start, time_end, last_known_speaker, clean_text, last_is_explicit])
        else:
            last_idx = 0
            for i, (start_pos, end_pos, spk_name, raw_m) in enumerate(speaker_tags):
                leading_text = block_text[last_idx:start_pos].strip()
                clean_leading = clean_dialogue_text_for_excel(leading_text)
                if clean_leading: data.append([time_start, time_end, last_known_speaker, clean_leading, last_is_explicit])

                next_start = speaker_tags[i+1][0] if i + 1 < len(speaker_tags) else len(block_text)
                segment_text = block_text[end_pos:next_start]
                clean_seg = clean_dialogue_text_for_excel(segment_text)
                if clean_seg:
                    data.append([time_start, time_end, spk_name, clean_seg, True])
                    last_known_speaker = spk_name
                    last_is_explicit = True
                else:
                    last_known_speaker = spk_name
                    last_is_explicit = True
                last_idx = next_start

    return pd.DataFrame(data, columns=['Start', 'End', 'Speaker', 'Dialogue', 'Is_Explicit'])

def check_resync_integrity(body_zone, srt_dialogues):
    input_tcs = []
    input_dialogues = []
    curr_tc = None
    for text in body_zone:
        if TIMECODE_REGEX.match(text) or SHORT_TIMECODE_REGEX.match(text):
            curr_tc = text; input_tcs.append(text)
        elif curr_tc and text and not text.lower().startswith("srt conversion") and not text.lower().startswith("vai:"):
            clean_txt = re.sub(r'</?[ibuIBU]>', '', text).strip()
            if clean_txt: input_dialogues.append({"timecode": curr_tc, "text": clean_txt})
                
    output_tcs = []
    output_dialogues = []
    for srt_block in srt_dialogues:
        lines = srt_block.strip().split('\n')
        if len(lines) >= 3:
            tc = lines[1]; txt = "\n".join(lines[2:])
            output_tcs.append(tc); output_dialogues.append({"timecode": tc, "text": txt})
            
    tc_in_cnt = len(input_tcs); tc_out_cnt = len(output_tcs)
    diff_issues = []
    
    if tc_in_cnt != tc_out_cnt:
        diff_issues.append({
            "Loại Cảnh Báo": "🔴 Mất Timecode",
            "Mốc Thời Gian": "Toàn file",
            "Chi Tiết So Sánh": f"File Edit có {tc_in_cnt} mốc timecode nhưng File Final có {tc_out_cnt} mốc."
        })
        
    min_len = min(len(input_dialogues), len(output_dialogues))
    for idx in range(min_len):
        in_item = input_dialogues[idx]; out_item = output_dialogues[idx]
        if in_item["timecode"] != out_item["timecode"]:
            diff_issues.append({
                "Loại Cảnh Báo": "🔵 Lệch Mốc Timecode",
                "Mốc Thời Gian": in_item["timecode"],
                "Chi Tiết So Sánh": f"File Edit: '{in_item['timecode']}' vs Final: '{out_item['timecode']}'"
            })
            break
        
        in_len = len(in_item["text"]); out_len = len(out_item["text"])
        if in_len > 10 and out_len < in_len * 0.7:
            diff_issues.append({
                "Loại Cảnh Báo": "⚠️ Sụt Giảm Ký Tự / Mất Chữ",
                "Mốc Thời Gian": in_item["timecode"],
                "Chi Tiết So Sánh": f"Gốc ({in_len} ký tự) -> Final ({out_len} ký tự): '{out_item['text'][:40]}...'"
            })

    return {
        "tc_in_cnt": tc_in_cnt,
        "tc_out_cnt": tc_out_cnt,
        "line_in_cnt": len(input_dialogues),
        "line_out_cnt": len(output_dialogues),
        "diff_issues": diff_issues
    }

def process_docx(uploaded_file, file_name_without_ext, enable_colors, enable_phonetic, enable_cast, is_resync=False, font_size_pt=12):
    speaker_color_map = {}; used_colors = [RGBColor(r, g, b) for r, g, b in FONT_COLORS_RGB_200]
    random.shuffle(used_colors); stats_counter = Counter(); seen_speakers_first_time = set()
    actor_dialogue_map = {}; qc_warnings = []        
    
    custom_speakers = st.session_state.get('custom_speakers', set())
    custom_non_speakers = st.session_state.get('custom_non_speakers', set())

    original_document = Document(io.BytesIO(uploaded_file.getvalue()))
    raw_paragraphs = [p for p in original_document.paragraphs]
    processed_strings = preprocess_raw_paragraphs(raw_paragraphs, custom_speakers, custom_non_speakers)
    
    first_timecode_idx = 0
    for idx, text in enumerate(processed_strings):
        if TIMECODE_REGEX.match(text) or SHORT_TIMECODE_REGEX.match(text): 
            first_timecode_idx = idx; break
            
    header_zone = processed_strings[:first_timecode_idx] if first_timecode_idx > 0 else []
    body_zone = processed_strings[first_timecode_idx:] if first_timecode_idx > 0 else processed_strings
    
    for h_line in header_zone:
        if ":" in h_line and not h_line.lower().startswith("srt conversion"):
            parts = h_line.split(":", 1)
            spk_k = parts[0].strip().upper(); act_v = parts[1].strip().upper()
            if is_valid_speaker_name(spk_k) and is_valid_actor_name_strict(act_v):
                st.session_state['custom_cast_mapping'][spk_k] = act_v

    document = Document()
    title_text_raw = file_name_without_ext.upper()
    title_text = title_text_raw.replace("CONVERTED_", "").replace("FORMATTED_", "").replace("_EDIT", "").replace("_RESYNC", "").replace("_FINAL", "").replace(" (GỐC)", "").strip()
    title_paragraph = document.add_paragraph(title_text)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.runs[0].font.name = 'Times New Roman'; title_paragraph.runs[0].font.size = Pt(20); title_paragraph.runs[0].bold = True
    
    unique_speakers = []; assigned_actors = []; unassigned_speakers = []
    for text in body_zone:
        if not text or text.lower().startswith("srt conversion") or text.lower().startswith("vai:"): continue 
        spk_tags = find_all_speaker_tags(text, custom_speakers, custom_non_speakers)
        for _, _, speaker_name, _ in spk_tags:
            if speaker_name not in unique_speakers:
                unique_speakers.append(speaker_name)
                act = st.session_state['custom_cast_mapping'].get(speaker_name.upper(), "").strip().upper()
                if act and is_valid_actor_name_strict(act):
                    if act not in assigned_actors: assigned_actors.append(act)
                else:
                    if speaker_name.upper() != "ALL": unassigned_speakers.append(speaker_name)
                    
    if unassigned_speakers: qc_warnings.append(f"⚠️ Phát hiện {len(unassigned_speakers)} nhân vật chưa gán diễn viên: {', '.join(unassigned_speakers)}")
            
    if unique_speakers:
        if enable_cast:
            header_vai = document.add_paragraph()
            r_vai = header_vai.add_run("VAI: ")
            r_vai.font.name = 'Times New Roman'; r_vai.font.size = Pt(font_size_pt); r_vai.font.bold = True
            header_vai.paragraph_format.space_before = Pt(0); header_vai.paragraph_format.space_after = Pt(0)

            for spk in unique_speakers:
                p_spk = document.add_paragraph()
                p_spk.paragraph_format.space_before = Pt(0); p_spk.paragraph_format.space_after = Pt(0)
                is_all = (spk.strip().upper() == "ALL")
                spk_color, spk_hl = get_speaker_color_and_highlight(spk, speaker_color_map, used_colors)
                
                r_spk_name = p_spk.add_run(f"{spk}: ")
                r_spk_name.font.name = 'Times New Roman'; r_spk_name.font.size = Pt(font_size_pt); r_spk_name.font.bold = True
                
                if is_all:
                    r_spk_name.font.color.rgb = RED_COLOR; r_spk_name.font.highlight_color = WD_COLOR_INDEX.YELLOW
                elif enable_colors:
                    if spk_color: r_spk_name.font.color.rgb = spk_color
                    if spk_hl:
                        apply_speaker_styling_to_run(r_spk_name, (spk_color[0], spk_color[1], spk_color[2]) if spk_color else None, (spk_hl[0], spk_hl[1], spk_hl[2]))
                    
                actor = st.session_state['custom_cast_mapping'].get(spk.upper(), "").strip().upper()
                if actor and not is_all and is_valid_actor_name_strict(actor):
                    r_actor = p_spk.add_run(actor)
                    r_actor.font.name = 'Times New Roman'; r_actor.font.size = Pt(font_size_pt); r_actor.font.bold = True
                    r_actor.font.color.rgb = RED_COLOR
        else:
            speaker_list_text = "VAI: " + ", ".join(unique_speakers)
            p = document.add_paragraph(speaker_list_text)
            p.runs[0].font.name = 'Times New Roman'; p.runs[0].font.size = Pt(font_size_pt); p.runs[0].bold = True
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(6)
    
    document.add_paragraph()
    start_index = len(document.paragraphs); total_paras = len(body_zone)
    progress_bar = st.progress(0); status_text = st.empty()
    ass_dialogues = []; srt_dialogues = []; current_timecode_line = None; srt_counter = 1; max_video_time_sec = 0.0

    for idx, text in enumerate(body_zone):
        if idx % max(1, total_paras // 10) == 0:
            progress = int((idx / total_paras) * 100)
            progress_bar.progress(progress)
            status_text.text(f"Đang phân tích & xử lý {idx}/{total_paras}...")

        if not text or text.upper() == title_text.upper(): continue
        if text.lower().startswith("srt conversion") or text.lower().startswith("vai:") or re.fullmatch(r"^\s*\d+\s*$", text): continue
            
        if TIMECODE_REGEX.match(text) or SHORT_TIMECODE_REGEX.match(text):
            current_timecode_line = text
            dur, t1, t2 = calculate_duration_sec(text)
            if t2 > max_video_time_sec: max_video_time_sec = t2
            new_paragraph = document.add_paragraph(text)
            new_paragraph.runs[0].font.bold = True; new_paragraph.runs[0].font.name = 'Times New Roman'; new_paragraph.runs[0].font.size = Pt(font_size_pt)
            new_paragraph.paragraph_format.space_before = Pt(0); new_paragraph.paragraph_format.space_after = Pt(0)
        else:
            cleaned_text = clean_and_normalize_text(text, strip_all_tags=False)
            if is_resync: cleaned_text = normalize_phonetics_in_text(cleaned_text)
            
            ass_formatted_line, pure_dialogue_text = format_and_split_dialogue(
                document, cleaned_text, enable_colors, enable_phonetic, enable_cast, 
                speaker_color_map, used_colors, stats_counter, seen_speakers_first_time,
                actor_dialogue_map, current_timecode_line, custom_speakers, custom_non_speakers,
                font_size_pt=font_size_pt
            )
            
            if current_timecode_line and pure_dialogue_text:
                dur, _, _ = calculate_duration_sec(current_timecode_line)
                clean_chars = len(re.sub(r'</?[ibuIBU]>', '', pure_dialogue_text).strip())
                cps = clean_chars / dur if dur > 0 else 0
                if cps > 20 and clean_chars > 0: 
                    qc_warnings.append(f"⏱️ **Tốc độ đọc nhanh ({cps:.1f} ký tự/s)** tại `{current_timecode_line}`: \"{pure_dialogue_text[:45]}...\"")
            
            if current_timecode_line and ass_formatted_line:
                start_ass, end_ass = srt_timecode_to_ass(current_timecode_line)
                if start_ass and end_ass:
                    ass_dialogues.append(f"Dialogue: 0,{start_ass},{end_ass},Default,,0,0,0,,{ass_formatted_line}")
                    srt_dialogues.append(f"{srt_counter}\n{current_timecode_line}\n{re.sub(r'{\\.*?}', '', ass_formatted_line)}\n")
                    srt_counter += 1

    progress_bar.progress(100); status_text.text("Xử lý hoàn tất!"); time.sleep(0.5)
    progress_bar.empty(); status_text.empty()
            
    for paragraph in document.paragraphs[start_index:]:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.space_before = Pt(0); paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            if run.font.size is None or is_resync: run.font.size = Pt(font_size_pt)
        
    docx_file = io.BytesIO(); document.save(docx_file); docx_file.seek(0)
    
    ass_header = f"""[Script Info]
Title: {title_text}
ScriptType: v4.00+
WrapStyle: 0
ScaledBorderAndShadow: yes
YCbCr Matrix: None
PlayResX: 1920
PlayResY: 1080

[V4+ Styles]
Format: Name, Fontname, Fontsize, PrimaryColour, SecondaryColour, OutlineColour, BackColour, Bold, Italic, Underline, StrikeOut, ScaleX, ScaleY, Spacing, Angle, BorderStyle, Outline, Shadow, Alignment, MarginL, MarginR, MarginV, Encoding
Style: Default,Times New Roman,45,&H00FFFFFF,&H000000FF,&H00000000,&H80000000,0,0,0,0,100,100,0,0,1,2,1,2,10,10,30,1

[Events]
Format: Layer, Start, End, Style, Name, MarginL, MarginR, MarginV, Effect, Text
"""
    ass_content = ass_header + "\n".join(ass_dialogues)
    ass_file = io.BytesIO(ass_content.encode('utf-8'))
    
    srt_content = "\n".join(srt_dialogues)
    srt_file = io.BytesIO(srt_content.encode('utf-8-sig'))
    
    actor_zip_bytes = io.BytesIO()
    with zipfile.ZipFile(actor_zip_bytes, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for act_name, dialogues in actor_dialogue_map.items():
            act_buf = generate_actor_docx(title_text, act_name, dialogues, font_size_pt=font_size_pt)
            zip_file.writestr(f"Kich_Ban_{act_name}.docx", act_buf.getvalue())
    actor_zip_bytes.seek(0)
    
    actor_stats_breakdown = {}
    for act_name, dialogues in actor_dialogue_map.items():
        line_cnt = len(dialogues)
        word_cnt = sum(len(re.sub(r'</?[ibuIBU]>', '', d['text']).split()) for d in dialogues)
        actor_stats_breakdown[act_name] = {"lines": line_cnt, "words": word_cnt}
    
    video_duration_min = round_seconds_to_int_minutes(max_video_time_sec)
    
    # Tính toán báo cáo so sánh độ toàn vẹn (Edit vs Final)
    integrity_report = check_resync_integrity(body_zone, srt_dialogues)
    
    stats = {
        "total_speakers": len(unique_speakers),
        "total_lines": sum(stats_counter.values()),
        "top_speaker": stats_counter.most_common(1)[0] if stats_counter else ("Không có", 0),
        "actors_list": assigned_actors,
        "actor_dialogue_map": actor_dialogue_map,
        "actor_stats_breakdown": actor_stats_breakdown,
        "video_duration_min": video_duration_min,
        "video_title": title_text,
        "qc_warnings": qc_warnings,
        "integrity_report": integrity_report
    }
    
    return docx_file, ass_file, srt_file, actor_zip_bytes, stats
